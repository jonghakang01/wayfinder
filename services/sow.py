import calendar
import io
import json
import os
import re
import tempfile
import urllib.error
import urllib.request
import uuid
from datetime import date, datetime, timedelta

from ._paths import DATA_ROOT

META = {
    "name": "Deal Desk",
    "path": "/sow",
    "icon": "🤝",
    "description": "Statement of Work drafting & payment schedule automation",
    "hidden": False,
    "admin_only": True,
}

CHEIL_ENTITY = "Cheil USA, Inc."
SAMSUNG_ENTITY = "Samsung Electronics America, Inc."

# Document types drill down: direction → document kind (SOW / MSA / NDA).
# One SOW template per direction (2026-07-22, 강프로: team/individual merged) —
# the rate model (monthly vs hourly resource table) is a toggle inside the
# editor instead of a separate document type.
TYPES = {
    "sea_sow": {
        "dir": "samsung", "kind": "sow", "mode": "hourly", "icon": "🤝",
        "label": "Statement of Work",
        "desc": "SOW under the Advertising Services Agreement — resources billed hourly or at monthly rates (pick inside).",
    },
    "agy_sow": {
        "dir": "agency", "kind": "sow", "mode": "hourly", "icon": "🤝",
        "label": "Statement of Work",
        "desc": "Vendor SOW under the vendor's MSA — resources billed hourly or at monthly rates (pick inside).",
    },
    "sea_est": {
        "dir": "samsung", "kind": "est", "mode": None, "icon": "🧮",
        "label": "Cost Estimation",
        "desc": "Pick people from the roster, set allocation and months — billing (and internal cost) computes itself. Exports the estimation xlsx.",
    },
    "agy_msa": {
        "dir": "agency", "kind": "msa", "mode": None, "icon": "📜",
        "label": "Master Services Agreement",
        "desc": "The standard Cheil MSA — pick the vendor and effective date; the full legal text exports as-is.",
    },
    "agy_nda": {
        "dir": "agency", "kind": "nda", "mode": None, "icon": "🤐",
        "label": "One-Way NDA",
        "desc": "Confidentiality & nondisclosure agreement signed before vendor talks — vendor + date fill-in.",
    },
}

_ASSETS = os.path.join(os.path.dirname(__file__), "sow_assets")

# Full executed documents rendered beside the template (2026-07-22 강프로:
# the whole Word document, title page through signatures, with the fill-in
# fields highlighted so drafting maps 1:1 to the editor's slots).
SAMPLES = {
    "sea_sow": {
        "src": "Samsung_ Data Engineer.docx (executed Dec 2025)",
        "title": "Data Engineer # 1",
        "date": "Dec 7, 2025", "project": "SEA eCom Data",
        "by": "Jongha Kang", "for": "Nanda Kumar",
        "summary": ("This Statement of Work outlines the provision of Data Engineering "
                    "support focused on building and maintaining the data infrastructure "
                    "that enables analytics, and reporting initiatives across organizations. "
                    "The role is centered on developing reliable, scalable, and high-quality "
                    "data pipelines and systems, while contributing to data platform "
                    "evolution and cross-functional data enablement."),
        "deliverables": [
            "Design, build, and maintain scalable data pipelines to ingest, transform, and process data from multiple sources",
            "Develop and manage ETL/ELT workflows to ensure timely, accurate, and efficient data availability for analytics and personalization use cases.",
            "Model and structure data for consumption by downstream systems, including BI tools, experimentation platforms, and marketing systems.",
            "Ensure data quality, integrity, and reliability through validation checks, monitoring frameworks, and automated alerting mechanisms.",
            "Partner with Data Product Managers, analysts, and business stakeholders to understand data requirements and translate them into scalable engineering solutions.",
            "Optimize data pipelines and queries for performance, cost efficiency, and scalability across growing data volumes.",
            "Manage and maintain data warehouse and/or lake environments, ensuring proper partitioning, indexing, and storage optimization.",
            "Implement and enforce data governance best practices, including schema management, access controls, and documentation standards.",
            "Support real-time and near-real-time data processing needs where required for personalization and campaign activation.",
            "Collaborate with engineering teams to improve data instrumentation, event tracking, and logging across digital platforms.",
            "Troubleshoot data issues, perform root cause analysis, and implement long-term fixes to prevent recurrence.",
            "Build reusable data frameworks, pipeline templates, and tooling to accelerate development and standardize engineering practices.",
            "Maintain clear documentation of data pipelines, schemas, transformations, and dependencies to support transparency and maintainability.",
        ],
        "stk": [["Durga R", "Jongha Kang"], ["durga.r@samsung.com", "jongha.kang@cheil.com"],
                ["Chennai, India", "Mountain View, USA"]],
        "start": "March 23, 2026", "end": "December 31, 2026",
        "res_mode": "hourly",
        "res_rows": [["Data Engineer", "India", "1", "9.3", "$25", "168", "$39,060"]],
        "fee": "$39,060",
        "schedule": [
            ["Mar-26", "$1,260", "1-Apr-26"], ["Apr-26", "$4,200", "1-May-26"],
            ["May-26", "$4,200", "1-Jun-26"], ["Jun-26", "$4,200", "1-Jul-26"],
            ["Jul-26", "$4,200", "1-Aug-26"], ["Aug-26", "$4,200", "1-Sep-26"],
            ["Sep-26", "$4,200", "1-Oct-26"], ["Oct-26", "$4,200", "1-Nov-26"],
            ["Nov-26", "$4,200", "1-Dec-26"], ["Dec-26", "$4,200", "1-Jan-27"],
        ],
    },
    "agy_sow": {
        "src": "eComm Sr. Data Analyst — Cheil-Invictus SOW (executed Jan 2025)",
        "title": "Sr. Data Analyst (Corp Marketing Dashboard Support Scope)",
        "date": "Jan 17, 2025", "project": "Corp Marketing Dashboard Support",
        "by": "", "for": "",
        "vendor_name": "Invictus Data, Inc.",
        "vendor_entity": ("Invictus Data Inc, with its principal place of business located "
                          "at 675 Shady Creek Ln, Los Altos, CA – 94024"),
        "msa_date": "September 28, 2023",
        "summary": ("A skilled Sr. Data Analyst for Bigdata platform to work as a part of a "
                    "growing team to provide business insights, analyze trends, build "
                    "dashboards, data mining using SQL, and other similar technologies."),
        "deliverables": [
            "Analyze business performance, pinpoint key challenges, and present insights using clear and concise visualizations (charts, graphs, tables, or summaries).",
            "Drive innovation by leveraging data to generate insights, develop business cases, and create scalable solutions that foster business growth.",
            "Collaborate with Category/Product Managers to guide product and business decisions through data-driven insights.",
            "Provide data analysis and support for Ecommerce Operations and Trade-in strategies, ensuring alignment with business goals.",
            "Take ownership of a key business area from a data perspective, conducting deep dives and delivering actionable insights.",
            "Partner with product managers, business team members, and engineers to implement accurate tracking and tagging for critical business metrics.",
            "Develop and refine processes to test, learn, and iterate, accelerating growth through continuous improvement.",
            "Extract actionable insights from data using SQL, Spark, Hive, and Tableau, summarizing findings for leadership teams.",
            "Design and maintain Tableau dashboards to meet business needs, while working with engineering teams to build data solutions.",
            "Regularly audit dashboards and business metrics to identify trends, discrepancies, or issues with data pipelines.",
            "Provide ad-hoc analytics and reporting support for executive presentations and decision-making.",
        ],
        "stk": [["Pranav Vishwanathan", ""], ["", ""], ["", ""]],
        "start": "December 1, 2024", "end": "December 31, 2025",
        "res_mode": "hourly",
        "res_rows": [["Pranav Vishwanathan", "Sr. Data Analyst", "1", "13", "$32", "168", "$69,888"]],
        "fee": "$69,888",
        "schedule": [
            ["Dec-24", "$5,824", "1-Jan-25"], ["Jan-25", "$5,824", "1-Feb-25"],
            ["Feb-25", "$5,824", "1-Mar-25"], ["Mar-25", "$5,824", "1-Apr-25"],
            ["Apr-25", "$5,824", "1-May-25"], ["May-25", "$5,824", "1-Jun-25"],
            ["Jun-25", "$5,824", "1-Jul-25"], ["Jul-25", "$5,824", "1-Aug-25"],
            ["Aug-25", "$5,824", "1-Sep-25"], ["Sep-25", "$5,824", "1-Oct-25"],
            ["Oct-25", "$5,824", "1-Nov-25"], ["Nov, Dec-25", "$5,824", "1-Dec-25"],
        ],
    },
    "agy_msa": {
        "src": "MSA Cheil-AIEnterprise (executed May 31, 2024)",
        "note": ("Executed fill-in: Effective Date = May 31, 2024 · Contractor = AIEnterprise Inc. "
                 "The template on the left IS the full executed document — the highlighted date and "
                 "vendor fields were the only vendor-specific edits."),
    },
    "agy_nda": {
        "src": "Cheil NY Vendor One-Way NDA (executed May 29, 2024)",
        "note": ("Executed fill-in: Effective Date = 05/29/2024 · Vendor = AIENTERPRISE INC. "
                 "(signed by Sudhanshu Mohan, CEO). The template on the left is the full executed "
                 "text — only the highlighted date and vendor name vary."),
    },
}

# One-way NDA text lifted from the executed "Cheil NY Vendor One-Way NDA" (slots
# for effective date and vendor name). Kept verbatim, including original quirks.
NDA_TITLE = "CONFIDENTIALITY AND NONDISCLOSURE AGREEMENT"
NDA_INTRO = (
    'This CONFIDENTIALITY AND NONDISCLOSURE AGREEMENT (the "Agreement"), is '
    'entered into as of {date} (the "Effective Date"), by and between {vendor} '
    '(the "Vendor"), and Cheil USA, Inc. (the "Cheil").'
)
NDA_BODY = [
    "WHEREAS, the Vendor and Cheil are engaged in, or may enter into, talks regarding "
    "a potential business relationship, and the Vendor understands that Cheil has "
    "disclosed or may disclose to the Vendor certain confidential and proprietary "
    "information which has commercial and other value in the business of Cheil.",
    "NOW, THEREFORE, in consideration of the foregoing, and the mutual covenants, "
    "terms and conditions set forth herein, and other good and valuable consideration, "
    "the receipt and sufficiency of which are hereby acknowledged, the parties hereto "
    "hereby agree as follows.",
    "1. For purposes of this Agreement, “Confidential Information” shall mean all "
    "technical and business information relating to Cheil’s products, clients, "
    "technology, software, processes, methods, services, research and development, "
    "pricing, future business plans and all other information of Cheil or its clients "
    "which may be disclosed by Cheil or to which the Vendor may be provided access by "
    "Cheil in accordance with this Agreement.",
    "2. Confidential Information shall not include any information that: (i) is or "
    "becomes (through no improper action or inaction by the Vendor) generally "
    "available to the public; (ii) was in its possession or known by it prior to "
    "receipt from Cheil; (iii) was rightfully disclosed to him/her by a third party "
    "without a breach of any confidentiality obligations; or (iv) was independently "
    "developed by the Vendor without reference to any Confidential Information.",
    "3. The Vendor agrees: (i) to hold the Confidential Information in confidence and "
    "to take all reasonable precautions to protect such Confidential Information; "
    "(ii) not to disclose any such Confidential Information or any information derived "
    "therefrom to any third person; and (iii) not to make any use whatsoever at any "
    "time of such Confidential Information except for the limited and sole internal "
    "business purposes for which is has been disclosed by Cheil.  Any employee given "
    "access to any such Confidential Information by the Vendor must have a legitimate "
    "“need to know” such Confidential Information.  The Vendor is liable for all acts "
    "and omissions of third parties to whom he/she discloses Confidential Information.  "
    "Further, the Vendor may make disclosures required by valid order of any court or "
    "other authorized governmental entity, provided the Vendor promptly notifies "
    "Cheil, uses reasonable efforts to limit disclosure and assists Cheil, at Cheil's "
    "expense, to obtain confidential treatment or a protective order for such "
    "Confidential Information.  All Confidential Information is provided “AS IS” and "
    "without any warranties, express, implied or otherwise, and no warranty is made "
    "regarding its accuracy or completeness.  The Vendor shall not reverse engineer, "
    "decompile, translate, adapt, or disassemble any software of the other party, or "
    "attempt to make derivative works from such software.  No licenses or rights under "
    "any patent, copyright, trademark or trade secret are granted, or are to be "
    "implied, by this Agreement.  The Confidential Information shall remain the sole "
    "property of Cheil and the Vendor shall not challenge or contest Cheil’s right to "
    "own and use the Confidential Information or other intellectual property.",
    "4. Immediately upon a request by Cheil at any time, the Vendor will turn over to "
    "Cheil all Confidential Information and all documents or media containing any such "
    "Confidential Information and any and all copies or extracts thereof.  The Vendor "
    "understands that nothing herein: (i) requires the disclosure of any Confidential "
    "Information by Cheil, which shall be disclosed, if at all, solely at the option "
    "of Cheil; or (ii) requires Cheil to proceed with any proposed transaction or "
    "other business relationship in connection with which Confidential Information "
    "may be disclosed.",
    "5. The Vendor acknowledges and agrees that due to the unique nature of the "
    "Confidential Information, there can be no adequate remedy at law for any breach "
    "of its obligations hereunder, that any such breach or any unauthorized use or "
    "release of any Confidential Information will allow the Vendor or third parties "
    "to unfairly compete with Cheil resulting in irreparable harm to Cheil and "
    "therefore, that upon any such breach or any threat thereof, Cheil shall be "
    "entitled to appropriate equitable relief in addition to whatever remedies it "
    "might have at law and to be indemnified by the Vendor from any loss or harm, "
    "including, without limitation, reasonable attorney’s fees and expenses, in "
    "connection with any breach or enforcement of the Vendor’s obligations hereunder "
    "or the unauthorized use or release of any such Confidential Information.  The "
    "Vendor will notify Cheil in writing immediately upon becoming aware of the "
    "occurrence of any such unauthorized release or other breach of confidentiality "
    "obligations hereunder.",
    "6. Neither party shall have the right to assign its rights or obligations under "
    "this Agreement, whether expressly or by operation of law, without the prior "
    "written consent of the other party.  This Agreement shall be binding on, and "
    "inure to the benefit of, each party and their permitted successors and assigns.",
    "7. This Agreement shall be governed by the laws of the State of New York, and "
    "each party irrevocably submits to the exclusive jurisdiction of the courts "
    "located in New York County, New York.  In the event that any of the provisions "
    "of this Agreement shall be held by a court or other tribunal of competent "
    "jurisdiction to be illegal, invalid or unenforceable, such provisions shall be "
    "limited or eliminated to the minimum extent necessary so that this Agreement "
    "shall otherwise remain in full force and effect.  This Agreement supersedes all "
    "prior discussions and writings and constitutes the entire agreement between the "
    "parties with respect to the limited subject matter set forth herein.  No waiver "
    "or modification of this Agreement will be binding upon either party unless made "
    "in writing and signed by a duly authorized representative of such party and no "
    "failure or delay in enforcing any right will be deemed a waiver.",
    "IN WITNESS WHEREOF, the parties have executed this Agreement as of the "
    "Effective Date.",
]

# Fixed legal boilerplate lifted verbatim from executed SOW samples.
PREAMBLE_SAMSUNG = (
    'This Statement of Work ("Statement of Work" or "SOW"), shall be construed and '
    'treated as a "Statement of Work" under, and as defined in, that certain '
    'Advertising Services Agreement ("Agreement") dated as of September 16, 2022 by '
    'and between Cheil USA Inc ("Agency" or "Cheil") and Samsung Electronics America, '
    'Inc., a New York corporation ("Samsung"). This SOW is valid and binding when '
    'signed on behalf of both parties and shall be effective as of the Start Date of '
    'Service ("Effective Date"). Capitalized terms used herein and not otherwise '
    "defined shall have the meaning given such terms in the Agreement."
)
PREAMBLE_AGENCY_1 = (
    'This Statement of Work ("Statement of Work" or "SOW") is made effective as of '
    '{sow_date} (the "Statement of Work Effective Date") by and between Cheil USA '
    "Inc., a Delaware corporation with its principal of business located at 837 "
    "Washington Street, 4th Floor, New York, NY 10014 on behalf of itself and its "
    'affiliates and subsidiaries ("Cheil") and {vendor_entity} ("Contractor").  '
    'Contractor and Cheil may each be referred to herein as a "Party", and, together '
    'as the "Parties".'
)
PREAMBLE_AGENCY_2 = (
    "This SOW is governed by, incorporated into, and made part of, that certain "
    'Master Services Agreement (the "Agreement"), dated as of {msa_date}, by and '
    "between Cheil and Contractor. This SOW defines the Services that Contractor "
    "shall provide to Cheil in accordance with the terms of the Agreement and this "
    "SOW.  The terms of this SOW are limited to the scope of this SOW, and shall not "
    "be applicable to any other SOWs, which may be executed and attached to the "
    "Agreement. Capitalized terms used herein and not otherwise defined shall have "
    "the meanings given them in the Agreement.  To the extent there is a conflict "
    "between the terms of this SOW and the Agreement, the terms of the Agreement "
    "shall control, except for terms where the Agreement expressly permits the SOW "
    "to control in the event of conflict with the Agreement."
)
OOP_SAMSUNG = (
    "All out of pocket expenses are a pass through cost as per the Agreement between "
    "Samsung and Cheil.\n\n"
    "No travel cost is included in the annual cost for providing this Service set "
    "forth in this SOW. If the Samsung team requires Agency personnel to travel, cost "
    "of approved travel will be reimbursed and invoiced separately within 30 days of "
    "completion of travel.\n\n"
    "Reimbursable travel-related expenses are limited to transportation, "
    "accommodation, and meals/subsistence costs, in each case that are directly "
    "related to work performed for Samsung pursuant to this SOW.  If (a) Agency "
    "reasonably requires any Agency personnel to travel to a specific site for the "
    "performance of the Services that is outside of such Agency personnel's local "
    "metropolitan area and (b) Samsung expressly approves such travel in writing, "
    "then the reasonable and documented expenses actually incurred as a result of "
    "such travel will be reimbursed by Samsung.  However, any such travel-related "
    "expenses will only be reimbursed where a travel expense maximum amount has been "
    "specifically authorized by the Samsung team in writing; and the total amount of "
    "travel expenses does not exceed the maximum amount authorized by the Samsung "
    "team. Travel-related expenses will be reimbursed by Samsung only if such "
    "expenses comply with Samsung's applicable travel policy, as provided from time "
    "to time by Samsung."
)
OOP_AGENCY = (
    "All out of pocket expenses are a pass through cost as per the most current "
    "Master Service Agreement between Cheil USA and Contractor.\n\n"
    "No travel cost is included in the cost for providing this service. If the Cheil "
    "team requires the Contractor to travel, cost of approved travel will be "
    "reimbursed and invoiced separately within 30 days of completion of travel."
)
PAYMENT_INTRO = (
    "Agency will invoice Samsung on a monthly basis at the beginning of the "
    "following month as set forth in the table below."
)
PAYMENT_INTRO_AGENCY = (
    "Contractor will invoice Cheil on a monthly basis at the end of the month as "
    "set forth in the table below."
)
CHANGE_ORDER_NOTE = (
    "Any additional costs for the Service beyond those outlined above must receive "
    "prior approval and be agreed upon through a change order to this Statement of "
    "Work executed by the parties."
)


# ── data ─────────────────────────────────────────────────────────────────────

def _data_path(user):
    return os.path.join(DATA_ROOT, user, "sow.json")


def _load(user):
    f = _data_path(user)
    if not os.path.exists(f):
        return {"sows": [], "vendors": [], "people": [], "contracts": []}
    try:
        with open(f) as fp:
            d = json.load(fp)
            d.setdefault("sows", [])
            d.setdefault("vendors", [])
            d.setdefault("people", [])
            d.setdefault("contracts", [])
            for p in d["people"]:
                _migrate_person(p)
            return d
    except Exception:
        return {"sows": [], "vendors": [], "people": [], "contracts": []}


def _migrate_person(p):
    """v1 roster fields → the 3-axis schema (Client↔Cheil / Cheil↔Partner /
    Cheil employee) requested 2026-07-22."""
    if "sell_hr" not in p:
        p["sell_hr"] = p.pop("rate", "") or ""
        p["cost_hr"] = p.pop("vendor_cost", "") or ""
        p["role_title"] = p.pop("function", "") or ""
        p["salary_mo"] = p.pop("salary", "") or ""
        # Seeded emails were @samsung.com working addresses.
        em = p.pop("email", "") or ""
        p["email_samsung"] = em if "samsung" in em else ""
        p["email_cheil"] = em if "cheil" in em else ""
    for k in ("project", "sell_hr", "sell_mo", "client_duration", "client_budget",
              "client_po", "cost_hr", "cost_mo", "partner_duration", "partner_cost",
              "partner_po", "role_title", "salary_mo", "cheil_since", "salary_oh",
              "email_cheil", "email_samsung", "pc", "svpn", "ebita", "location"):
        p.setdefault(k, "")
    p.setdefault("linked_sows", [])
    p.setdefault("linked_contracts", [])
    return p


# Roster columns worth editing in place — the ones retyped most often.
# Anything structural (name, affiliation, links) still goes through the form.
_PP_EDITABLE = {"role_title", "project", "sell_hr", "sell_mo", "client_budget",
                "client_po", "cost_hr", "cost_mo", "partner_cost", "partner_po",
                "salary_mo", "salary_oh", "location"}
_PP_MONEY = {"sell_hr", "sell_mo", "client_budget", "cost_hr", "cost_mo",
             "partner_cost", "salary_mo", "salary_oh"}


_PEOPLE_HEADERS = {
    "name": ("resource", "resource name", "name", "person", "employee",
             "team member", "member", "이름", "성명"),
    "role": ("function", "role", "title", "profile", "position", "job title",
             "role · title", "직무", "역할"),
    "email": ("email", "email id", "e-mail", "mail", "email address"),
    "location": ("location", "region", "country", "site", "base", "지역"),
}
# rate columns, most specific first — a sheet often carries an hourly rate AND
# a derived monthly one, and the hourly is the figure worth keeping
_RATE_HEADERS = (
    ("hour", ("hourly rate", "rate/hr", "rate / hr", "hourly", "rate per hour",
              "unit rate", "rate", "시급")),
    ("month", ("monthly rate", "rate/month", "rate / month", "rate/mo",
               "monthly", "per month", "월단가", "salary")),
)
_RATE_SKIP = ("monthly cost", "total", "allocation", "alloc", "qty", "quantity",
              "hrs", "hours", "months")


def _norm_header(v):
    return re.sub(r"\s+", " ", str(v or "").strip().lower()).strip(" :·|")


def _sheet_rows(content, ext):
    """Rows of strings out of an uploaded table — xlsx, csv or a pasted range.
    Everything downstream works on text, so numbers become plain strings."""
    ext = (ext or "").lower()
    if ext in ("xlsx", "xlsm"):
        import openpyxl
        wb = openpyxl.load_workbook(io.BytesIO(content), read_only=True, data_only=True)
        out = []
        for row in wb.active.iter_rows(values_only=True):
            out.append(["" if c is None else
                        (f"{c:g}" if isinstance(c, float) else str(c)) for c in row])
        wb.close()
        return out
    text = content.decode("utf-8-sig", errors="replace") if isinstance(content, bytes) else content
    return _text_rows(text)


def _text_rows(text):
    """A pasted Excel range arrives tab-separated; a saved export, comma."""
    import csv as _csv
    sample = "\n".join((text or "").splitlines()[:5])
    delim = "\t" if sample.count("\t") >= sample.count(",") and "\t" in sample else ","
    return [list(r) for r in _csv.reader((text or "").splitlines(), delimiter=delim)]


_MONTH_WORDS = {}
for _i in range(1, 13):
    _MONTH_WORDS[calendar.month_abbr[_i].lower()] = _i
    _MONTH_WORDS[calendar.month_name[_i].lower()] = _i

_TOTAL_WORDS = ("total", "totals", "sum", "subtotal", "grand total",
                "합계", "계", "총계", "소계")


def _yy(s):
    """'26' → 2026, '1998' → 1998. Billing sheets abbreviate the year."""
    n = int(s)
    if n >= 1000:
        return n
    return 2000 + n if n < 70 else 1900 + n


def _parse_month_label(v):
    """(year, month) out of a billing-schedule cell — 'Jan-26', 'January 2026',
    '2026-01', '2026/01/01', '1/2026'. None for anything else, so an amount
    column never gets mistaken for a month."""
    s = re.sub(r"\s+", " ", str(v or "").strip()).strip(" :·|")
    if not s:
        return None
    m = re.match(r"^(\d{4})[-/. ](\d{1,2})(?:[-/. ]\d{1,2})?(?:[ T].*)?$", s)
    if m and 1 <= int(m.group(2)) <= 12:
        return (int(m.group(1)), int(m.group(2)))
    m = re.match(r"^(\d{1,2})[-/. ](\d{4})$", s)
    if m and 1 <= int(m.group(1)) <= 12:
        return (int(m.group(2)), int(m.group(1)))
    m = re.match(r"^([A-Za-z]{3,9})\.?[-/. ,]*'?(\d{2,4})$", s)
    if m:
        mo = _MONTH_WORDS.get(m.group(1).lower())
        if mo:
            return (_yy(m.group(2)), mo)
    # '2026 Jan' / '26-Jan' — a 2-digit lead is only a year when it can't be a day
    m = re.match(r"^'?(\d{2,4})[-/. ,]*([A-Za-z]{3,9})\.?$", s)
    if m:
        mo = _MONTH_WORDS.get(m.group(2).lower())
        if mo and (len(m.group(1)) == 4 or int(m.group(1)) > 12):
            return (_yy(m.group(1)), mo)
    return None


def _is_total_row(row):
    for cell in row:
        t = _norm_header(cell)
        if t:
            return t in _TOTAL_WORDS or t.startswith("total")
    return False


def _col_letter(j):
    s = ""
    j += 1
    while j:
        j, r = divmod(j - 1, 26)
        s = chr(65 + r) + s
    return s


def _month_table_read(rows):
    """Everything a monthly billing sheet says, without deciding for the user
    (강프로 2026-07-28 — "15% cut simulation_1.xlsx" has BOTH an Original Cost
    and an Adjusted Cost column, and picking one silently is a wrong answer
    dressed as a right one):

      {"shape": "long"|"wide", "note": str, "warnings": [str],
       "columns": [{"key","label","total","months":{"YYYY-MM": v}}],
       "chosen": key}

    Two layouts, because both arrive:
      long — one row per month, one or more money columns beside it. Every
             money column is offered; date columns (an Invoice Month next to
             the billing month) fall out on their own, since a date is not a
             number.
      wide — months across a header row, figures underneath: the per-line rows
             are summed into one column and a Total row is skipped.
    Repeated months are reported, never quietly added together.
    """
    body = [r for r in rows if r and not _is_total_row(r)]
    # which column holds the months? whichever one parses as months most often
    counts = {}
    for r in body:
        for j, cell in enumerate(r):
            if _parse_month_label(cell):
                counts[j] = counts.get(j, 0) + 1
    mcol = max(counts, key=lambda j: (counts[j], -j)) if counts else None

    if mcol is not None and counts[mcol] >= 2:
        return _read_long(rows, body, mcol)
    wide = _read_wide(rows)
    if wide["columns"]:
        return wide
    if mcol is not None and counts[mcol] == 1:
        return _read_long(rows, body, mcol)          # a one-month change
    return {"shape": "", "note": "", "warnings": [], "columns": [], "chosen": ""}


def _read_long(rows, body, mcol):
    data = [r for r in body if mcol < len(r) and _parse_month_label(r[mcol])]
    # the row above the first month row usually names the columns
    first = rows.index(data[0])
    header = rows[first - 1] if first > 0 else []
    if any(_parse_month_label(c) for c in header):
        header = []

    def label(j):
        h = str(header[j]).strip() if j < len(header) else ""
        return h or f"Column {_col_letter(j)}"

    cols, seen = [], {}
    for j in range(max(len(r) for r in data)):
        if j == mcol:
            continue
        months, dups = {}, set()
        for r in data:
            n = _num_or_none(r[j]) if j < len(r) else None
            if n is None:
                continue
            ym = _parse_month_label(r[mcol])
            key = f"{ym[0]:04d}-{ym[1]:02d}"
            if key in months:
                dups.add(key)
            months[key] = months.get(key, 0.0) + n
        if not months:
            continue
        seen.update({k: 1 for k in dups})
        cols.append({"key": str(j), "label": label(j),
                     "total": sum(months.values()), "months": months})
    warn = []
    if seen:
        names = ", ".join(date(int(k[:4]), int(k[5:]), 1).strftime("%b %Y")
                          for k in sorted(seen))
        warn.append(f"{names} appears more than once in column "
                    f"{_col_letter(mcol)} — those rows were added together. "
                    f"Check the sheet if that is not what you meant.")
    if len(cols) > 1:
        warn.append(f"{len(cols)} money columns found — pick the one this "
                    f"change should bill.")
    note = _months_note(cols[0]["months"], "one row per month") if cols else ""
    return {"shape": "long", "note": note, "warnings": warn,
            "columns": cols, "chosen": cols[0]["key"] if cols else ""}


def _read_wide(rows):
    hdr_i, cols = None, {}
    for i, row in enumerate(rows[:15]):
        found = {}
        for j, cell in enumerate(row):
            ym = _parse_month_label(cell)
            if ym and ym not in found.values():
                found[j] = ym
        if len(found) > len(cols):
            hdr_i, cols = i, found
    months = {}
    if hdr_i is not None:
        for row in rows[hdr_i + 1:]:
            if not row or _is_total_row(row):
                continue
            for j, ym in cols.items():
                n = _num_or_none(row[j]) if j < len(row) else None
                if n is not None:
                    key = f"{ym[0]:04d}-{ym[1]:02d}"
                    months[key] = months.get(key, 0.0) + n
    if not months:
        return {"shape": "", "note": "", "warnings": [], "columns": [], "chosen": ""}
    col = {"key": "*", "label": "Monthly figures", "total": sum(months.values()),
           "months": months}
    return {"shape": "wide",
            "note": _months_note(months, "months across the header row"),
            "warnings": [], "columns": [col], "chosen": "*"}


def _month_amounts_from_table(rows):
    """({(y,m): amount}, note) for the sheet's default money column — the
    single-answer view of _month_table_read, used where no choice is offered."""
    read = _month_table_read(rows)
    if not read["columns"]:
        return {}, ""
    col = next(c for c in read["columns"] if c["key"] == read["chosen"])
    return ({_parse_month_label(k): v for k, v in col["months"].items()},
            read["note"])


def _months_note(months, shape):
    ks = sorted(_parse_month_label(k) if isinstance(k, str) else k for k in months)
    span = (f"{date(*ks[0], 1).strftime('%b %Y')} – {date(*ks[-1], 1).strftime('%b %Y')}"
            if ks else "")
    return f"{len(ks)} month(s) · {span} · {shape}"


def _find_header_row(rows):
    """(index, {field: column}) for the first row that reads like a header.
    Sheets in the wild start with a title row and a blank column, so the header
    is found by content, never by position."""
    best = None
    for i, row in enumerate(rows[:15]):
        cols, rate_col, rate_basis = {}, None, ""
        for j, cell in enumerate(row):
            h = _norm_header(cell)
            if not h:
                continue
            for field, names in _PEOPLE_HEADERS.items():
                if field not in cols and h in names:
                    cols[field] = j
            if rate_col is None and not any(s in h for s in _RATE_SKIP):
                for basis, names in _RATE_HEADERS:
                    if h in names:
                        rate_col, rate_basis = j, basis
                        break
        if "name" not in cols:
            continue
        score = len(cols) + (1 if rate_col is not None else 0)
        if best is None or score > best[2]:
            if rate_col is not None:
                cols["rate"] = rate_col
            best = (i, cols, score, rate_basis)
    if best is None:
        return None, {}, ""
    return best[0], best[1], best[3]


def _people_from_table(rows):
    """[{name, role, location, rate, rate_basis, email}] from a team sheet.
    Same shape the contract extractor produces, so both feed one review list."""
    hdr_i, cols, basis = _find_header_row(rows)
    if hdr_i is None:
        return [], ""
    out = []
    for row in rows[hdr_i + 1:]:
        def cell(field):
            j = cols.get(field)
            return str(row[j]).strip() if j is not None and j < len(row) else ""
        name = cell("name")
        if not name or _norm_header(name) in ("total", "totals", "sum", "합계", "계"):
            continue
        rate = cell("rate")
        n = _num_or_none(rate)
        out.append({"name": name[:80], "role": cell("role")[:80],
                    "location": cell("location")[:60],
                    "email": cell("email")[:120],
                    "rate": (_money(n) if n is not None else rate[:40]),
                    "rate_basis": _rate_basis(rate, basis)})
    note = ", ".join(k for k in ("name", "role", "rate", "location", "email") if k in cols)
    return out, note


def _merge_pending(existing, found):
    """Fold a fresh read into what is already staged — one entry per person,
    new values filling blanks rather than replacing the earlier read."""
    by_name = {}
    order = []
    for e in list(existing or []) + list(found or []):
        key = (e.get("name") or "").strip().lower()
        if not key:
            continue
        if key not in by_name:
            by_name[key] = dict(e)
            order.append(key)
            continue
        cur = by_name[key]
        for k, v in e.items():
            if v and not cur.get(k):
                cur[k] = v
    return [by_name[k] for k in order]


def _rate_basis(rate, stated=""):
    """hour | month. The document usually says; when it doesn't, magnitude
    decides — nobody bills $12,000 an hour or $25 a month."""
    if stated in ("hour", "month"):
        return stated
    n = _num_or_none(rate)
    if n is None:
        return ""
    return "hour" if n < 1000 else "month"


def _apply_extracted_person(p, ext, c):
    """Fold a resource row read out of a contract into a roster person
    (강프로 2026-07-27). Never overwrites a value that is already there — the
    user's own typing outranks the extractor — and the rate lands on the axis
    the contract direction dictates: a vendor contract states what Cheil PAYS,
    a SEA contract what Cheil BILLS."""
    axis = "cost" if c.get("side") == "vendor" else "sell"
    fill = {"role_title": (ext.get("role") or "").strip(),
            "location": (ext.get("location") or "").strip(),
            "project": (c.get("project_name") or "").strip()}
    basis = _rate_basis(ext.get("rate"), ext.get("rate_basis") or "")
    if basis:
        fill[f"{axis}_{'hr' if basis == 'hour' else 'mo'}"] = (ext.get("rate") or "").strip()
    period = " ~ ".join(x for x in (c.get("period_start"), c.get("period_end")) if x)
    if period:
        fill["partner_duration" if axis == "cost" else "client_duration"] = period
    for k, v in fill.items():
        if v and not (p.get(k) or "").strip():
            p[k] = v
    return p


def _num_or_none(v):
    try:
        s = str(v).replace(",", "").replace("$", "").strip()
        return float(s) if s else None
    except (TypeError, ValueError):
        return None


def _person_ebita(p):
    """EBITA by Cheil — manual value wins; else Client budget − Partner cost
    (vendor personnel). Cheil-employee EBITA needs salary+OH × duration and
    stays manual until durations are numeric."""
    manual = _num_or_none(p.get("ebita"))
    if manual is not None:
        return manual, False
    budget = _num_or_none(p.get("client_budget"))
    cost = _num_or_none(p.get("partner_cost"))
    if budget is not None and cost is not None:
        return budget - cost, True
    return None, True


def _save(user, data):
    f = _data_path(user)
    os.makedirs(os.path.dirname(f), exist_ok=True)
    with open(f, "w") as fp:
        json.dump(data, fp, ensure_ascii=False, indent=2)


def _sow_type(sow):
    """Resolve a record's type key. Legacy team/role types collapse into the
    merged per-direction SOW type (the rate model lives on the record)."""
    t = sow.get("type")
    if t in TYPES:
        return t
    d = "agy" if sow.get("direction") == "agency" else "sea"
    return f"{d}_sow"


# ── schedule math ────────────────────────────────────────────────────────────

def _parse_date(s):
    try:
        return date.fromisoformat((s or "").strip())
    except Exception:
        return None


def _month_spans(start, end):
    """[(year, month, fraction)] covering start..end. Partial months use the
    days/30 convention observed in executed SOWs (Mar 23-31 → 0.3)."""
    if not start or not end or end < start:
        return []
    spans, y, m = [], start.year, start.month
    while (y, m) <= (end.year, end.month):
        last = calendar.monthrange(y, m)[1]
        s = start if (y, m) == (start.year, start.month) else date(y, m, 1)
        e = end if (y, m) == (end.year, end.month) else date(y, m, last)
        full = s.day == 1 and e.day == last
        frac = 1.0 if full else min(1.0, round(((e - s).days + 1) / 30, 2))
        spans.append((y, m, frac))
        y, m = (y + 1, 1) if m == 12 else (y, m + 1)
    return spans


def _monthly_amount(sow):
    total = 0.0
    for r in sow.get("resources", []):
        try:
            if sow.get("res_mode") == "hourly":
                total += float(r.get("hourly") or 0) * float(r.get("hrs") or 0) * float(r.get("qty") or 1)
            else:
                total += float(r.get("rate") or 0)
        except (TypeError, ValueError):
            pass
    return round(total, 2)


def _build_schedule(sow):
    start, end = _parse_date(sow.get("start")), _parse_date(sow.get("end"))
    spans = _month_spans(start, end)
    monthly = _monthly_amount(sow)
    rule = sow.get("invoice_rule") or "next_first"
    overrides = sow.get("schedule_overrides") or {}
    rows = []
    for y, m, frac in spans:
        label = date(y, m, 1).strftime("%b-%y")
        amount = round(monthly * frac, 2)
        if label in overrides:
            try:
                amount = round(float(overrides[label]), 2)
            except (TypeError, ValueError):
                pass
        if rule == "month_end":
            inv = date(y, m, calendar.monthrange(y, m)[1])
        else:
            inv = date(y + 1, 1, 1) if m == 12 else date(y, m + 1, 1)
        rows.append({
            "label": label,
            "amount": amount,
            "invoice": inv.strftime("%-d-%b-%y"),
        })
    # months comes back exact; round only for display so per-line Cost stays
    # equal to the schedule total (a $-level mismatch reads as an error in a
    # finance document).
    return rows, round(sum(r["amount"] for r in rows), 2), sum(f for _, _, f in spans)


def _money(x):
    try:
        x = float(x)
    except (TypeError, ValueError):
        return "$0"
    return f"${x:,.0f}" if abs(x - round(x)) < 0.005 else f"${x:,.2f}"


def _fmt_long(iso):
    d = _parse_date(iso)
    return d.strftime("%B %-d, %Y") if d else (iso or "")


# ── docx ─────────────────────────────────────────────────────────────────────

def _build_docx(sow, vendor):
    from docx import Document
    from docx.shared import Pt

    schedule, fee, months_total = _build_schedule(sow)
    is_agency = sow.get("direction") == "agency"
    counterpart = (vendor or {}).get("name") or "Contractor" if is_agency else SAMSUNG_ENTITY

    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10)

    def p(text="", bold=False, size=None):
        par = doc.add_paragraph()
        run = par.add_run(text)
        run.bold = bold
        if size:
            run.font.size = Pt(size)
        return par

    def h(text):
        p(text, bold=True, size=12)

    def table(headers, rows, total_row=None):
        t = doc.add_table(rows=1, cols=len(headers))
        t.style = "Table Grid"
        for i, hd in enumerate(headers):
            cell = t.rows[0].cells[i]
            cell.text = ""
            cell.paragraphs[0].add_run(hd).bold = True
        for row in rows:
            cells = t.add_row().cells
            for i, v in enumerate(row):
                cells[i].text = str(v)
        if total_row:
            cells = t.add_row().cells
            for i, v in enumerate(total_row):
                cells[i].text = ""
                cells[i].paragraphs[0].add_run(str(v)).bold = True
        return t

    # SEA SOWs open with a cover (head) page: Cheil logo + title + meta block,
    # then a page break into the legal preamble. Agency SOWs start directly.
    if not is_agency:
        logo = os.path.join(_ASSETS, "cheil_logo.png")
        if os.path.exists(logo):
            from docx.shared import Inches
            doc.add_picture(logo, width=Inches(3.2))
        p()
    p(sow.get("title") or sow.get("project_name") or "", bold=True, size=14)
    p("STATEMENT OF WORK", bold=True, size=13)
    p()
    p(f"DATE: {_fmt_long(sow.get('date'))}")
    p(f"CLIENT: {SAMSUNG_ENTITY if not is_agency else CHEIL_ENTITY}")
    p(f"PROJECT NAME: {sow.get('project_name') or ''}")
    p(f"PREPARED BY: {sow.get('prepared_by') or ''}")
    p(f"PREPARED FOR: {sow.get('prepared_for') or ''}")
    if not is_agency:
        doc.add_page_break()
    else:
        p()

    if is_agency:
        p(PREAMBLE_AGENCY_1.format(
            sow_date=_fmt_long(sow.get("date")),
            vendor_entity=(vendor or {}).get("entity_line") or counterpart,
        ))
        p()
        p(PREAMBLE_AGENCY_2.format(msa_date=_fmt_long((vendor or {}).get("msa_date")) or "the MSA date"))
    else:
        p(PREAMBLE_SAMSUNG)
    p()

    h("Executive Summary")
    for line in (sow.get("exec_summary") or "").splitlines():
        if line.strip():
            p(line.strip())
    p()

    h("Deliverables:" if not is_agency else "Service Description:")
    for line in (sow.get("deliverables") or "").splitlines():
        if line.strip():
            doc.add_paragraph(line.strip(), style="List Bullet")
    p()

    h("Project Stakeholders:")
    client_label = ((vendor or {}).get("name") or "Contractor") + " POC" if is_agency \
        else "Samsung Manager for this Role"
    table(
        ["", client_label, "Cheil Project Management & SOW Owner"],
        [
            ["Name", sow.get("stk_c_name") or "", sow.get("stk_a_name") or ""],
            ["Email", sow.get("stk_c_email") or "", sow.get("stk_a_email") or ""],
            ["Location", sow.get("stk_c_loc") or "", sow.get("stk_a_loc") or ""],
        ],
    )
    p()

    h("Service Period")
    p(f"Start Date : {_fmt_long(sow.get('start'))}")
    p(f"End Date : {_fmt_long(sow.get('end'))}")
    p()

    h("Resource Management" if not is_agency else "Resource Planning")
    payer, payee = ("Cheil", "Contractor") if is_agency else ("Samsung", "Cheil")
    p(
        "In consideration for the provision of the Services and Deliverables under "
        f"this SOW, {payer} shall pay {payee} in accordance with the following rates "
        "and fees, subject to the applicable terms and conditions of the Agreement:"
    )
    if sow.get("res_mode") == "hourly":
        rows = []
        for r in sow.get("resources", []):
            qty = float(r.get("qty") or 1)
            cost = float(r.get("hourly") or 0) * float(r.get("hrs") or 0) * qty * months_total
            rows.append([
                r.get("profile") or "", r.get("location") or "", int(qty), round(months_total, 1),
                _money(r.get("hourly")), r.get("hrs") or "", _money(cost),
            ])
        table(
            ["Profile", "Location", "# of Resources", "# of Months", "Hourly Cost",
             "# of Anticipated Hrs/Month", "Cost"],
            rows, total_row=["Total Cost", "", "", "", "", "", _money(fee)],
        )
    else:
        rows = []
        for i, r in enumerate(sow.get("resources", []), 1):
            rows.append([
                i, r.get("name") or "", r.get("role") or "", r.get("level") or "",
                r.get("region") or "", _money(r.get("rate")),
            ])
        table(["No.", "Name", "Role", "Level", "Region", "Rate/Month (USD)"], rows)
    p()

    h("Cost and Payment Schedule")
    p(f"Fee : {_money(fee)}", bold=True)
    p()
    p("Payment Schedule :")
    p(PAYMENT_INTRO_AGENCY if is_agency else PAYMENT_INTRO)
    table(
        ["Month", "Amount", "Invoice Date"],
        [[r["label"], _money(r["amount"]), r["invoice"]] for r in schedule],
        total_row=["Total", _money(fee), ""],
    )
    p()
    p(CHANGE_ORDER_NOTE)
    p()

    h("Out-of-pocket Expense")
    for para in (OOP_AGENCY if is_agency else OOP_SAMSUNG).split("\n\n"):
        p(para)
        p()

    h("Signatures")
    p(
        "IN WITNESS WHEREOF, the parties have caused this Statement of Work to be "
        "duly executed by their authorized representatives as set forth below."
    )
    p()
    left = CHEIL_ENTITY if is_agency else SAMSUNG_ENTITY
    right = ((vendor or {}).get("name") or "Contractor") if is_agency else CHEIL_ENTITY
    sig = doc.add_table(rows=5, cols=2)
    sig.rows[0].cells[0].paragraphs[0].add_run(left).bold = True
    sig.rows[0].cells[1].paragraphs[0].add_run(right).bold = True
    for i, lbl in enumerate(["Signature:", "Name:", "Title:", "Date:"], 1):
        for c in range(2):
            sig.rows[i].cells[c].text = f"{lbl} _______________________"

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _replace_in_par(par, old, new):
    """Replace `old` with `new` across a paragraph's runs, splicing at run
    boundaries so surrounding formatting survives (docx runs split words)."""
    full = "".join(r.text for r in par.runs)
    idx = full.find(old)
    if idx < 0:
        return False
    end = idx + len(old)
    pos = 0
    for r in par.runs:
        r_start, r_end = pos, pos + len(r.text)
        if r_end <= idx or r_start >= end:
            pos = r_end
            continue
        keep_head = r.text[: max(0, idx - r_start)]
        keep_tail = r.text[max(0, min(len(r.text), end - r_start)):]
        if r_start <= idx:
            r.text = keep_head + new + keep_tail
        else:
            r.text = keep_tail
        pos = r_end
    return True


def _build_msa_docx(sow, vendor):
    """Fill the stored MSA template (executed-format docx) with the vendor
    name and effective date; every other clause exports byte-identical."""
    from docx import Document
    doc = Document(os.path.join(_ASSETS, "msa_template.docx"))
    reps = [
        ("XXX XX, 2026", _fmt_long(sow.get("date")) or "____________"),
        ("(Your Company Name)", (vendor or {}).get("name") or "____________________"),
    ]
    def walk(pars):
        for par in pars:
            for old, new in reps:
                while _replace_in_par(par, old, new):
                    pass
    walk(doc.paragraphs)
    for t in doc.tables:
        for row in t.rows:
            for c in row.cells:
                walk(c.paragraphs)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _build_nda_docx(sow, vendor):
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt

    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(11)
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run(NDA_TITLE).bold = True
    doc.add_paragraph()
    doc.add_paragraph(NDA_INTRO.format(
        date=_fmt_long(sow.get("date")) or "____________",
        vendor=(vendor or {}).get("name") or "______________________",
    ))
    for par in NDA_BODY:
        doc.add_paragraph()
        doc.add_paragraph(par)
    doc.add_paragraph()
    sig = doc.add_table(rows=4, cols=2)
    sig.rows[0].cells[0].paragraphs[0].add_run(CHEIL_ENTITY).bold = True
    sig.rows[0].cells[1].paragraphs[0].add_run(
        (vendor or {}).get("name") or "[_____________________]").bold = True
    for i, lbl in enumerate(["By: _________________________", "Name:", "Title:"], 1):
        for c in range(2):
            sig.rows[i].cells[c].text = lbl
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ── html shells ──────────────────────────────────────────────────────────────

def _esc(s):
    return (str(s or "").replace("&", "&amp;").replace("<", "&lt;")
            .replace(">", "&gt;").replace('"', "&quot;"))


_CSS = """
/* ── pill tab bar (cardconv pattern) ── */
.dd-tabbar{position:sticky;top:52px;z-index:90;background:var(--bg-deep);padding:12px 0 10px;margin:-12px 0 10px}
.dd-tabs{display:inline-flex;align-items:center;gap:2px;padding:3px;background:var(--surface-2);border:1px solid var(--border);border-radius:var(--radius-md);flex-wrap:wrap;max-width:100%}
.dd-tab{display:inline-flex;align-items:center;padding:7px 16px;font-size:.82rem;font-weight:600;color:var(--text-muted);border-radius:var(--radius-sm);text-decoration:none;transition:background .15s,color .15s}
.dd-tab:hover{color:var(--text)}
.dd-tab.active{background:var(--accent);color:var(--on-accent)}
/* ── landing stat cards ── */
.dd-stats{display:grid;grid-template-columns:repeat(4,1fr);gap:12px;margin:0 0 8px}
.dd-stat{background:var(--surface);border:1px solid var(--border);border-radius:var(--radius-lg);padding:16px 18px;text-decoration:none;color:var(--text);transition:.15s;display:flex;flex-direction:column;gap:2px;min-width:0}
.dd-stat .lb,.dd-stat .sub{min-width:0;overflow-wrap:anywhere}
.dd-stat:hover{border-color:var(--accent);transform:translateY(-2px);box-shadow:var(--shadow-md)}
.dd-stat b{font-size:1.45rem;font-variant-numeric:tabular-nums;letter-spacing:-.02em}
.dd-stat .lb{font-size:.72rem;color:var(--text-muted);font-weight:600}
.dd-stat .sub{font-size:.68rem;color:var(--text-muted)}
.dd-stat .sub.pos{color:var(--success)}
.dd-stat .sub.neg{color:var(--danger)}
.dd-stat .cash-pos{color:var(--success)}
.dd-stat .cash-pay{color:var(--group-4)}
.dd-stat .cash-neg{color:var(--danger)}
button.dd-tab{border:none;background:transparent;cursor:pointer;font-family:inherit}
/* ── Home To Do alert ── */
.dd-todo{background:rgba(245,158,11,.06);border:1px solid rgba(245,158,11,.35);border-radius:var(--radius-lg);padding:14px 16px;margin-bottom:16px}
.dd-todo-hd{font-size:.84rem;font-weight:800;margin-bottom:8px;display:flex;align-items:center;gap:8px;flex-wrap:wrap}
.dd-todo-item{display:flex;gap:10px;align-items:baseline;padding:7px 4px;font-size:.82rem;color:var(--text);text-decoration:none;border-top:1px dashed var(--border)}
.dd-todo-item:hover{background:rgba(245,158,11,.08)}
.tab-badge{display:inline-flex;align-items:center;justify-content:center;min-width:16px;height:16px;border-radius:8px;font-size:.62rem;font-weight:700;color:#fff;padding:0 5px}
.sow-hero{display:grid;grid-template-columns:1fr 1fr;gap:16px;margin:20px 0 34px}
.dir-card{display:flex;flex-direction:column;gap:10px;background:var(--surface);border:1px solid var(--border);border-radius:var(--radius-xl);padding:30px 28px;text-decoration:none;color:var(--text);transition:.2s;position:relative;overflow:hidden}
.dir-card::before{content:"";position:absolute;top:0;left:0;right:0;height:3px;background:var(--dir-color,var(--accent));opacity:.85}
.dir-card:hover{border-color:var(--dir-color,var(--accent));transform:translateY(-3px);box-shadow:var(--shadow-lg)}
.dir-icon{font-size:2.2rem}
.dir-name{font-size:1.15rem;font-weight:800;letter-spacing:-.02em}
.dir-desc{font-size:.82rem;color:var(--text-muted);line-height:1.5}
.type-grid{display:grid;grid-template-columns:1fr 1fr;gap:14px;margin:20px 0}
.type-card{display:flex;gap:14px;align-items:flex-start;background:var(--surface);border:1px solid var(--border);border-radius:var(--radius-lg);padding:20px;text-decoration:none;color:var(--text);transition:.2s}
.type-card:hover{border-color:var(--accent);transform:translateY(-2px);box-shadow:var(--shadow-md)}
.type-icon{font-size:1.6rem}
.type-name{font-weight:800;font-size:.95rem;margin-bottom:4px}
.type-desc{font-size:.78rem;color:var(--text-muted);line-height:1.5}
.sow-list{display:flex;flex-direction:column;gap:10px}
.sow-row{display:flex;align-items:center;gap:14px;background:var(--surface);border:1px solid var(--border);border-radius:var(--radius-lg);padding:14px 18px;flex-wrap:wrap}
.sow-row:hover{border-color:var(--accent)}
.sow-title{font-weight:700;color:var(--text);font-size:.92rem}
.sow-meta{font-size:.76rem;color:var(--text-muted)}
.sow-fee{font-weight:700;color:var(--success);font-variant-numeric:tabular-nums;margin-left:auto}
.dir-chip{font-size:.64rem;font-weight:700;padding:2px 8px;border-radius:10px;white-space:nowrap}
/* ── vendors registry rows ── */
.vnd-form{display:grid;grid-template-columns:1fr 2fr 150px auto auto;gap:10px;align-items:center;padding:14px 18px}
.vnd-form.new{grid-template-columns:1fr 2fr 150px auto;border-style:dashed}
.dir-samsung{color:var(--accent);background:var(--accent-glow)}
.dir-agency{color:var(--group-4);background:rgba(251,146,60,.12)}
/* ── document preview editor ── */
.doc-bar{position:sticky;top:52px;z-index:60;display:flex;align-items:center;gap:10px;background:var(--surface-3);border:1px solid var(--border-bright);border-radius:var(--radius-md);padding:10px 14px;margin-bottom:18px;flex-wrap:wrap}
.doc-bar .spacer{flex:1}
.paper{background:var(--surface);border:1px solid var(--border);border-radius:var(--radius-lg);padding:46px 52px;max-width:860px;margin:0 auto 60px;line-height:1.55;font-size:.9rem;color:var(--text)}
.paper h2{font-size:1.05rem;font-weight:800;margin:26px 0 8px}
.paper .doc-title{font-size:1.25rem;font-weight:800}
.paper .legal{color:var(--text-muted);font-size:.8rem;margin:10px 0}
.paper table{width:100%;border-collapse:collapse;margin:10px 0;font-size:.84rem}
.paper th{border:1px solid var(--border-bright);padding:7px 9px;font-size:.7rem;text-transform:uppercase;letter-spacing:.04em;color:var(--text-muted);text-align:left;background:var(--surface-2)}
.paper td{border:1px solid var(--border);padding:6px 9px}
.paper .num{text-align:right;font-variant-numeric:tabular-nums}
.slot{background:rgba(56,189,248,.07);border:none;border-bottom:1.5px dashed rgba(56,189,248,.55);border-radius:4px 4px 0 0;color:var(--text);font:inherit;padding:2px 6px;outline:none;min-width:60px}
.slot:focus{background:rgba(56,189,248,.14);border-bottom-color:var(--accent)}
textarea.slot{width:100%;border:1.5px dashed rgba(56,189,248,.45);border-radius:8px;padding:8px 10px;resize:vertical;line-height:1.5}
select.slot{cursor:pointer}
.paper td .slot{width:100%;min-width:0;padding:3px 5px}
.meta-line{display:flex;gap:8px;align-items:baseline;margin:2px 0}
.meta-line b{font-size:.8rem;letter-spacing:.03em;white-space:nowrap}
.ro-note{font-size:.68rem;color:var(--text-muted);text-transform:uppercase;letter-spacing:.06em;margin-top:26px}
.row-del{background:none;border:none;color:var(--text-muted);cursor:pointer;font-size:.9rem;padding:4px 8px}
.row-del:hover{color:var(--danger)}
.add-row-btn{margin-top:6px}
/* ── side-by-side executed example ── */
.ed-wrap{display:grid;grid-template-columns:minmax(0,1fr) minmax(0,1fr);gap:18px;align-items:start;max-width:1760px;margin:0 auto}
.ed-wrap>.paper{margin:0 0 60px;max-width:none}
.ed-wrap.ex-hidden{grid-template-columns:minmax(0,1fr);max-width:900px}
.ed-wrap.ex-hidden .ex-col{display:none}
.ex-col{position:sticky;top:118px;max-height:calc(100vh - 140px);overflow-y:auto;border-radius:var(--radius-lg)}
.ex-paper{background:var(--surface);border:1px dashed var(--border-bright);border-radius:var(--radius-lg);padding:26px 30px;font-size:.8rem;line-height:1.55;color:var(--text-muted)}
.ex-paper h3{font-size:.85rem;font-weight:800;color:var(--text);margin:16px 0 4px}
.ex-badge{display:inline-flex;align-items:center;gap:6px;background:rgba(52,211,153,.1);border:1px solid rgba(52,211,153,.3);color:var(--success);border-radius:99px;padding:3px 12px;font-size:.7rem;font-weight:700;margin-bottom:14px}
.ex-paper table{width:100%;border-collapse:collapse;margin:6px 0;font-size:.72rem}
.ex-paper th{border:1px solid var(--border);padding:4px 6px;font-size:.62rem;text-transform:uppercase;color:var(--text-muted);text-align:left;background:var(--surface-2)}
.ex-paper td{border:1px solid var(--border);padding:4px 6px}
.ex-paper ul{padding-left:18px;margin:4px 0}
.ex-paper li{margin-bottom:4px}
.ex-paper .num{text-align:right;font-variant-numeric:tabular-nums}
.ex-mark{background:rgba(56,189,248,.14);border-bottom:1.5px dashed rgba(56,189,248,.55);border-radius:3px;padding:0 3px;color:var(--text)}
@media(max-width:1100px){.ed-wrap{grid-template-columns:1fr}.ex-col{position:static;max-height:none;order:2}}
@media(max-width:768px){
  .dd-tabbar{padding:8px 0}
  .dd-tabs{display:flex;flex-wrap:nowrap;overflow-x:auto;scrollbar-width:none;-webkit-overflow-scrolling:touch}
  .dd-tabs::-webkit-scrollbar{display:none}
  .dd-tab{flex:0 0 auto;padding:10px 16px}
  .dd-stats{grid-template-columns:repeat(2,1fr)}
  .dd-stat{padding:14px 12px}
  .dd-stat b{font-size:1.05rem;letter-spacing:-.03em;white-space:nowrap}
  /* registry rows stack; the two actions share the last line (44px thumb zone) */
  .vnd-form,.vnd-form.new{grid-template-columns:1fr 1fr;padding:12px 14px;gap:8px}
  .vnd-form>input{grid-column:1/-1}
  .vnd-form>span{grid-column:1/-1;text-align:right}
  .vnd-form .btn{width:100%;min-height:44px}
  .vnd-form.new .btn{grid-column:1/-1}
  .cf-details>summary{padding:10px 2px}
  .ctr-chip{min-height:32px;display:inline-flex;align-items:center}
  .sow-hero,.type-grid{grid-template-columns:1fr}
  .paper{padding:22px 16px}
  .doc-bar{top:0;position:static}
  .paper .table-wrap{overflow-x:auto}
  .paper .table-wrap table{min-width:620px}
  .sow-fee{margin-left:0;flex-basis:100%}
}
"""


_TABS = [
    ("home", "/sow", "🏠 Home"),
    ("contracts", "/sow/contracts", "📎 Contracts"),
    ("docs", "/sow/docs", "📄 Documents"),
    ("people", "/sow/people", "👥 People"),
    ("vendors", "/sow/vendors", "🏢 Vendors"),
]


def _tab_bar(active):
    links = "".join(
        f'<a href="{href}" class="dd-tab{" active" if k == active else ""}">{label}</a>'
        for k, href, label in _TABS)
    # keep the pill pinned just below the live nav height (nav wraps on mobile)
    sync = ('<script>(function(){var n=document.querySelector("nav"),'
            't=document.querySelector(".dd-tabbar");if(!n||!t)return;'
            'var f=function(){t.style.top=n.offsetHeight+"px"};'
            'f();window.addEventListener("resize",f);'
            'var a=t.querySelector(".dd-tab.active"),r=a&&a.parentElement;'
            'if(r&&r.scrollWidth>r.clientWidth)r.scrollLeft=a.offsetLeft-(r.clientWidth-a.offsetWidth)/2;'
            '})();</script>')
    return f'<div class="dd-tabbar"><div class="dd-tabs">{links}</div></div>{sync}'


def _shell(user, title, body, wide=False, tab=None):
    """Page frame. Every tabbed page runs at the full app width (강프로
    2026-07-28): People was the only wide one, so switching to any other tab
    snapped the content into a 1000px column and read as boxed in. Narrow stays
    for the leaf forms reached from a tab — a person, a document type — where a
    single column is the point."""
    tabs = _tab_bar(tab) if tab else ""
    wide = wide or bool(tab)
    return f"""<!doctype html><html lang="en"><head>
<meta charset="utf-8"><meta name="viewport" content="width=device-width, initial-scale=1">
<title>🤝 {_esc(title)} · Wayfinder</title><link rel="stylesheet" href="/static/style.css">
<style>{_CSS}{_CTR_CSS}</style></head><body>
<nav><span class="nav-brand">🤝 Deal Desk</span>
<span class="nav-user">👤 {_esc(user)} &nbsp;·&nbsp; <a href="/logout">Logout</a></span></nav>
<div class="container" style="max-width:{'1800px' if wide else '1000px'}">{tabs}{body}</div></body></html>"""


def _mk(v, field):
    """Highlighted fill-in value in the sample document — maps 1:1 to an
    editable slot in the template on the left."""
    return f'<mark class="ex-mark" title="Fill-in field: {_esc(field)}">{_esc(v)}</mark>'


def _render_example(type_key):
    ex = SAMPLES.get(type_key)
    if not ex:
        return ""
    head = f'<div class="ex-badge">📖 Executed document · {_esc(ex["src"])}</div>'
    if ex.get("note"):
        return ('<div class="ex-col"><div class="ex-paper">' + head +
                f'<p>{_esc(ex["note"])}</p></div></div>')

    is_agency = type_key.startswith("agy")
    p = []
    p.append(head)
    p.append('<div style="font-size:.66rem;color:var(--text-muted);text-transform:uppercase;'
             'letter-spacing:.06em;margin-bottom:14px">Highlighted = the fields you fill in on the left</div>')
    # ── cover / head ──
    if not is_agency:
        p.append('<img src="/sow/asset/logo" alt="Cheil" style="max-width:220px;background:#fff;'
                 'padding:8px 12px;border-radius:6px;margin-bottom:16px">')
    p.append(f'<div style="font-size:1.05rem;font-weight:800;color:var(--text)">{_mk(ex["title"], "Title")}</div>')
    p.append('<div style="font-weight:800;color:var(--text);margin:6px 0 12px">STATEMENT OF WORK</div>')
    p.append(f'<p><b>DATE:</b> {_mk(ex["date"], "SOW Date")}<br>'
             f'<b>CLIENT:</b> {_esc(CHEIL_ENTITY if is_agency else SAMSUNG_ENTITY)}<br>'
             f'<b>PROJECT NAME:</b> {_mk(ex["project"], "Project Name")}<br>'
             + (f'<b>PREPARED BY:</b> {_mk(ex["by"], "Prepared By")}<br>'
                f'<b>PREPARED FOR:</b> {_mk(ex["for"], "Prepared For")}' if ex.get("by") else "")
             + '</p>')
    if not is_agency:
        p.append('<div style="border-top:2px dashed var(--border-bright);margin:16px -30px;position:relative">'
                 '<span style="position:absolute;top:-8px;left:50%;transform:translateX(-50%);'
                 'background:var(--surface);padding:0 8px;font-size:.6rem;color:var(--text-muted);'
                 'text-transform:uppercase;letter-spacing:.08em">Page 2</span></div>')
    # ── preamble ──
    if is_agency:
        p.append("<p>" + _esc(PREAMBLE_AGENCY_1).replace("{sow_date}", _mk(ex["date"], "SOW Date"))
                 .replace("{vendor_entity}", _mk(ex["vendor_entity"], "Vendor entity line")) + "</p>")
        p.append("<p>" + _esc(PREAMBLE_AGENCY_2).replace("{msa_date}", _mk(ex["msa_date"], "Vendor MSA date")) + "</p>")
    else:
        p.append(f"<p>{_esc(PREAMBLE_SAMSUNG)}</p>")
    # ── summary / deliverables ──
    p.append(f'<h3>Executive Summary</h3><p>{_mk(ex["summary"], "Executive Summary")}</p>')
    p.append(f'<h3>{"Service Description" if is_agency else "Deliverables"}</h3><ul>'
             + "".join(f"<li>{_mk(d, 'Deliverables (one per line)')}</li>" for d in ex["deliverables"])
             + "</ul>")
    # ── stakeholders ──
    stk_head = (ex.get("vendor_name", "Contractor") + " POC") if is_agency else "Samsung Manager for this Role"
    labels = ["Name", "Email", "Location"]
    stk_rows = "".join(
        f"<tr><td><b>{labels[i]}</b></td><td>{_mk(row[0], 'Stakeholder') if row[0] else ''}</td>"
        f"<td>{_mk(row[1], 'Stakeholder') if row[1] else ''}</td></tr>"
        for i, row in enumerate(ex["stk"]))
    p.append(f'<h3>Project Stakeholders</h3><div style="overflow-x:auto"><table>'
             f'<tr><th></th><th>{_esc(stk_head)}</th><th>Cheil Project Management &amp; SOW Owner</th></tr>'
             f'{stk_rows}</table></div>')
    # ── period ──
    p.append(f'<h3>Service Period</h3><p>Start Date : {_mk(ex["start"], "Start Date")}<br>'
             f'End Date : {_mk(ex["end"], "End Date")}</p>')
    # ── resources ──
    payer = "Cheil shall pay Contractor" if is_agency else "Samsung shall pay Cheil"
    p.append(f'<h3>{"Resource Planning" if is_agency else "Resource Management"}</h3>'
             f'<p>In consideration for the provision of the Services and Deliverables under this SOW, '
             f'{payer} in accordance with the following rates and fees, subject to the applicable '
             f'terms and conditions of the Agreement:</p>')
    res_head = ["Profile", "Location", "Qty", "# Months", "Hourly", "Hrs/Mo", "Cost"]
    res_rows = "".join("<tr>" + "".join(f"<td>{_mk(c, 'Resources table')}</td>" for c in r) + "</tr>"
                       for r in ex["res_rows"])
    p.append('<div style="overflow-x:auto"><table><tr>'
             + "".join(f"<th>{h}</th>" for h in res_head) + f"</tr>{res_rows}</table></div>")
    # ── payment ──
    p.append(f'<h3>Cost and Payment Schedule</h3><p><b>Fee : {_mk(ex["fee"], "auto-computed Fee")}</b></p>')
    p.append(f"<p>{_esc(PAYMENT_INTRO_AGENCY if is_agency else PAYMENT_INTRO)}</p>")
    sched_rows = "".join(
        f"<tr><td>{_esc(r[0])}</td><td class=\"num\">{_mk(r[1], 'Monthly amount (auto, editable)')}</td>"
        f"<td>{_esc(r[2])}</td></tr>" for r in ex["schedule"])
    p.append('<div style="overflow-x:auto"><table><tr><th>Month</th><th>Amount</th><th>Invoice Date</th></tr>'
             + sched_rows + f'<tr><td><b>Total</b></td><td class="num"><b>{_mk(ex["fee"], "Total")}</b></td><td></td></tr></table></div>')
    p.append(f"<p>{_esc(CHANGE_ORDER_NOTE)}</p>")
    # ── OOP + signatures ──
    p.append("<h3>Out-of-pocket Expense</h3>" + "".join(
        f"<p>{_esc(par)}</p>" for par in (OOP_AGENCY if is_agency else OOP_SAMSUNG).split("\n\n")))
    left = CHEIL_ENTITY if is_agency else SAMSUNG_ENTITY
    right = ex.get("vendor_name", CHEIL_ENTITY) if is_agency else CHEIL_ENTITY
    p.append('<h3>Signatures</h3><p>IN WITNESS WHEREOF, the parties have caused this Statement of Work '
             'to be duly executed by their authorized representatives as set forth below.</p>'
             f'<div style="overflow-x:auto"><table><tr><th>{_esc(left)}</th><th>{_mk(right, "Counterpart") if is_agency else _esc(right)}</th></tr>'
             '<tr><td>Signature: ______________</td><td>Signature: ______________</td></tr>'
             '<tr><td>Name / Title / Date</td><td>Name / Title / Date</td></tr></table></div>')
    return '<div class="ex-col"><div class="ex-paper">' + "".join(p) + "</div></div>"


_EX_TOGGLE_JS = """<script>
(function(){
  var w = document.getElementById('edWrap'), b = document.getElementById('exToggle');
  if(!w || !b) return;
  function apply(off){
    w.classList.toggle('ex-hidden', off);
    b.classList.toggle('btn-ghost', off);
    b.classList.toggle('btn-secondary', !off);
  }
  var off = localStorage.getItem('sowExampleOff') === '1';
  apply(off);
  b.addEventListener('click', function(){
    off = !off;
    localStorage.setItem('sowExampleOff', off ? '1' : '0');
    apply(off);
  });
})();
</script>"""


def _sow_rows(user, data):
    rows = []
    for s in sorted(data["sows"], key=lambda x: x.get("updated", ""), reverse=True):
        t = TYPES[_sow_type(s)]
        if t["kind"] == "sow":
            _, fee, _ = _build_schedule(s)
            fee = _money(fee)
        elif t["kind"] == "est":
            _, _, tot, _ = _est_rows_computed(s)
            fee = _money(tot)
        else:
            fee = "—"
        d = s.get("direction", "samsung")
        chip = f'<span class="dir-chip dir-{d}">{"SEA" if d == "samsung" else "Agency"} · {_esc(t["label"])}</span>'
        rows.append(
            f'<div class="sow-row">'
            f'<div><div class="sow-title">{_esc(s.get("title") or s.get("project_name") or "(untitled)")}</div>'
            f'<div class="sow-meta">{chip} &nbsp;{_esc(s.get("start") or s.get("date") or "")}{(" ~ " + _esc(s.get("end"))) if s.get("end") else ""}</div></div>'
            f'<span class="sow-fee">{fee}</span>'
            f'<span style="display:flex;gap:8px">'
            f'<a class="btn btn-secondary btn-sm" href="/sow/edit?id={s["id"]}">✎ Edit</a>'
            f'<a class="btn btn-primary btn-sm" href="/sow/docx?id={s["id"]}">⬇ docx</a>'
            f'<button class="btn btn-danger btn-sm" onclick="delSow(\'{s["id"]}\')">🗑</button>'
            f'</span></div>'
        )
    return "".join(rows)


_DEL_JS = """<script>
function delSow(id){
  if(!confirm('Delete this SOW?')) return;
  fetch('/sow/delete', {method:'POST', headers:{'Content-Type':'application/x-www-form-urlencoded'}, body:'id='+encodeURIComponent(id)})
    .then(function(){ location.reload(); });
}
</script>"""


def _landing_stats(data):
    """Clickable stat cards — the top status row of the Home dashboard."""
    today = date.today()
    contracts = data.get("contracts", [])
    active = 0
    for c in contracts:
        if _lifecycle(data, c) != "active" or c.get("amends_id"):
            continue  # a chain counts once, under its base contract
        _, s, e = _chain_effective(data, c)
        if s and e and s <= today <= e:
            active += 1
    n_vendor_ppl = sum(1 for p in data.get("people", [])
                       if (p.get("affiliation") or "Cheil").strip().lower() != "cheil")
    n_ppl = len(data.get("people", []))
    cards = [
        ("/sow/contracts", active, "📎 Active contracts",
         f'<span class="sub">{len(contracts)} total on file</span>'),
        ("/sow/people", n_ppl, "👥 People",
         f'<span class="sub">{n_ppl - n_vendor_ppl} Cheil · {n_vendor_ppl} partner</span>'),
        ("/sow/vendors", len(data.get("vendors", [])), "🏢 Vendors",
         '<span class="sub">Contractor registry</span>'),
        ("/sow/docs", len(data.get("sows", [])), "📄 Documents",
         '<span class="sub">SOW · MSA · NDA · Estimation</span>'),
    ]
    return '<div class="dd-stats">' + "".join(
        f'<a class="dd-stat" href="{href}"><span class="lb">{lb}</span><b>{n}</b>{sub}</a>'
        for href, n, lb, sub in cards) + '</div>'


def _dd_year_cashflow(data, year):
    """12-slot arrays (bill to SEA, vendor payouts, salaries) for one year.
    Contract totals spread with the days/30 convention; salaries are the
    roster's monthly salary+OH (falling back to salary) held constant."""
    bill, pay = [0.0] * 12, [0.0] * 12
    for c in data.get("contracts", []):
        if _lifecycle(data, c) != "active":
            continue
        m = _contract_month_amounts(c, data)
        if not m:
            continue
        tgt = bill if c.get("side") == "sea" else pay
        for (y, mo), v in m.items():
            if y == year:
                tgt[mo - 1] += v
    sal_mo = 0.0
    for p in data.get("people", []):
        v = _num_or_none(p.get("salary_oh"))
        if v is None:
            v = _num_or_none(p.get("salary_mo"))
        if v:
            sal_mo += v
    return bill, pay, [sal_mo] * 12


def _dd_cashflow_section(data, year):
    bill, pay, sal = _dd_year_cashflow(data, year)
    if not any(bill) and not any(pay) and not any(sal):
        return ('<div class="sow-meta" style="padding:26px;text-align:center">'
                'No dated contracts yet — upload one in the Contracts tab to '
                'light up the cashflow view.</div>')
    net = [b - p - s for b, p, s in zip(bill, pay, sal)]

    # scope datasets: 전체 + active SEA 그룹별 (salaries are company-wide, so
    # per-contract scopes show Net(margin) without the salary row)
    def _year_slice(mmap):
        out = [0.0] * 12
        for (y, mo), v in (mmap or {}).items():
            if y == year:
                out[mo - 1] += v
        return out

    scopes = [("all", "All contracts", bill, pay, sal, True)]
    for sea, kids in _contract_groups(data)[0]:
        if _lifecycle(data, sea) != "active":
            continue
        b = _year_slice(_contract_month_amounts(sea, data))
        for a in _amendments_of(data, sea):
            ab = _year_slice(_contract_month_amounts(a, data))
            b = [x + y for x, y in zip(b, ab)]
        p12 = [0.0] * 12
        for k in kids:
            if _lifecycle(data, k) == "active":
                kb = _year_slice(_contract_month_amounts(k, data))
                p12 = [a + x for a, x in zip(p12, kb)]
        if any(b) or any(p12):
            scopes.append((sea["id"],
                           sea.get("project_name") or sea.get("filename") or sea["id"],
                           b, p12, [0.0] * 12, False))

    def tile(lb, v, cls=""):
        return (f'<div class="dd-stat" style="cursor:default">'
                f'<span class="lb">{lb}</span>'
                f'<b class="{cls}" style="font-size:1.25rem">{_money(v)}</b>'
                f'<span class="sub">{year} full year</span></div>')

    tiles = ('<div class="dd-stats">'
             + tile("🔵 Billed to SEA", sum(bill), "cash-pos")
             + tile("🟠 Vendor payouts", sum(pay), "cash-pay")
             + tile("💼 Salaries (roster)", sum(sal))
             + tile("∑ Net after salaries", sum(net),
                    "cash-pos" if sum(net) >= 0 else "cash-neg")
             + '</div>')

    def row(lb, vals, cls=""):
        cells = "".join(f'<td class="num {cls}">{_money(v)}</td>' for v in vals)
        tot_cls = cls or ("pos" if sum(vals) >= 0 else "neg")
        return (f'<tr><td class="pin">{lb}</td>{cells}'
                f'<td class="num tot {tot_cls}">{_money(sum(vals))}</td></tr>')

    def net_row(vals, label="Net after salaries"):
        cells = "".join(
            f'<td class="num {"pos" if v >= 0 else "neg"}">{_money(v)}</td>' for v in vals)
        return (f'<tr class="net"><td class="pin">{label}</td>{cells}'
                f'<td class="num tot {"pos" if sum(vals) >= 0 else "neg"}">{_money(sum(vals))}</td></tr>')

    def table(headers, b, p, s, with_sal):
        n = [x - y - z for x, y, z in zip(b, p, s)]
        head = ('<tr><th class="pin"></th>'
                + "".join(f'<th>{h}</th>' for h in headers)
                + f'<th>{year}</th></tr>')
        return (f'<div class="cf-wrap"><table class="cf-table"><thead>{head}</thead><tbody>'
                + row("🔵 Billed to SEA", b, "bill")
                + row("🟠 Vendor payouts", p, "pay")
                + (row("💼 Salaries", s) if with_sal else "")
                + net_row(n, "Net after salaries" if with_sal else "Net (margin)")
                + '</tbody></table></div>')

    q = lambda arr, i: sum(arr[i * 3:(i + 1) * 3])
    panes, opts = [], []
    qh = ["Q1", "Q2", "Q3", "Q4"]
    mh = [date(year, m, 1).strftime("%b") for m in range(1, 13)]
    for sid, label, b, p12, s12, with_sal in scopes:
        opts.append(f'<option value="{sid}">{_esc(label)}</option>')
        panes.append(
            f'<div class="cf-pane" data-scope="{sid}" data-freq="q">'
            + table(qh, [q(b, i) for i in range(4)], [q(p12, i) for i in range(4)],
                    [q(s12, i) for i in range(4)], with_sal) + '</div>'
            f'<div class="cf-pane" data-scope="{sid}" data-freq="m" style="display:none">'
            + table(mh, b, p12, s12, with_sal) + '</div>')
    return f"""
{tiles}
<div style="display:flex;align-items:center;gap:10px;margin:18px 0 10px;flex-wrap:wrap">
  <h2 style="font-size:1rem;font-weight:800;margin:0">📅 {year} Billing vs payout</h2>
  <select id="cfScope" class="slot" style="max-width:340px" onchange="cfRender()">{''.join(opts)}</select>
  <span style="flex:1"></span>
  <span class="dd-tabs" style="padding:2px">
    <button class="dd-tab active" id="cfQBtn" onclick="cfFreq='q';cfRender()">Quarterly</button>
    <button class="dd-tab" id="cfMBtn" onclick="cfFreq='m';cfRender()">Monthly</button>
  </span>
</div>
{''.join(panes)}
<div class="sow-meta" style="margin-top:8px">Contract totals spread evenly across each term (days/30 partial months) · salaries = roster monthly salary+OH, shown in the All-contracts scope only · an amended document stops where its amendment takes effect; cancelled ones are excluded.</div>
<script>
var cfFreq='q';
function cfRender(){{
  var scope=document.getElementById('cfScope').value;
  document.querySelectorAll('.cf-pane').forEach(function(el){{
    el.style.display=(el.dataset.scope===scope&&el.dataset.freq===cfFreq)?'':'none';
  }});
  document.getElementById('cfQBtn').classList.toggle('active',cfFreq==='q');
  document.getElementById('cfMBtn').classList.toggle('active',cfFreq==='m');
}}
cfRender();
</script>"""


def _render_landing(user):
    data = _load(user)
    _migrate_contract_texts(user, data)
    year = date.today().year
    todos = _contract_todos(data)
    todo_html = ""
    if todos:
        items = "".join(
            f'<a href="/sow/contracts?newc={c["id"]}" class="dd-todo-item">'
            f'<span>⚠️</span><span style="flex:1"><b>{_esc(c.get("project_name") or c.get("filename") or "(untitled contract)")}</b>'
            f' — {reason}</span><span style="color:var(--accent);font-weight:700;white-space:nowrap">Open →</span></a>'
            for c, reason in todos)
        todo_html = (f'<div class="dd-todo"><div class="dd-todo-hd">📋 To Do '
                     f'<span class="tab-badge" style="background:var(--warn)">{len(todos)}</span>'
                     f' <span style="font-weight:400;color:var(--text-muted);font-size:.74rem">— every contract needs a confirmed SEA ↔ vendor (or Cheil-USA-self) mapping</span></div>'
                     f'{items}</div>')
    body = f"""
<h1 style="margin:8px 0 4px">Deal Desk</h1>
<p style="color:var(--text-muted);font-size:.86rem;margin-bottom:14px">Status at a glance — work happens in the tabs above.</p>
{todo_html}
{_landing_stats(data)}
{_dd_cashflow_section(data, year)}"""
    return _shell(user, "Deal Desk", body, tab="home")


def _render_contracts_page(user):
    data = _load(user)
    _migrate_contract_texts(user, data)
    body = f"""
<h1 style="margin:8px 0 4px">📎 Contracts</h1>
<p style="color:var(--text-muted);font-size:.86rem;margin-bottom:14px">Executed contracts, grouped by SEA ↔ Cheil deal — vendor contracts align underneath with the margin rolled up.</p>
{_render_contracts_section(user, data)}
<div class="cmodal-ov" id="cmodalOv"><div class="cmodal" id="cmodal"></div></div>
{_CTR_JS}"""
    return _shell(user, "Contracts", body, tab="contracts")


def _render_docs_page(user):
    data = _load(user)
    rows = _sow_rows(user, data)
    body = f"""
<div style="display:flex;align-items:center;gap:12px;margin:8px 0 14px;flex-wrap:wrap">
  <h1 style="margin:0">📄 Documents</h1>
  <span style="flex:1"></span>
  <a class="btn btn-secondary btn-sm" href="/sow/types?dir=sea">🔵 New with SEA</a>
  <a class="btn btn-secondary btn-sm" href="/sow/types?dir=agy">🟠 New with Vendor</a>
</div>
<div class="sow-list">{rows or '<div class="sow-meta" style="padding:36px;text-align:center">No documents yet — start one from the buttons above.</div>'}</div>
{_DEL_JS}"""
    return _shell(user, "Documents", body, tab="docs")


def _person_docs(data, person):
    """Documents tied to this person: explicit links + auto-detected
    (estimate rows by person_id, SOW resource rows by name)."""
    linked = set(person.get("linked_sows") or [])
    auto = set()
    pname = (person.get("name") or "").strip().lower()
    for s in data["sows"]:
        kind = TYPES.get(_sow_type(s), {}).get("kind")
        if kind == "est":
            if any(r.get("person_id") == person["id"] for r in s.get("rows", [])):
                auto.add(s["id"])
        elif kind == "sow":
            for r in s.get("resources", []):
                nm = (r.get("profile") or r.get("name") or "").strip().lower()
                if pname and nm == pname:
                    auto.add(s["id"])
                    break
    ids = linked | auto
    return [s for s in data["sows"] if s["id"] in ids], auto


def _person_doc_count(data, person):
    docs, _ = _person_docs(data, person)
    return len(docs)


def _fmt_money_cell(v):
    n = _num_or_none(v)
    return _money(n) if n is not None else (_esc(v) if v else "–")


def _money_input(v):
    """Same figure, formatted for an editable cell — money stays scannable
    while it is being edited; blank stays blank (no em dash to delete)."""
    n = _num_or_none(v)
    return _money(n) if n is not None else (str(v or ""))


def _render_people_by_contract(data):
    """People grouped per contract (active first) with each person's monthly
    money and a totals row — the contract-level cost breakdown (강프로
    2026-07-24). Membership = person.linked_contracts."""
    today = date.today()

    def is_active(c):
        if _lifecycle(data, c) != "active":
            return False
        s = _parse_any_date(c.get("period_start"))
        e = _effective_end(data, c)
        return bool(s and e and s <= today <= e)

    contracts = sorted([c for c in data.get("contracts", [])
                        if not c.get("amends_id")],
                       key=lambda c: (not is_active(c),
                                      (c.get("project_name") or "").lower()))
    assigned = set()
    sections = []
    for c in contracts:
        members = [p for p in data.get("people", [])
                   if c["id"] in (p.get("linked_contracts") or [])]
        active = is_active(c)
        if not members and not active:
            continue
        assigned.update(p["id"] for p in members)
        label, chip, _, icon = _SIDE_META.get(c.get("side"), _SIDE_META["sea"])
        badge = ('<span class="ctr-chip pos">● Active</span>' if active
                 else '<span class="ctr-chip">○ Inactive</span>')
        period = ""
        if c.get("period_start") or c.get("period_end"):
            period = f'{_esc(c.get("period_start") or "…")} ~ {_esc(c.get("period_end") or "…")}'
        rows, sell_sum, cost_sum = [], 0.0, 0.0
        for p in sorted(members, key=lambda x: (x.get("name") or "").lower()):
            sell = _num_or_none(p.get("sell_mo"))
            cost = _num_or_none(p.get("cost_mo"))
            cost_src = "contract"
            if cost is None:
                cost = _num_or_none(p.get("salary_oh")) or _num_or_none(p.get("salary_mo"))
                cost_src = "salary"
            sell_sum += sell or 0.0
            cost_sum += cost or 0.0
            rows.append(
                f'<tr><td class="pin"><a href="/sow/person?id={p["id"]}" '
                f'style="color:var(--text);text-decoration:none;font-weight:700">{_esc(p.get("name"))}</a></td>'
                f'<td>{_esc(p.get("role_title") or "–")}</td>'
                f'<td>{_esc(p.get("affiliation") or "Cheil")}</td>'
                f'<td class="num bill">{_money(sell) if sell is not None else "–"}</td>'
                f'<td class="num pay">{_money(cost) if cost is not None else "–"}'
                + (f' <span style="font-size:.6rem;color:var(--text-muted)">{cost_src}</span>' if cost is not None else "")
                + f'</td>'
                f'<td class="num {"pos" if (sell or 0) - (cost or 0) >= 0 else "neg"}">'
                f'{_money((sell or 0) - (cost or 0)) if (sell is not None or cost is not None) else "–"}</td></tr>')
        net = sell_sum - cost_sum
        total_row = (
            f'<tr class="net"><td class="pin">Total / month</td><td></td><td></td>'
            f'<td class="num bill">{_money(sell_sum)}</td>'
            f'<td class="num pay">{_money(cost_sum)}</td>'
            f'<td class="num {"pos" if net >= 0 else "neg"}">{_money(net)}</td></tr>'
            if rows else "")
        body_rows = "".join(rows) or (
            '<tr><td colspan="6" style="text-align:center;color:var(--text-muted);'
            'padding:14px;font-size:.78rem">No people linked yet — open a person '
            'and tick this contract under "Linked contracts".</td></tr>')
        sections.append(f"""
<div class="ctr-group" style="margin-bottom:16px">
  <div class="ctr-group-hd">
    <span class="dir-chip {chip}">{icon} {label}</span>
    <span class="ctr-group-name">{_esc(c.get("project_name") or c.get("filename") or "(untitled contract)")}</span>
    {badge}
    <span class="sow-meta">{period}{(" · " + _esc(c.get("amount"))) if c.get("amount") else ""}</span>
  </div>
  <div class="cf-wrap"><table class="cf-table" style="min-width:760px">
    <thead><tr><th class="pin">Name</th><th>Role</th><th>Affiliation</th>
      <th>Selling/mo</th><th>Cost/mo</th><th>Net/mo</th></tr></thead>
    <tbody>{body_rows}{total_row}</tbody>
  </table></div>
</div>""")
    rest = [p for p in data.get("people", []) if p["id"] not in assigned]
    if rest:
        chips = "".join(
            f'<a href="/sow/person?id={p["id"]}" class="btn btn-secondary btn-sm" '
            f'style="margin:3px">{_esc(p.get("name"))}</a>' for p in
            sorted(rest, key=lambda x: (x.get("name") or "").lower()))
        sections.append(
            f'<div class="ctr-orphans"><div class="ctr-orphan-hd">Not linked to any contract '
            f'({len(rest)}) <span>· open a person and tick their contract</span></div>'
            f'<div>{chips}</div></div>')
    return "".join(sections) or (
        '<div class="sow-meta" style="padding:30px;text-align:center">'
        'No contracts on file yet.</div>')


def _render_people(user, saved=False, view="roster"):
    """Excel-style summary grid (강프로 2026-07-22): one row per person, the
    3 money axes as banded column groups, first column pinned, horizontal
    scroll inside the wrapper per the design-system reference-table rule.
    view='contract' swaps the grid for the per-contract cost breakdown."""
    data = _load(user)

    if view == "contract":
        body = f"""
<div style="display:flex;align-items:center;gap:12px;margin:8px 0 14px;flex-wrap:wrap">
  <h1 style="margin:0">👥 People</h1>
  <span class="dd-tabs" style="padding:2px">
    <a class="dd-tab" href="/sow/people">Roster</a>
    <a class="dd-tab active" href="/sow/people?view=contract">By contract</a>
  </span>
  <span style="flex:1"></span>
  <a class="btn btn-primary btn-sm" href="/sow/person">+ Add person</a>
</div>
<p style="color:var(--text-muted);font-size:.86rem">Each contract's team and monthly money — selling vs cost per person (cost falls back to salary+OH for Cheil staff), active contracts first. Link people from their profile page.</p>
{_render_people_by_contract(data)}"""
        return _shell(user, "People", body, wide=True, tab="people")

    def mo_auto(hr_v, mo_v):
        if mo_v:
            return _fmt_money_cell(mo_v)
        n = _num_or_none(hr_v)
        return f"${n * 168:,.0f}" if n is not None else "–"

    def cell(person, key, num=False, ph=""):
        """An editable roster cell — saves on blur, no page reload. Typing must
        never rebuild the input, so nothing here re-renders on input."""
        return (f'<input class="pp-cell{" num" if num else ""}" '
                f'data-pid="{person["id"]}" data-field="{key}" '
                f'value="{_esc(_money_input(person.get(key)) if num else person.get(key))}" '
                f'placeholder="{_esc(ph if ph and ph != "–" else "–")}"'
                f'{" inputmode=decimal" if num else ""}>')

    rows = []
    for p in data["people"]:
        docs, _ = _person_docs(data, p)
        ebita, ebita_auto = _person_ebita(p)
        aff = p.get("affiliation") or "Cheil"
        aff_chip = (f'<span class="dir-chip dir-samsung">{_esc(aff)}</span>' if aff == "Cheil"
                    else f'<span class="dir-chip dir-agency">{_esc(aff)}</span>')
        if ebita is not None:
            col = "var(--success)" if ebita >= 0 else "var(--danger)"
            ebita_html = (f'<span style="color:{col};font-weight:700">{_money(ebita)}</span>'
                          + ('<span style="font-size:.6rem;color:var(--text-muted)"> a</span>' if ebita_auto else ''))
        else:
            ebita_html = "–"
        doc_html = (f'<a href="/sow/person?id={p["id"]}" style="color:var(--accent);text-decoration:none">'
                    f'{len(docs)} 📄</a>' if docs else "–")
        rows.append(
            '<tr>'
            f'<td class="pp-pin"><a href="/sow/person?id={p["id"]}" style="font-weight:700;color:var(--text);text-decoration:none">{_esc(p.get("name"))}</a><br>{aff_chip}</td>'
            f'<td>{cell(p, "project")}</td>'
            f'<td>{cell(p, "role_title")}</td>'
            f'<td class="num">{cell(p, "sell_hr", num=True)}</td>'
            f'<td class="num">{cell(p, "sell_mo", num=True, ph=mo_auto(p.get("sell_hr"), p.get("sell_mo")))}</td>'
            f'<td>{_esc(p.get("client_duration") or "–")}</td>'
            f'<td class="num">{cell(p, "client_budget", num=True)}</td>'
            f'<td>{cell(p, "client_po")}</td>'
            f'<td class="num">{cell(p, "cost_hr", num=True)}</td>'
            f'<td class="num">{cell(p, "cost_mo", num=True, ph=mo_auto(p.get("cost_hr"), p.get("cost_mo")))}</td>'
            f'<td>{_esc(p.get("partner_duration") or "–")}</td>'
            f'<td class="num">{cell(p, "partner_cost", num=True)}</td>'
            f'<td>{cell(p, "partner_po")}</td>'
            f'<td class="num">{cell(p, "salary_mo", num=True)}</td>'
            f'<td>{_esc(p.get("cheil_since") or "–")}</td>'
            f'<td class="num">{cell(p, "salary_oh", num=True)}</td>'
            f'<td class="num">{ebita_html}</td>'
            f'<td style="text-align:center">{doc_html}</td>'
            '</tr>')

    saved_banner = ('<div style="color:var(--success);font-size:.85rem;margin-bottom:12px">✓ Saved</div>'
                    if saved else "")
    body = f"""
<style>
.pp-wrap{{overflow-x:auto;overscroll-behavior-x:contain;border:1px solid var(--border);border-radius:var(--radius-lg);background:var(--surface)}}
.pp-table{{border-collapse:collapse;font-size:.8rem;min-width:1750px;width:100%}}
.pp-table th{{border:1px solid var(--border);padding:6px 9px;font-size:.62rem;text-transform:uppercase;letter-spacing:.04em;color:var(--text-muted);background:var(--surface-2);text-align:left;white-space:nowrap}}
.pp-table th.grp{{text-align:center;font-weight:800;color:var(--text)}}
.pp-table th.g-sell{{background:rgba(56,189,248,.10)}}
.pp-table th.g-cost{{background:rgba(251,146,60,.10)}}
.pp-table th.g-emp{{background:rgba(52,211,153,.10)}}
.pp-table th.g-ebita{{background:rgba(129,140,248,.12)}}
.pp-table td{{border:1px solid var(--border);padding:7px 9px;white-space:nowrap}}
.pp-table td.num{{text-align:right;font-variant-numeric:tabular-nums}}
.pp-table .pp-pin,.pp-table th:first-child{{position:sticky;left:0;background:var(--surface);z-index:2;min-width:150px;box-shadow:2px 0 6px rgba(0,0,0,.25)}}
.pp-table thead th:first-child{{background:var(--surface-2);z-index:3}}
.pp-table tbody tr:hover td{{background:var(--surface-2)}}
.pp-table tbody tr:hover td.pp-pin{{background:var(--surface-2)}}
.pp-table td:has(.pp-cell){{padding:0}}
.pp-cell{{width:100%;min-width:70px;background:transparent;border:1px solid transparent;
  border-radius:var(--radius-sm);color:var(--text);font:inherit;padding:7px 9px}}
.pp-cell::placeholder{{color:var(--text-dim)}}
.pp-cell.num{{text-align:right;font-variant-numeric:tabular-nums}}
.pp-cell:hover{{border-color:var(--border-bright)}}
.pp-cell:focus{{outline:none;border-color:var(--accent);background:var(--bg-deep)}}
.pp-cell.saving{{border-color:var(--warn)}}
.pp-cell.saved{{border-color:var(--success)}}
</style>
<div style="display:flex;align-items:center;gap:12px;margin:8px 0 4px;flex-wrap:wrap">
  <h1 style="margin:0">👥 People</h1>
  <span class="dd-tabs" style="padding:2px">
    <a class="dd-tab{' active' if view != 'contract' else ''}" href="/sow/people">Roster</a>
    <a class="dd-tab{' active' if view == 'contract' else ''}" href="/sow/people?view=contract">By contract</a>
  </span>
  <span style="flex:1"></span>
  <a class="btn btn-primary btn-sm" href="/sow/person">+ Add person</a>
</div>
<p style="color:var(--text-muted);font-size:.86rem">One row per person — selling side, cost side, employee comp and EBITA in one sweep (scroll sideways; the name column stays pinned). Click a name for the full profile and linked SOWs.</p>
{saved_banner}
<div class="pp-wrap"><table class="pp-table">
  <thead>
    <tr>
      <th rowspan="2">Name / Affiliation</th><th rowspan="2">Project / SOW</th><th rowspan="2">Role · Title</th>
      <th class="grp g-sell" colspan="5">Client → Cheil</th>
      <th class="grp g-cost" colspan="5">Cheil → Partner</th>
      <th class="grp g-emp" colspan="3">Cheil employee</th>
      <th class="grp g-ebita" rowspan="2">EBITA<br>by Cheil</th>
      <th rowspan="2">Docs</th>
    </tr>
    <tr>
      <th class="g-sell">Selling/hr</th><th class="g-sell">Selling/mo</th><th class="g-sell">Duration</th><th class="g-sell">Budget</th><th class="g-sell">PO</th>
      <th class="g-cost">Contract/hr</th><th class="g-cost">Contract/mo</th><th class="g-cost">Duration</th><th class="g-cost">Cost</th><th class="g-cost">PO</th>
      <th class="g-emp">Salary/mo</th><th class="g-emp">Since</th><th class="g-emp">Salary+OH</th>
    </tr>
  </thead>
  <tbody>{''.join(rows) or '<tr><td colspan="19" style="text-align:center;padding:30px;color:var(--text-muted)">No people yet.</td></tr>'}</tbody>
</table></div>
<script>
// Inline roster edit: save the cell that changed, nothing else. No re-render,
// so the caret never jumps out from under the user.
document.addEventListener('change', function(e){{
  var c = e.target.closest('.pp-cell'); if(!c) return;
  c.classList.remove('saved'); c.classList.add('saving');
  var b = 'id='+encodeURIComponent(c.dataset.pid)
        + '&field='+encodeURIComponent(c.dataset.field)
        + '&value='+encodeURIComponent(c.value);
  fetch('/sow/person/cell',{{method:'POST',
      headers:{{'Content-Type':'application/x-www-form-urlencoded'}},body:b}})
    .then(function(r){{ return r.json(); }})
    .then(function(j){{
      c.classList.remove('saving');
      if(!j.ok){{ c.style.borderColor='var(--danger)'; return; }}
      if(typeof j.display === 'string' && j.display !== c.value) c.value = j.display;
      c.classList.add('saved');
      setTimeout(function(){{ c.classList.remove('saved'); }}, 1200);
    }})
    .catch(function(){{ c.classList.remove('saving'); c.style.borderColor='var(--danger)'; }});
}});
// Enter commits and moves down the column — the way a spreadsheet behaves
document.addEventListener('keydown', function(e){{
  if(e.key !== 'Enter') return;
  var c = e.target.closest('.pp-cell'); if(!c) return;
  e.preventDefault(); c.blur();
  var col = [].slice.call(document.querySelectorAll(
    '.pp-cell[data-field="'+c.dataset.field+'"]'));
  var nx = col[col.indexOf(c)+1]; if(nx) nx.focus();
}});
</script>"""
    return _shell(user, "People", body, wide=True, tab="people")


def _render_person_detail(user, person, saved=False):
    data = _load(user)
    p = person or {}
    docs, auto_ids = _person_docs(data, p) if p.get("id") else ([], set())
    vend_names = [v["name"] for v in data["vendors"]]
    aff_opts = ["Cheil"] + vend_names + ["TBD"]
    cur_aff = p.get("affiliation") or "Cheil"
    if cur_aff not in aff_opts:
        aff_opts.append(cur_aff)
    aff_sel = "".join(f'<option{" selected" if a == cur_aff else ""}>{_esc(a)}</option>' for a in aff_opts)

    def fld(label, name, ph="", typ="text", wide=False, dl=""):
        return (f'<div class="f-cell{" f-wide" if wide else ""}">'
                f'<div style="font-size:.64rem;text-transform:uppercase;letter-spacing:.05em;color:var(--text-muted);margin-bottom:3px">{label}</div>'
                f'<input class="slot" style="width:100%" type="{typ}" name="{name}" '
                f'value="{_esc(p.get(name))}" placeholder="{_esc(ph)}"'
                f'{f" list={dl}" if dl else ""}></div>')

    def _datalist(dl_id, key):
        """Values already in the roster — typing the 13th person should be
        picking, not retyping (강프로 2026-07-27)."""
        vals = sorted({(x.get(key) or "").strip() for x in data.get("people", [])
                       if (x.get(key) or "").strip()})
        return (f'<datalist id="{dl_id}">'
                + "".join(f'<option value="{_esc(v)}">' for v in vals) + '</datalist>')

    ebita, ebita_auto = _person_ebita(p)
    ebita_hint = (f'auto = Client budget − Partner cost = {_money(ebita)}'
                  if (ebita is not None and ebita_auto) else 'auto needs Client budget + Partner cost — or type a value')

    # linked docs: explicit links are checkboxes over every SOW/estimate doc;
    # auto-detected ones are pre-noted.
    doc_opts = []
    linked = set(p.get("linked_sows") or [])
    for s in sorted(data["sows"], key=lambda x: x.get("updated", ""), reverse=True):
        k = TYPES.get(_sow_type(s), {})
        if k.get("kind") not in ("sow", "est"):
            continue
        checked = " checked" if s["id"] in linked else ""
        auto = ' <span style="font-size:.64rem;color:var(--success)">auto-linked</span>' if s["id"] in auto_ids else ""
        doc_opts.append(
            f'<label style="display:flex;align-items:center;gap:8px;padding:5px 0;font-size:.82rem;cursor:pointer">'
            f'<input type="checkbox" name="linked_sows" value="{s["id"]}"{checked}>'
            f'{k.get("icon","")} {_esc(s.get("title") or "(untitled)")}'
            f'<span style="color:var(--text-muted);font-size:.72rem">{_esc(s.get("start") or "")}{(" ~ " + _esc(s.get("end"))) if s.get("end") else ""}</span>'
            f'{auto}'
            f'<a href="/sow/edit?id={s["id"]}" style="margin-left:auto;color:var(--accent);font-size:.74rem;text-decoration:none">open →</a></label>')

    # linked contracts: explicit checkboxes — powers the People "By contract"
    # breakdown (강프로 2026-07-24)
    ctr_opts = []
    linked_c = set(p.get("linked_contracts") or [])
    _today = date.today()
    for c in data.get("contracts", []):
        s0 = _parse_any_date(c.get("period_start"))
        e0 = _effective_end(data, c)
        act = ('<span style="font-size:.64rem;color:var(--success)">● Active</span>'
               if (s0 and e0 and s0 <= _today <= e0) else
               '<span style="font-size:.64rem;color:var(--text-muted)">○ Inactive</span>')
        _, chip, _, icon = _SIDE_META.get(c.get("side"), _SIDE_META["sea"])
        checked = " checked" if c["id"] in linked_c else ""
        ctr_opts.append(
            f'<label style="display:flex;align-items:center;gap:8px;padding:5px 0;font-size:.82rem;cursor:pointer">'
            f'<input type="checkbox" name="linked_contracts" value="{c["id"]}"{checked}>'
            f'{icon} {_esc(c.get("project_name") or c.get("filename") or "(untitled contract)")}'
            f'<span style="color:var(--text-muted);font-size:.72rem">{_esc(c.get("vendor") or c.get("client") or "")}</span>'
            f'{act}</label>')

    # 투입 프로젝트 = contract-driven dropdown (강프로 2026-07-24): one option
    # per CONTRACT (several vendor deals can share a project name — the party
    # tells them apart). Picking one auto-links that contract (and its SEA
    # parent) via the checkboxes below; the server mirrors the same rule.
    cur_proj = (p.get("project") or "").strip()
    linked_now = set(p.get("linked_contracts") or [])
    sel_val, proj_items = None, []
    for c in data.get("contracts", []):
        pn = (c.get("project_name") or "").strip()
        if not pn or c.get("amends_id"):
            continue
        _, _, _, icon2 = _SIDE_META.get(c.get("side"), _SIDE_META["sea"])
        party = c.get("vendor") or c.get("client") or ""
        val = "ctr:" + c["id"]
        if sel_val is None and pn == cur_proj and (not linked_now or c["id"] in linked_now):
            sel_val = val
        period = " ~ ".join(x for x in (c.get("period_start"), c.get("period_end")) if x)
        proj_items.append(
            f'<option value="{val}" data-cid="{c["id"]}" '
            f'data-parent="{_esc(c.get("linked_id") or "")}" '
            f'data-period="{_esc(period)}" data-side="{_esc(c.get("side") or "sea")}">'
            f'{icon2} {(_esc(party) + " — ") if party else ""}{_esc(pn)}</option>')
    custom = ""
    if cur_proj and sel_val is None:
        custom = f'<option value="__keep__">{_esc(cur_proj)} (manual)</option>'
        sel_val = "__keep__"
    proj_opts = ('<option value="">— none —</option>' + custom + "".join(proj_items))         .replace(f'value="{sel_val}"', f'value="{sel_val}" selected', 1)         if sel_val else '<option value="" selected>— none —</option>' + "".join(proj_items)

    saved_banner = ('<div style="color:var(--success);font-size:.85rem;margin-bottom:12px">✓ Saved</div>'
                    if saved else "")
    body = f"""
<div style="display:flex;align-items:center;gap:12px;margin:8px 0 16px">
  <a class="btn btn-ghost btn-sm" href="/sow/people">←</a>
  <h1 style="margin:0">{_esc(p.get('name') or 'New person')}</h1>
</div>
{saved_banner}
<form method="post" action="/sow/person/save">
<input type="hidden" name="id" value="{_esc(p.get('id') or '')}">
<style>.f-grid3{{display:grid;grid-template-columns:repeat(auto-fit,minmax(170px,1fr));gap:12px}}.f-wide{{grid-column:1/-1}}</style>

<div class="sow-card">
  <div class="sec-title" style="font-size:.8rem;font-weight:800;margin-bottom:12px">Basics</div>
  <div class="f-grid3">
    {fld('Name', 'name', 'required')}
    <div class="f-cell"><div style="font-size:.64rem;text-transform:uppercase;letter-spacing:.05em;color:var(--text-muted);margin-bottom:3px">Affiliation</div>
      <select class="slot" style="width:100%" name="affiliation">{aff_sel}</select></div>
    {fld('Role / Title', 'role_title', 'e.g. Sr. Developer', dl='dlRole')}
    <div class="f-cell f-wide"><div style="font-size:.64rem;text-transform:uppercase;letter-spacing:.05em;color:var(--text-muted);margin-bottom:3px">Assigned project / SOW</div>
      <select class="slot" style="width:100%" name="project_pick" onchange="projPick(this)">{proj_opts}</select>
      <div style="font-size:.68rem;color:var(--text-muted);margin-top:3px">Picking a contract project also ticks it under Linked contracts below (a vendor contract links its SEA deal too).</div></div>
    {fld('Location', 'location', 'US / India', dl='dlLoc')}
    {_datalist('dlRole', 'role_title')}{_datalist('dlLoc', 'location')}
  </div>
</div>

<div class="sow-card">
  <div class="sec-title" style="font-size:.8rem;font-weight:800;margin-bottom:12px">Client → Cheil <span style="font-weight:500;color:var(--text-muted)">(selling side)</span></div>
  <div class="f-grid3">
    {fld('Selling Rate (hr)', 'sell_hr', '$/h')}
    {fld('Selling Rate (Mo)', 'sell_mo', 'auto: hr × 168')}
    {fld('Client–Cheil Duration', 'client_duration', 'e.g. Apr – Oct 2026')}
    {fld('Contracted Budget', 'client_budget', '$')}
    {fld('Client–Cheil PO', 'client_po', 'PO #')}
  </div>
</div>

<div class="sow-card">
  <div class="sec-title" style="font-size:.8rem;font-weight:800;margin-bottom:12px">Cheil → Partner <span style="font-weight:500;color:var(--text-muted)">(cost side, vendor personnel)</span></div>
  <div class="f-grid3">
    {fld('Contract Rate (hr)', 'cost_hr', '$/h')}
    {fld('Contract Rate (Mo)', 'cost_mo', 'auto: hr × 168')}
    {fld('Cheil–Partner Duration', 'partner_duration')}
    {fld('Contracted Cost', 'partner_cost', '$')}
    {fld('Cheil–Partner PO (if any)', 'partner_po')}
  </div>
</div>

<div class="sow-card">
  <div class="sec-title" style="font-size:.8rem;font-weight:800;margin-bottom:12px">Cheil employee</div>
  <div class="f-grid3">
    {fld('Salary (Mo)', 'salary_mo', '$')}
    {fld('Cheil Since', 'cheil_since', 'e.g. 2024-03')}
    {fld('Cheil Salary + OH', 'salary_oh', '$')}
  </div>
</div>

<div class="sow-card">
  <div class="sec-title" style="font-size:.8rem;font-weight:800;margin-bottom:12px">EBITA by Cheil</div>
  <div class="f-grid3">
    {fld('EBITA (manual override)', 'ebita', ebita_hint)}
  </div>
</div>

<div class="sow-card">
  <div class="sec-title" style="font-size:.8rem;font-weight:800;margin-bottom:12px">Accounts &amp; equipment</div>
  <div class="f-grid3">
    {fld('Cheil.com Email', 'email_cheil')}
    {fld('Samsung.com Email', 'email_samsung')}
    {fld('PC', 'pc', 'asset / model')}
    {fld('SVPN', 'svpn', 'account / status')}
  </div>
</div>

<div class="sow-card">
  <div class="sec-title" style="font-size:.8rem;font-weight:800;margin-bottom:12px">Linked SOWs &amp; estimates</div>
  {''.join(doc_opts) or '<div style="font-size:.8rem;color:var(--text-muted)">No documents yet.</div>'}
</div>

<div class="sow-card">
  <div class="sec-title" style="font-size:.8rem;font-weight:800;margin-bottom:12px">Linked contracts</div>
  <div style="font-size:.72rem;color:var(--text-muted);margin-bottom:8px">Drives the People → By contract breakdown.</div>
  {''.join(ctr_opts) or '<div style="font-size:.8rem;color:var(--text-muted)">No contracts uploaded yet.</div>'}
</div>

<div style="display:flex;gap:10px;margin:16px 0 40px">
  <button type="submit" class="btn btn-primary btn-lg">💾 Save</button>
  {f'<button type="button" class="btn btn-secondary" onclick="dupPerson(&quot;{p.get("id")}&quot;)" title="Same rates, project and links — new name, blank personal details">⧉ Duplicate</button>' if p.get('id') else ''}
  {f'<button type="button" class="btn btn-danger" onclick="delPerson()">🗑 Delete</button>' if p.get('id') and not docs else ''}
</div>
</form>
<script>
function projPick(sel){{
  var o = sel.options[sel.selectedIndex];
  if(!o) return;
  [o.dataset.cid, o.dataset.parent].forEach(function(cid){{
    if(!cid) return;
    var cb = document.querySelector('input[name="linked_contracts"][value="'+cid+'"]');
    if(cb) cb.checked = true;
  }});
  // the contract already states the period — fill the matching duration box
  // if it is still empty (never overwrite what was typed)
  var per = o.dataset.period;
  if(per){{
    var f = document.querySelector('input[name="'
      + (o.dataset.side === 'vendor' ? 'partner_duration' : 'client_duration') + '"]');
    if(f && !f.value.trim()) f.value = per;
  }}
}}
function dupPerson(id){{
  fetch('/sow/person/duplicate',{{method:'POST',headers:{{'Content-Type':'application/x-www-form-urlencoded'}},
    body:'id='+encodeURIComponent(id)}})
    .then(function(r){{ location.href = r.url || '/sow/people'; }});
}}
function delPerson(){{
  if(!confirm('Delete this person from the roster?')) return;
  fetch('/sow/person/delete', {{method:'POST', headers:{{'Content-Type':'application/x-www-form-urlencoded'}}, body:'id={_esc(p.get("id") or "")}'}})
    .then(function(){{ location.href = '/sow/people'; }});
}}
</script>"""
    return _shell(user, p.get("name") or "New person", body)


def _est_rows_computed(sow):
    """Per-row computed figures + totals for an estimate document."""
    months = float(sow.get("months") or 0)
    out, tot_monthly, tot_total, tot_cost = [], 0.0, 0.0, 0.0
    for r in sow.get("rows", []):
        rate = float(r.get("rate") or 0)
        alloc = float(r.get("alloc") or 1)
        monthly = rate * 168
        mcost = monthly * alloc
        total = mcost * months
        vc = float(r.get("vendor_cost") or 0)
        out.append({**r, "monthly": monthly, "monthly_cost": mcost, "total": total})
        tot_monthly += mcost
        tot_total += total
        tot_cost += vc * 168 * alloc * months
    return out, round(tot_monthly, 2), round(tot_total, 2), round(tot_cost, 2)


def _build_est_xlsx(sow):
    """Estimation workbook in the executed 'Cost Estimation' layout:
    Resource | Function | Email ID | Location | Rate | Monthly | Allocation |
    Monthly Cost | <period total> — with the totals row at the bottom."""
    import openpyxl
    from openpyxl.styles import Font, Alignment, Border, Side

    rows, tot_monthly, tot_total, _ = _est_rows_computed(sow)
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Cost Estimation"
    period = sow.get("period_label") or f"{sow.get('months') or 0} months (Total)"
    headers = ["", "Resource", "Function", "Email ID", "Location", "Rate",
               "Monthly", "Allocation", "Monthly Cost", period]
    thin = Border(bottom=Side(style="thin", color="CCCCCC"))
    ws.append([])
    ws.append(headers)
    for c in ws[2]:
        c.font = Font(bold=True)
        c.border = thin
    for r in rows:
        ws.append(["", r.get("name") or "", r.get("function") or "",
                   r.get("email") or "", r.get("location") or "",
                   float(r.get("rate") or 0), r["monthly"],
                   float(r.get("alloc") or 1), r["monthly_cost"], r["total"]])
    ws.append(["", "", "", "", "", "", "", "", tot_monthly, tot_total])
    for c in ws[ws.max_row]:
        c.font = Font(bold=True)
    widths = [3, 22, 10, 26, 10, 8, 10, 11, 13, 16]
    for i, w in enumerate(widths, 1):
        ws.column_dimensions[openpyxl.utils.get_column_letter(i)].width = w
    for col in ("F", "G", "I", "J"):
        for cell in ws[col]:
            cell.number_format = "#,##0"
            cell.alignment = Alignment(horizontal="right")
    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


_EST_JS = """<script>
var EROWS = __ROWS__;
var CFG = __CFG__;
function $id(x){ return document.getElementById(x); }
function money(x){
  if(isNaN(x)) return '$0';
  var r = Math.round(x*100)/100;
  return '$' + (Math.abs(r-Math.round(r))<0.005 ? Math.round(r).toLocaleString('en-US')
              : r.toLocaleString('en-US', {minimumFractionDigits:2}));
}
function months(){ return parseFloat($id('fMonths').value) || 0; }
function renderRows(){
  var tb = $id('estBody'); tb.innerHTML = '';
  EROWS.forEach(function(r, i){
    var tr = document.createElement('tr');
    var opts = '<option value="">— pick —</option>' + CFG.people.map(function(p){
      return '<option value="' + p.id + '"' + (r.person_id === p.id ? ' selected' : '') + '>' + p.name + '</option>';
    }).join('');
    tr.innerHTML =
      '<td><select class="slot" data-k="person_id" data-i="' + i + '">' + opts + '</select></td>'
      + cell(r,'function',i) + cell(r,'location',i) + cell(r,'rate',i)
      + '<td class="num c-monthly"></td>' + cell(r,'alloc',i)
      + '<td class="num c-mcost"></td><td class="num c-total"></td>'
      + '<td style="border:none"><button type="button" class="row-del" onclick="rmRow('+i+')">✕</button></td>';
    tb.appendChild(tr);
  });
  paint();
}
function cell(r, k, i){
  var v = r[k] != null ? String(r[k]).replace(/"/g,'&quot;') : '';
  return '<td><input class="slot" data-k="'+k+'" data-i="'+i+'" value="'+v+'"></td>';
}
function paint(){
  var m = months(), totM = 0, totT = 0, totC = 0, anyVc = false;
  document.querySelectorAll('#estBody tr').forEach(function(tr, i){
    var r = EROWS[i] || {};
    var monthly = (parseFloat(r.rate)||0)*168;
    var mcost = monthly*(parseFloat(r.alloc)||1);
    var total = mcost*m;
    totM += mcost; totT += total;
    var vc = parseFloat(r.vendor_cost)||0;
    if(vc){ anyVc = true; totC += vc*168*(parseFloat(r.alloc)||1)*m; }
    tr.querySelector('.c-monthly').textContent = money(monthly);
    tr.querySelector('.c-mcost').textContent = money(mcost);
    tr.querySelector('.c-total').textContent = money(total);
  });
  $id('totMonthly').textContent = money(totM);
  $id('totTotal').textContent = money(totT);
  var mg = $id('marginLine');
  if(anyVc && totT){
    mg.style.display = '';
    mg.textContent = 'internal cost ' + money(totC) + ' \\u00b7 margin ' + money(totT - totC)
      + ' (' + Math.round((totT - totC)/totT*100) + '%)';
  } else mg.style.display = 'none';
}
function addRow(){ EROWS.push({alloc: 1}); renderRows(); }
function rmRow(i){ EROWS.splice(i,1); renderRows(); }
document.addEventListener('input', function(e){
  var k = e.target.dataset && e.target.dataset.k;
  if(!k) { if(e.target.id === 'fMonths') paint(); return; }
  EROWS[parseInt(e.target.dataset.i)][k] = e.target.value;
  paint();
});
document.addEventListener('change', function(e){
  if(e.target.dataset && e.target.dataset.k === 'person_id'){
    var i = parseInt(e.target.dataset.i);
    var p = CFG.people.find(function(x){ return x.id === e.target.value; });
    if(p){
      EROWS[i] = {person_id: p.id, name: p.name, function: p.function, email: p.email,
                  location: p.location, rate: p.rate, vendor_cost: p.vendor_cost,
                  alloc: EROWS[i].alloc || 1};
      renderRows();
    }
  }
});
document.getElementById('sowForm').addEventListener('submit', function(){
  $id('rowsJson').value = JSON.stringify(EROWS.filter(function(r){
    return r.person_id || String(r.rate||'').trim();
  }));
});
if(!EROWS.length) EROWS.push({alloc: 1});
renderRows();
</script>"""


def _render_est_editor(user, sow, type_key, saved=False):
    data = _load(user)
    t = TYPES[type_key]
    saved_note = ('<span style="color:var(--success);font-size:.8rem;font-weight:700">✓ Saved</span>'
                  if saved else "")
    people = [{"id": p["id"], "name": p.get("name") or "", "function": p.get("role_title") or "",
               "email": p.get("email_samsung") or p.get("email_cheil") or "",
               "location": p.get("location") or "",
               "rate": p.get("sell_hr") or "", "vendor_cost": p.get("cost_hr") or ""}
              for p in data["people"]]
    cfg = {"people": people}
    body = f"""
<form method="post" action="/sow/save" id="sowForm">
<input type="hidden" name="id" value="{_esc(sow.get('id') or '')}">
<input type="hidden" name="type" value="{type_key}">
<input type="hidden" name="rows_json" id="rowsJson">
<div class="doc-bar">
  <a class="btn btn-ghost btn-sm" href="/sow" title="All documents">←</a>
  <span class="dir-chip dir-samsung">SEA</span>
  <span style="font-size:.82rem;font-weight:700">🧮 Cost Estimation</span>
  <a class="btn btn-ghost btn-sm" href="/sow/people">👥 People</a>
  <span class="spacer"></span>
  {saved_note}
  {f'<a class="btn btn-secondary btn-sm" href="/sow/xlsx?id={sow["id"]}">⬇ xlsx</a>' if sow.get('id') else ''}
  <button type="submit" class="btn btn-primary btn-sm">💾 Save</button>
</div>
<div class="paper" style="max-width:1200px">
  <div class="doc-title">{_slot('title', sow.get('title'), 'Estimation title — e.g. AEM Bridge 2', 'style="width:100%;font-weight:800;font-size:1.15rem"')}</div>
  <div style="font-weight:800;margin:8px 0 18px">COST ESTIMATION</div>
  <div class="meta-line"><b>PROJECT:</b> {_slot('project_name', sow.get('project_name'), 'project', 'style="flex:1"')}</div>
  <div class="meta-line"><b>PERIOD:</b> {_slot('period_label', sow.get('period_label'), 'e.g. From Aug til Dec (Total)')}
    <b style="margin-left:14px"># MONTHS:</b> <input class="slot" type="number" step="0.1" min="0" name="months" id="fMonths" value="{_esc(sow.get('months') or 5)}" style="width:80px"></div>

  <h2>Resources <span style="font-size:.7rem;color:var(--text-muted);font-weight:500">Monthly = rate × 168h · Monthly Cost = Monthly × allocation · Total = Monthly Cost × months</span></h2>
  <div class="table-wrap"><table>
    <thead><tr><th>Resource</th><th>Function</th><th>Location</th><th>Rate/h</th><th>Monthly</th><th>Allocation</th><th>Monthly Cost</th><th>Total</th><th style="border:none"></th></tr></thead>
    <tbody id="estBody"></tbody>
    <tfoot><tr><td colspan="6" style="text-align:right"><b>Totals</b></td><td class="num"><b id="totMonthly">$0</b></td><td class="num"><b id="totTotal">$0</b></td><td style="border:none"></td></tr></tfoot>
  </table></div>
  <button type="button" class="btn btn-ghost btn-sm add-row-btn" onclick="addRow()">+ Add resource</button>
  <div id="marginLine" style="display:none;margin-top:10px;font-size:.76rem;color:var(--text-muted)"></div>
  <div class="ro-note">Pick people from the roster — rate/function auto-fill (editable per row). Figures recompute live.</div>
</div>
</form>
""" + _EST_JS.replace("__ROWS__", json.dumps(sow.get("rows", [])).replace("</", "<\\/")) \
             .replace("__CFG__", json.dumps(cfg).replace("</", "<\\/")) + _EX_TOGGLE_JS
    return _shell(user, "Cost Estimation", body, wide=True)


def _render_vendors(user, saved=False):
    data = _load(user)
    counts = {}
    for s in data["sows"]:
        vid = s.get("vendor_id")
        if vid:
            counts[vid] = counts.get(vid, 0) + 1
    today = date.today()
    rows = []
    for v in data["vendors"]:
        n = counts.get(v["id"], 0)
        del_btn = (f'<button type="button" class="btn btn-danger btn-sm" onclick="delVendor(\'{v["id"]}\')">🗑</button>'
                   if n == 0 else
                   f'<span style="font-size:.72rem;color:var(--text-muted)" title="Referenced by {n} document(s) — delete those first">{n} doc(s)</span>')
        # mapping line: this vendor's contracts (name-matched) + its people
        # (affiliation match or linked to those contracts) — 강프로 2026-07-24
        vn_norm = _norm_tokens(v.get("name"))
        v_low = (v.get("name") or "").strip().lower()
        v_ctrs = [c for c in data.get("contracts", [])
                  if c.get("side") == "vendor"
                  and ((vn_norm and _norm_tokens(c.get("vendor")) == vn_norm)
                       or (c.get("vendor") or "").strip().lower() == v_low)]
        cids = {c["id"] for c in v_ctrs}
        v_ppl = [p for p in data.get("people", [])
                 if (vn_norm and _norm_tokens(p.get("affiliation")) == vn_norm)
                 or (set(p.get("linked_contracts") or []) & cids)]
        map_bits = []
        for c in v_ctrs:
            s0 = _parse_any_date(c.get("period_start"))
            e0 = _effective_end(data, c)
            act = ' <span style="color:var(--success)">●</span>' if (s0 and e0 and s0 <= today <= e0) else ''
            map_bits.append(
                f'<a href="/sow/contracts?newc={c["id"]}" class="ctr-chip" '
                f'style="text-decoration:none">📎 {_esc(c.get("project_name") or c.get("filename") or "contract")}'
                f'{(" · " + _esc(c.get("amount"))) if c.get("amount") else ""}{act}</a>')
        for p in sorted(v_ppl, key=lambda x: (x.get("name") or "").lower()):
            map_bits.append(
                f'<a href="/sow/person?id={p["id"]}" class="ctr-chip" '
                f'style="text-decoration:none">👤 {_esc(p.get("name"))}'
                f'{(" · " + _esc(p.get("role_title"))) if p.get("role_title") else ""}</a>')
        map_line = ""
        if map_bits:
            map_line = ('<div style="display:flex;gap:6px;flex-wrap:wrap;align-items:center;'
                        'padding:0 18px 12px;margin-top:-6px">' + "".join(map_bits) + '</div>')
        rows.append(f"""
<div class="sow-card" style="padding:0">
<form method="post" action="/sow/vendor/save" class="vnd-form">
  <input type="hidden" name="id" value="{v['id']}">
  <input class="slot" name="name" value="{_esc(v.get('name'))}" placeholder="Vendor name">
  <input class="slot" name="entity_line" value="{_esc(v.get('entity_line'))}" placeholder="Entity line (name + address, used in the SOW/MSA preamble)">
  <input class="slot" type="date" name="msa_date" value="{_esc(v.get('msa_date'))}" title="MSA date">
  <button type="submit" class="btn btn-secondary btn-sm">💾 Save</button>
  {del_btn}
</form>
{map_line}
</div>""")
    saved_banner = ('<div style="color:var(--success);font-size:.85rem;margin-bottom:12px">✓ Saved</div>'
                    if saved else "")
    body = f"""
<div style="display:flex;align-items:center;gap:12px;margin:8px 0 4px">
  <h1 style="margin:0">🏢 Vendors</h1>
</div>
<p style="color:var(--text-muted);font-size:.86rem">One row per contractor — name, the entity line that lands in document preambles, and the MSA date. Referenced by Agency-side SOWs, MSAs and NDAs.</p>
{saved_banner}
<form method="post" action="/sow/vendor/save" class="sow-card vnd-form new">
  <input class="slot" name="name" placeholder="+ New vendor name — e.g. Invictus Data, Inc." required>
  <input class="slot" name="entity_line" placeholder="Entity line — Invictus Data Inc, with its principal place of business located at ...">
  <input class="slot" type="date" name="msa_date" title="MSA date">
  <button type="submit" class="btn btn-primary btn-sm">+ Add</button>
</form>
<div style="display:flex;flex-direction:column;gap:0">{''.join(rows) or '<div class="sow-meta" style="padding:30px;text-align:center">No vendors yet.</div>'}</div>
<script>
function delVendor(id){{
  if(!confirm('Delete this vendor?')) return;
  fetch('/sow/vendor/delete', {{method:'POST', headers:{{'Content-Type':'application/x-www-form-urlencoded'}}, body:'id='+encodeURIComponent(id)}})
    .then(function(){{ location.reload(); }});
}}
</script>"""
    return _shell(user, "Vendors", body, tab="vendors")


def _render_types(user, dir_key):
    is_agy = dir_key == "agy"
    name = "with Agency (Vendor)" if is_agy else "with Samsung (SEA)"

    def cards(kind):
        out = []
        for key, t in TYPES.items():
            if key.startswith(dir_key) and t["kind"] == kind:
                out.append(
                    f'<a class="type-card" href="/sow/new?type={key}">'
                    f'<span class="type-icon">{t["icon"]}</span>'
                    f'<span><div class="type-name">{_esc(t["label"])}</div>'
                    f'<div class="type-desc">{_esc(t["desc"])}</div></span></a>'
                )
        return "".join(out)

    sections = f"""
<h2 style="font-size:.92rem;font-weight:800;margin:22px 0 0">Statement of Work</h2>
<div class="type-grid">{cards('sow')}</div>"""
    if not is_agy:
        sections += f"""
<h2 style="font-size:.92rem;font-weight:800;margin:10px 0 0">Estimation</h2>
<div class="type-grid">{cards('est')}</div>"""
    if is_agy:
        sections += f"""
<h2 style="font-size:.92rem;font-weight:800;margin:10px 0 0">Agreements</h2>
<div class="type-grid">{cards('msa')}{cards('nda')}</div>"""
    else:
        sections += ('<p style="color:var(--text-muted);font-size:.78rem">Master agreement with Samsung '
                     '(Advertising Services Agreement, Sep 16 2022) already exists — SOWs reference it; '
                     'MSA/NDA drafting lives under <a href="/sow/types?dir=agy" style="color:var(--accent)">with Agency</a>.</p>')
    body = f"""
<div style="display:flex;align-items:center;gap:12px;margin:8px 0 4px">
  <a class="btn btn-ghost btn-sm" href="/sow">←</a>
  <h1 style="margin:0">{name}</h1>
</div>
<p style="color:var(--text-muted);font-size:.86rem">Pick the document type — the template opens as a live preview; fill in the highlighted fields.</p>
{sections}"""
    return _shell(user, "Document Type", body)


# ── document editor ──────────────────────────────────────────────────────────

_EDITOR_JS = """<script>
var RES = __RES__;
var OV = __OV__;   /* per-month amount overrides: {'Mar-26': 1260, ...} */
var CFG = __CFG__;
function $id(x){ return document.getElementById(x); }
function mode(){ return $id('fMode').value; }
function money(x){
  if(isNaN(x)) return '$0';
  var r = Math.round(x*100)/100;
  return '$' + (Math.abs(r-Math.round(r))<0.005 ? Math.round(r).toLocaleString('en-US')
              : r.toLocaleString('en-US', {minimumFractionDigits:2}));
}
function monthSpans(s, e){
  if(!s || !e) return [];
  var sd = new Date(s+'T00:00:00'), ed = new Date(e+'T00:00:00');
  if(ed < sd) return [];
  var out = [], y = sd.getFullYear(), m = sd.getMonth();
  while(y < ed.getFullYear() || (y === ed.getFullYear() && m <= ed.getMonth())){
    var last = new Date(y, m+1, 0).getDate();
    var d1 = (y===sd.getFullYear() && m===sd.getMonth()) ? sd.getDate() : 1;
    var d2 = (y===ed.getFullYear() && m===ed.getMonth()) ? ed.getDate() : last;
    var frac = (d1===1 && d2===last) ? 1 : Math.min(1, Math.round((d2-d1+1)/30*100)/100);
    out.push([y, m, frac]);
    m++; if(m===12){ m=0; y++; }
  }
  return out;
}
var MON = ['Jan','Feb','Mar','Apr','May','Jun','Jul','Aug','Sep','Oct','Nov','Dec'];
function monthlyAmt(){
  var t = 0;
  RES.forEach(function(r){
    if(mode() === 'hourly') t += (parseFloat(r.hourly)||0)*(parseFloat(r.hrs)||0)*(parseFloat(r.qty)||1);
    else t += parseFloat(r.rate)||0;
  });
  return t;
}
var _months = 0;
function schedRows(){
  var spans = monthSpans($id('fStart').value, $id('fEnd').value);
  var monthly = monthlyAmt(), rule = $id('fRule').value;
  _months = 0;
  return spans.map(function(sp){
    var y = sp[0], m = sp[1], frac = sp[2];
    _months += frac;
    var label = MON[m] + '-' + String(y).slice(2);
    var auto = Math.round(monthly*frac*100)/100;
    var amt = (OV[label] != null && OV[label] !== '') ? parseFloat(OV[label]) || 0 : auto;
    var inv;
    if(rule === 'month_end') inv = new Date(y, m+1, 0);
    else inv = (m===11) ? new Date(y+1, 0, 1) : new Date(y, m+1, 1);
    return {label: label, amt: amt, auto: auto, edited: OV[label] != null && OV[label] !== '',
            inv: inv.getDate() + '-' + MON[inv.getMonth()] + '-' + String(inv.getFullYear()).slice(2)};
  });
}
function feeOf(rows){ return rows.reduce(function(a, r){ return a + r.amt; }, 0); }
function paintTotals(rows){
  var fee = feeOf(rows);
  document.querySelectorAll('.feeOut').forEach(function(el){ el.textContent = money(fee); });
  $id('monthsOut') && ($id('monthsOut').textContent = Math.round(_months*10)/10);
  var t = $id('schedTotal'); if(t) t.textContent = money(fee);
  return fee;
}
function renderSched(){
  var rows = schedRows();
  var tb = $id('schedBody'); tb.innerHTML = '';
  rows.forEach(function(r){
    var tr = document.createElement('tr');
    tr.innerHTML = '<td>' + r.label + '</td>'
      + '<td class="num"><input class="slot sched-amt" data-label="' + r.label + '" value="' + r.amt + '"'
      + ' title="' + (r.edited ? 'Edited — auto value ' + money(r.auto) : 'Auto from period × rates; type to override') + '"'
      + (r.edited ? ' style="border-bottom-color:var(--warn)"' : '') + '></td>'
      + '<td>' + r.inv + '</td>';
    tb.appendChild(tr);
  });
  var tr = document.createElement('tr');
  tr.innerHTML = '<td><b>Total</b></td><td class="num"><b id="schedTotal"></b></td>'
    + '<td><button type="button" class="btn btn-ghost btn-sm" onclick="OV={};renderSched()" title="Drop manual edits, back to auto">↺ auto</button></td>';
  tb.appendChild(tr);
  paintTotals(rows);
}
function updateResComputed(){
  if(mode() !== 'hourly') return;
  document.querySelectorAll('#resBody tr').forEach(function(tr, i){
    var r = RES[i] || {};
    var cost = (parseFloat(r.hourly)||0)*(parseFloat(r.hrs)||0)*(parseFloat(r.qty)||1)*_months;
    var tds = tr.querySelectorAll('td.num-ro');
    if(tds[0]) tds[0].textContent = Math.round(_months*10)/10;
    if(tds[1]) tds[1].textContent = money(cost);
  });
}
var HEADS = {
  hourly:  ['Profile','Location','Qty','# Months','Hourly USD','Hrs/Month','Cost'],
  monthly: ['No.','Name','Role','Level','Region','Rate/Month USD']
};
function renderRes(){
  var thead = document.querySelector('#resTableHead');
  thead.innerHTML = '<tr>' + HEADS[mode()].map(function(h){ return '<th>' + h + '</th>'; }).join('') + '<th style="border:none"></th></tr>';
  var tb = $id('resBody'); tb.innerHTML = '';
  RES.forEach(function(r, i){
    var tr = document.createElement('tr'), html = '';
    if(mode() === 'hourly'){
      html = cell(r,'profile',i) + cell(r,'location',i) + cell(r,'qty',i) +
             '<td class="num num-ro"></td>' + cell(r,'hourly',i) + cell(r,'hrs',i) +
             '<td class="num num-ro"></td>';
    } else {
      html = '<td class="num">' + (i+1) + '</td>' + cell(r,'name',i) + cell(r,'role',i) +
             cell(r,'level',i) + cell(r,'region',i) + cell(r,'rate',i);
    }
    tr.innerHTML = html + '<td style="border:none"><button type="button" class="row-del" onclick="rmRow('+i+')">✕</button></td>';
    tb.appendChild(tr);
  });
  updateResComputed();
}
function cell(r, k, i){
  var v = r[k] != null ? String(r[k]).replace(/"/g,'&quot;') : '';
  var dl = (k === 'profile' || k === 'name') ? ' list="peopleList"' : '';
  return '<td><input class="slot" data-k="'+k+'" data-i="'+i+'" value="'+v+'"'+dl+'></td>';
}
function rosterFill(i, name){
  var p = (CFG.people || []).find(function(x){ return x.name === name; });
  if(!p) return;
  var r = RES[i];
  if(mode() === 'hourly'){
    if(!r.hourly && p.rate) r.hourly = p.rate;
    if(!r.location && p.location) r.location = p.location;
  } else {
    if(!r.role && p.function) r.role = p.function;
    if(!r.region && p.location) r.region = p.location;
    if(!r.rate && p.rate) r.rate = String(Math.round(parseFloat(p.rate)*168));
  }
  var tr = document.querySelectorAll('#resBody tr')[i];
  if(tr) tr.querySelectorAll('input[data-k]').forEach(function(inp){
    if(inp.dataset.k !== 'profile' && inp.dataset.k !== 'name')
      inp.value = r[inp.dataset.k] != null ? r[inp.dataset.k] : '';
  });
}
function addRow(){ RES.push({}); renderRes(); renderSched(); updateResComputed(); }
function rmRow(i){ RES.splice(i,1); renderRes(); renderSched(); updateResComputed(); }
function modeSync(){ renderRes(); renderSched(); updateResComputed(); }
/* Typing must never rebuild the input being typed in — update data + computed
   cells in place; full re-renders only on structural events. */
document.addEventListener('input', function(e){
  if(e.target.classList && e.target.classList.contains('sched-amt')){
    OV[e.target.dataset.label] = e.target.value;
    paintTotals(schedRows());
    return;
  }
  var k = e.target.dataset && e.target.dataset.k;
  if(k){
    var i = parseInt(e.target.dataset.i);
    RES[i][k] = e.target.value;
    if(k === 'profile' || k === 'name') rosterFill(i, e.target.value);
    renderSched(); updateResComputed();
  }
});
document.addEventListener('change', function(e){
  if(e.target.id === 'fStart' || e.target.id === 'fEnd' || e.target.id === 'fRule' || e.target.id === 'fDate'){
    renderSched(); updateResComputed(); dateSyncPreamble();
  }
});
function dateSyncPreamble(){
  var dEl = $id('preamDate');
  if(dEl){
    var dv = $id('fDate').value;
    dEl.textContent = dv ? new Date(dv+'T00:00:00').toLocaleDateString('en-US',{year:'numeric',month:'long',day:'numeric'}) : '(date)';
  }
}
var vSel = $id('fVendor');
function vendorSync(){
  if(!vSel) return;
  var v = CFG.vendors[vSel.value] || {};
  var en = $id('preamVendor'), md = $id('preamMsa'), nm = document.querySelectorAll('.vendorName');
  if(en) en.textContent = v.entity_line || v.name || '(vendor — register below)';
  if(md) md.textContent = v.msa_date_long || '(MSA date)';
  nm.forEach(function(el){ el.textContent = v.name || 'Contractor'; });
}
if(vSel) vSel.addEventListener('change', vendorSync);
document.getElementById('sowForm').addEventListener('submit', function(){
  $id('resJson').value = JSON.stringify(RES.filter(function(r){
    return Object.keys(r).some(function(k){ return String(r[k]||'').trim(); });
  }));
  var clean = {};
  Object.keys(OV).forEach(function(k){ if(String(OV[k]).trim() !== '') clean[k] = parseFloat(OV[k]) || 0; });
  $id('ovJson').value = JSON.stringify(clean);
});
if(!RES.length) RES.push({});
vendorSync(); dateSyncPreamble(); renderRes(); renderSched(); updateResComputed();
</script>"""


def _slot(name, value, ph="", extra="", tag="input"):
    return (f'<input class="slot" name="{name}" value="{_esc(value)}" '
            f'placeholder="{_esc(ph)}" {extra}>')


def _render_doc_editor(user, sow, type_key, saved=False):
    data = _load(user)
    t = TYPES[type_key]
    is_agency = t["dir"] == "agency"
    mode = sow.get("res_mode") or t["mode"]
    vendors = {v["id"]: v for v in data["vendors"]}
    for v in vendors.values():
        v["msa_date_long"] = _fmt_long(v.get("msa_date"))
    cur_vendor = vendors.get(sow.get("vendor_id")) or {}
    vend_opts = "".join(
        f'<option value="{vid}"{" selected" if vid == sow.get("vendor_id") else ""}>{_esc(v["name"])}</option>'
        for vid, v in vendors.items()
    )
    dir_chip = ('<span class="dir-chip dir-agency">Agency</span>' if is_agency
                else '<span class="dir-chip dir-samsung">SEA</span>')
    saved_note = ('<span style="color:var(--success);font-size:.8rem;font-weight:700">✓ Saved</span>'
                  if saved else "")
    vendor_bar = ""
    if is_agency:
        vendor_bar = f"""
  <span style="display:flex;align-items:center;gap:6px;font-size:.78rem;color:var(--text-muted)">Vendor
    <select class="slot" name="vendor_id" id="fVendor" style="min-width:150px">
      <option value="">— pick —</option>{vend_opts}
    </select>
  </span>
  <details style="position:relative"><summary class="btn btn-ghost btn-sm" style="list-style:none">+ New vendor</summary>
    <div style="position:absolute;top:calc(100% + 8px);left:0;z-index:80;background:var(--surface-3);border:1px solid var(--border-bright);border-radius:10px;padding:14px;display:flex;flex-direction:column;gap:8px;min-width:300px;box-shadow:var(--shadow-lg)">
      <input class="slot" name="v_name" placeholder="Vendor name — e.g. Invictus Data, Inc.">
      <input class="slot" name="v_entity" placeholder="Entity line (name + address for preamble)">
      <input class="slot" type="date" name="v_msa" title="MSA date">
      <span style="font-size:.7rem;color:var(--text-muted)">Saved together on 💾 Save</span>
    </div>
  </details>"""

    # ── document body ──
    if is_agency:
        preamble = f"""
<p class="legal">This Statement of Work ("Statement of Work" or "SOW") is made effective as of <b id="preamDate">{_esc(_fmt_long(sow.get('date')) or '(date)')}</b> (the "Statement of Work Effective Date") by and between Cheil USA Inc., a Delaware corporation with its principal of business located at 837 Washington Street, 4th Floor, New York, NY 10014 on behalf of itself and its affiliates and subsidiaries ("Cheil") and <b id="preamVendor">{_esc(cur_vendor.get('entity_line') or cur_vendor.get('name') or '(vendor — register above)')}</b> ("Contractor").  Contractor and Cheil may each be referred to herein as a "Party", and, together as the "Parties".</p>
<p class="legal">This SOW is governed by, incorporated into, and made part of, that certain Master Services Agreement (the "Agreement"), dated as of <b id="preamMsa">{_esc(cur_vendor.get('msa_date_long') or '(MSA date)')}</b>, by and between Cheil and Contractor. This SOW defines the Services that Contractor shall provide to Cheil in accordance with the terms of the Agreement and this SOW. […] To the extent there is a conflict between the terms of this SOW and the Agreement, the terms of the Agreement shall control, except for terms where the Agreement expressly permits the SOW to control in the event of conflict with the Agreement.</p>"""
    else:
        preamble = f'<p class="legal">{_esc(PREAMBLE_SAMSUNG)}</p>'

    stk_client_label = (f'<span class="vendorName">{_esc(cur_vendor.get("name") or "Contractor")}</span> POC'
                        if is_agency else "Samsung Manager for this Role")
    client_line = CHEIL_ENTITY if is_agency else SAMSUNG_ENTITY
    oop_html = "".join(f'<p class="legal">{_esc(par)}</p>'
                       for par in (OOP_AGENCY if is_agency else OOP_SAMSUNG).split("\n\n"))
    sig_left = CHEIL_ENTITY if is_agency else SAMSUNG_ENTITY
    sig_right = (f'<span class="vendorName">{_esc(cur_vendor.get("name") or "Contractor")}</span>'
                 if is_agency else CHEIL_ENTITY)

    people_cfg = [{"name": pp.get("name") or "", "function": pp.get("role_title") or "",
                   "location": pp.get("location") or "", "rate": pp.get("sell_hr") or ""}
                  for pp in data["people"]]
    cfg = {"vendors": vendors, "people": people_cfg}
    body = f"""
<form method="post" action="/sow/save" id="sowForm">
<input type="hidden" name="id" value="{_esc(sow.get('id') or '')}">
<input type="hidden" name="type" value="{type_key}">
<input type="hidden" name="resources_json" id="resJson">
<input type="hidden" name="schedule_overrides" id="ovJson">
<div class="doc-bar">
  <a class="btn btn-ghost btn-sm" href="/sow" title="All SOWs">←</a>
  {dir_chip}<span style="font-size:.82rem;font-weight:700">{t['icon']} {_esc(t['label'])}</span>
  {vendor_bar}
  <span class="spacer"></span>
  {saved_note}
  <button type="button" class="btn btn-secondary btn-sm" id="exToggle" title="Show/hide the executed example">📖 Example</button>
  {f'<a class="btn btn-secondary btn-sm" href="/sow/docx?id={sow["id"]}">⬇ docx</a>' if sow.get('id') else ''}
  <button type="submit" class="btn btn-primary btn-sm">💾 Save</button>
</div>

<div class="ed-wrap" id="edWrap">
<div class="paper">
  {'<img src="/sow/asset/logo" alt="Cheil × Samsung" style="max-width:300px;margin-bottom:22px;background:#fff;padding:10px 14px;border-radius:8px">' if not is_agency else ''}
  <div class="doc-title">{_slot('title', sow.get('title'), 'SOW title — e.g. Data Engineer # 1', 'style="width:100%;font-weight:800;font-size:1.15rem"')}</div>
  <div style="font-weight:800;margin:8px 0 18px">STATEMENT OF WORK</div>

  <div class="meta-line"><b>DATE:</b> <input class="slot" type="date" name="date" id="fDate" value="{_esc(sow.get('date') or date.today().isoformat())}"></div>
  <div class="meta-line"><b>CLIENT:</b> <span>{client_line}</span></div>
  <div class="meta-line"><b>PROJECT NAME:</b> {_slot('project_name', sow.get('project_name'), 'e.g. SEA eCom Data', 'style="flex:1"')}</div>
  <div class="meta-line"><b>PREPARED BY:</b> {_slot('prepared_by', sow.get('prepared_by'), 'name')}</div>
  <div class="meta-line"><b>PREPARED FOR:</b> {_slot('prepared_for', sow.get('prepared_for'), 'name')}</div>
  {'<div style="border-top:2px dashed var(--border-bright);margin:26px -52px;position:relative"><span style="position:absolute;top:-9px;left:50%;transform:translateX(-50%);background:var(--surface);padding:0 10px;font-size:.64rem;color:var(--text-muted);text-transform:uppercase;letter-spacing:.08em">Page 2</span></div>' if not is_agency else ''}

  {preamble}

  <h2>Executive Summary</h2>
  <textarea class="slot" name="exec_summary" rows="3" placeholder="What service does this SOW provide?">{_esc(sow.get('exec_summary'))}</textarea>

  <h2>{'Service Description' if is_agency else 'Deliverables'}</h2>
  <textarea class="slot" name="deliverables" rows="7" placeholder="One item per line — rendered as bullets in the document">{_esc(sow.get('deliverables'))}</textarea>

  <h2>Project Stakeholders</h2>
  <div class="table-wrap"><table>
    <tr><th></th><th>{stk_client_label}</th><th>Cheil Project Management &amp; SOW Owner</th></tr>
    <tr><td><b>Name</b></td><td><input class="slot" name="stk_c_name" value="{_esc(sow.get('stk_c_name'))}"></td><td><input class="slot" name="stk_a_name" value="{_esc(sow.get('stk_a_name'))}"></td></tr>
    <tr><td><b>Email</b></td><td><input class="slot" name="stk_c_email" value="{_esc(sow.get('stk_c_email'))}"></td><td><input class="slot" name="stk_a_email" value="{_esc(sow.get('stk_a_email'))}"></td></tr>
    <tr><td><b>Location</b></td><td><input class="slot" name="stk_c_loc" value="{_esc(sow.get('stk_c_loc'))}"></td><td><input class="slot" name="stk_a_loc" value="{_esc(sow.get('stk_a_loc'))}"></td></tr>
  </table></div>

  <h2>Service Period</h2>
  <div class="meta-line"><b>Start Date :</b> <input class="slot" type="date" name="start" id="fStart" value="{_esc(sow.get('start'))}" required></div>
  <div class="meta-line"><b>End Date :</b> <input class="slot" type="date" name="end" id="fEnd" value="{_esc(sow.get('end'))}" required></div>

  <h2>{'Resource Planning' if is_agency else 'Resource Management'}
    <select class="slot" name="res_mode" id="fMode" onchange="modeSync()" style="font-size:.76rem;margin-left:10px;vertical-align:middle" title="Rate model — switches the resource table columns">
      <option value="hourly"{' selected' if mode == 'hourly' else ''}>Hourly (rate × hrs/month)</option>
      <option value="monthly"{' selected' if mode == 'monthly' else ''}>Monthly rate per member</option>
    </select></h2>
  <p class="legal">In consideration for the provision of the Services and Deliverables under this SOW, {'Cheil shall pay Contractor' if is_agency else 'Samsung shall pay Cheil'} in accordance with the following rates and fees, subject to the applicable terms and conditions of the Agreement:</p>
  <div class="table-wrap"><table>
    <thead id="resTableHead"></thead>
    <tbody id="resBody"></tbody>
  </table></div>
  <button type="button" class="btn btn-ghost btn-sm add-row-btn" onclick="addRow()">+ Add resource</button>

  <h2>Cost and Payment Schedule</h2>
  <p><b>Fee : <span class="feeOut">$0</span></b> <span style="color:var(--text-muted);font-size:.78rem">(<span id="monthsOut">0</span> months · amounts are auto-filled, type in any month to override)</span></p>
  <p class="legal">{PAYMENT_INTRO_AGENCY if is_agency else PAYMENT_INTRO}
    <br>Invoice dates: <select class="slot" name="invoice_rule" id="fRule">
      <option value="next_first"{'' if sow.get('invoice_rule') == 'month_end' else ' selected'}>1st of following month</option>
      <option value="month_end"{' selected' if sow.get('invoice_rule') == 'month_end' else ''}>last day of service month</option>
    </select></p>
  <div class="table-wrap"><table>
    <thead><tr><th>Month</th><th>Amount</th><th>Invoice Date</th></tr></thead>
    <tbody id="schedBody"></tbody>
  </table></div>
  <p class="legal">{_esc(CHANGE_ORDER_NOTE)}</p>

  <h2>Out-of-pocket Expense</h2>
  {oop_html}

  <h2>Signatures</h2>
  <p class="legal">IN WITNESS WHEREOF, the parties have caused this Statement of Work to be duly executed by their authorized representatives as set forth below.</p>
  <div class="table-wrap"><table>
    <tr><th>{sig_left}</th><th>{sig_right}</th></tr>
    <tr><td>Signature: ____________________</td><td>Signature: ____________________</td></tr>
    <tr><td>Name: ____________________</td><td>Name: ____________________</td></tr>
    <tr><td>Title: ____________________</td><td>Title: ____________________</td></tr>
    <tr><td>Date: ____________________</td><td>Date: ____________________</td></tr>
  </table></div>
  <div class="ro-note">Highlighted fields are editable · everything else exports as-is to .docx</div>
</div>
{_render_example(type_key)}
</div>
<datalist id="peopleList">{"".join(f'<option value="{_esc(pp["name"])}">' for pp in people_cfg)}</datalist>
</form>
""" + _EDITOR_JS.replace("__RES__", json.dumps(sow.get("resources", [])).replace("</", "<\\/")) \
                .replace("__OV__", json.dumps(sow.get("schedule_overrides") or {}).replace("</", "<\\/")) \
                .replace("__CFG__", json.dumps(cfg).replace("</", "<\\/")) + _EX_TOGGLE_JS
    return _shell(user, "SOW Editor", body, wide=True)


_MSA_PARS = None


def _msa_paragraphs():
    """Extract the MSA template's full text once: [(kind, text)] where kind is
    'title' | 'h' (numbered section heading) | 'p'. Falls back to [] when
    python-docx is unavailable (preview then shows the summary box)."""
    global _MSA_PARS
    if _MSA_PARS is not None:
        return _MSA_PARS
    try:
        from docx import Document
        doc = Document(os.path.join(_ASSETS, "msa_template.docx"))
        out = []
        for p in doc.paragraphs:
            t = " ".join(p.text.split())
            if not t:
                continue
            bold = any(r.bold for r in p.runs if r.text.strip())
            if t == "MASTER SERVICES AGREEMENT":
                kind = "title"
            elif p.style.name == "List Paragraph" and bold:
                kind = "h"
            else:
                kind = "p"
            out.append((kind, t))
        _MSA_PARS = out
    except Exception:
        _MSA_PARS = []
    return _MSA_PARS


_AGREEMENT_JS = """<script>
var CFG = __CFG__;
function $id(x){ return document.getElementById(x); }
function vendorSync(){
  var v = CFG.vendors[$id('fVendor').value] || {};
  document.querySelectorAll('.vendorName').forEach(function(el){
    el.textContent = v.name || '______________________';
  });
}
function dateSync(){
  var dv = $id('fDate').value, el = $id('agrDate');
  if(el) el.textContent = dv ? new Date(dv+'T00:00:00').toLocaleDateString('en-US',{year:'numeric',month:'long',day:'numeric'}) : '____________';
}
document.addEventListener('change', function(){ vendorSync(); dateSync(); });
vendorSync(); dateSync();
</script>"""


def _render_agreement_editor(user, sow, type_key, saved=False):
    data = _load(user)
    t = TYPES[type_key]
    kind = t["kind"]
    vendors = {v["id"]: v for v in data["vendors"]}
    cur_vendor = vendors.get(sow.get("vendor_id")) or {}
    vend_opts = "".join(
        f'<option value="{vid}"{" selected" if vid == sow.get("vendor_id") else ""}>{_esc(v["name"])}</option>'
        for vid, v in vendors.items()
    )
    saved_note = ('<span style="color:var(--success);font-size:.8rem;font-weight:700">✓ Saved</span>'
                  if saved else "")
    vname = f'<b class="vendorName">{_esc(cur_vendor.get("name") or "______________________")}</b>'
    date_slot = f'<input class="slot" type="date" name="date" id="fDate" value="{_esc(sow.get("date") or date.today().isoformat())}">'

    if kind == "msa":
        doc_title = "MASTER SERVICES AGREEMENT"
        pars = _msa_paragraphs()
        if pars:
            chunks, sec_n = [], 0
            for pk, txt in pars:
                if pk == "title":
                    continue  # rendered separately above the date slot
                html = _esc(txt)
                html = html.replace("XXX XX, 2026",
                                    f'<b id="agrDate">{_esc(_fmt_long(sow.get("date")) or "____________")}</b>')
                html = html.replace("(Your Company Name)", vname)
                if pk == "h":
                    sec_n += 1
                    chunks.append(f'<h2 style="font-size:.95rem">{sec_n}. {html}</h2>')
                else:
                    chunks.append(f'<p class="legal">{html}</p>')
            body_doc = "".join(chunks)
        else:
            body_doc = f"""
  <p class="legal">This Master Services Agreement (this "Agreement"), dated as of <b id="agrDate">{_esc(_fmt_long(sow.get('date')) or '____________')}</b> (the "Effective Date"), is made by and between Cheil USA Inc., a Delaware corporation ("Cheil"), and {vname} ("Contractor").</p>
  <div style="background:var(--surface-2);border:1px dashed var(--border-bright);border-radius:10px;padding:16px 18px;margin:18px 0;font-size:.8rem;color:var(--text-muted);line-height:1.7">
    ⚑ Full-text preview needs python-docx on this host — the export still contains the complete
    legal text verbatim from the executed Cheil MSA template.
  </div>"""
    else:
        clauses = "".join(f'<p class="legal">{_esc(par)}</p>' for par in NDA_BODY)
        body_doc = f"""
  <p class="legal">This CONFIDENTIALITY AND NONDISCLOSURE AGREEMENT (the "Agreement"), is entered into as of <b id="agrDate">{_esc(_fmt_long(sow.get('date')) or '____________')}</b> (the "Effective Date"), by and between {vname} (the "Vendor"), and Cheil USA, Inc. (the "Cheil").</p>
  {clauses}
  <div class="table-wrap"><table>
    <tr><th>Cheil USA, Inc.</th><th><span class="vendorName">{_esc(cur_vendor.get('name') or '[_____________________]')}</span></th></tr>
    <tr><td>By: ____________________</td><td>By: ____________________</td></tr>
    <tr><td>Name:</td><td>Name:</td></tr>
    <tr><td>Title:</td><td>Title:</td></tr>
  </table></div>"""
        doc_title = NDA_TITLE

    people_cfg = [{"name": pp.get("name") or "", "function": pp.get("role_title") or "",
                   "location": pp.get("location") or "", "rate": pp.get("sell_hr") or ""}
                  for pp in data["people"]]
    cfg = {"vendors": vendors, "people": people_cfg}
    body = f"""
<form method="post" action="/sow/save" id="sowForm">
<input type="hidden" name="id" value="{_esc(sow.get('id') or '')}">
<input type="hidden" name="type" value="{type_key}">
<div class="doc-bar">
  <a class="btn btn-ghost btn-sm" href="/sow" title="All documents">←</a>
  <span class="dir-chip dir-agency">Agency</span>
  <span style="font-size:.82rem;font-weight:700">{t['icon']} {_esc(t['label'])}</span>
  <span style="display:flex;align-items:center;gap:6px;font-size:.78rem;color:var(--text-muted)">Vendor
    <select class="slot" name="vendor_id" id="fVendor" style="min-width:150px">
      <option value="">— pick —</option>{vend_opts}
    </select>
  </span>
  <details style="position:relative"><summary class="btn btn-ghost btn-sm" style="list-style:none">+ New vendor</summary>
    <div style="position:absolute;top:calc(100% + 8px);left:0;z-index:80;background:var(--surface-3);border:1px solid var(--border-bright);border-radius:10px;padding:14px;display:flex;flex-direction:column;gap:8px;min-width:300px;box-shadow:var(--shadow-lg)">
      <input class="slot" name="v_name" placeholder="Vendor name — e.g. Invictus Data, Inc.">
      <input class="slot" name="v_entity" placeholder="Entity line (name + address)">
      <input class="slot" type="date" name="v_msa" title="MSA date">
      <span style="font-size:.7rem;color:var(--text-muted)">Saved together on 💾 Save</span>
    </div>
  </details>
  <span class="spacer"></span>
  {saved_note}
  <button type="button" class="btn btn-secondary btn-sm" id="exToggle" title="Show/hide the executed example">📖 Example</button>
  {f'<a class="btn btn-secondary btn-sm" href="/sow/docx?id={sow["id"]}">⬇ docx</a>' if sow.get('id') else ''}
  <button type="submit" class="btn btn-primary btn-sm">💾 Save</button>
</div>
<div class="ed-wrap" id="edWrap">
<div class="paper">
  <div style="text-align:center;font-weight:800;font-size:1.05rem;margin-bottom:20px">{doc_title}</div>
  <div class="meta-line" style="margin-bottom:14px"><b>EFFECTIVE DATE:</b> {date_slot}</div>
  {body_doc}
  <div class="ro-note">Highlighted fields are editable · everything else exports as-is to .docx</div>
</div>
{_render_example(type_key)}
</div>
</form>
""" + _AGREEMENT_JS.replace("__CFG__", json.dumps(cfg).replace("</", "<\\/")) + _EX_TOGGLE_JS
    return _shell(user, t["label"], body, wide=True)


# ── routing ──────────────────────────────────────────────────────────────────

def _f(body, key, default=""):
    v = body.get(key, default)
    if isinstance(v, list):
        v = v[0] if v else default
    return (v or "").strip()


# ══════════════════════════════════════════════════════════════════════════
# Uploaded contracts — SEA↔Cheil (upstream) aligned to Cheil↔Vendor (downstream)
# ══════════════════════════════════════════════════════════════════════════

_CONTRACT_MIME = {
    "pdf": "application/pdf",
    "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "doc": "application/msword",
    "txt": "text/plain",
}
_MONTHS = ("January|February|March|April|May|June|July|August|September|"
           "October|November|December|Jan|Feb|Mar|Apr|Jun|Jul|Aug|Sep|Sept|Oct|Nov|Dec")
_DATE_RE = re.compile(
    r"(?:%s)\.?\s+\d{1,2},?\s+\d{4}" % _MONTHS + r"|\b\d{1,2}/\d{1,2}/\d{2,4}\b",
    re.IGNORECASE)
_AMOUNT_RE = re.compile(r"(?:US)?\$\s?([\d]{1,3}(?:,\d{3})+(?:\.\d{2})?|\d+(?:\.\d{2})?)")


def _contracts_dir(user):
    return os.path.join(DATA_ROOT, user, "sow_contracts")


def _contract_file_path(user, c):
    return os.path.join(_contracts_dir(user), f"{c['id']}.{c.get('ext', 'bin')}")


def _contract_text_path(user, c):
    return os.path.join(_contracts_dir(user), f"{c['id']}.txt")


def _contract_text(user, c):
    """Extracted text lives in a sidecar .txt (older records carried it inline
    in sow.json as raw_text — migrated lazily on save paths)."""
    if c.get("raw_text"):
        return c["raw_text"]
    try:
        with open(_contract_text_path(user, c), encoding="utf-8") as fp:
            return fp.read()
    except OSError:
        return ""


def _store_contract_text(user, c, text):
    """Write the sidecar and keep the record lean (no raw_text in sow.json —
    a few 200k-char bodies made every _load reparse megabytes of JSON)."""
    os.makedirs(_contracts_dir(user), exist_ok=True)
    with open(_contract_text_path(user, c), "w", encoding="utf-8") as fp:
        fp.write(text[:200000])
    c.pop("raw_text", None)


def _migrate_contract_texts(user, data):
    """One-time move of inline raw_text blobs out of sow.json."""
    dirty = False
    for c in data.get("contracts", []):
        if "raw_text" in c:
            _store_contract_text(user, c, c.get("raw_text") or "")
            dirty = True
    if dirty:
        _save(user, data)


_ALLOWED_EXT = {"pdf", "docx", "doc", "txt", "eml", "msg"}


def _sheet_ext(fn):
    """Extension used only to pick a table parser. Spreadsheets are read and
    thrown away, never stored, so they are outside the _ALLOWED_EXT whitelist
    that guards on-disk contract files — which would flatten .xlsx to 'bin'
    and silently hand a workbook to the CSV reader."""
    ext = fn.rsplit(".", 1)[-1].lower() if "." in fn else ""
    return re.sub(r"[^a-z0-9]", "", ext)[:5]


def _safe_filename(fn):
    """Display/storage-safe filename: no path, no control chars/quotes."""
    fn = (fn or "").replace("\r", " ").replace("\n", " ").replace('"', " ").replace("\\", "/")
    fn = os.path.basename(fn).strip()
    return fn[:120] or "contract"


def _safe_ext(fn):
    """Whitelisted, alnum-only extension — keeps it out of the filesystem path
    as anything but a plain suffix (no '/', no '..')."""
    ext = fn.rsplit(".", 1)[-1].lower() if "." in fn else ""
    ext = re.sub(r"[^a-z0-9]", "", ext)[:5]
    return ext if ext in _ALLOWED_EXT else "bin"


def _read_multipart(raw_handler):
    """Parse a multipart body once; return (fields, files) where fields is
    {name: str} for the plain parts and files is [(filename, bytes, mime)]."""
    try:
        ct = raw_handler.headers.get("Content-Type", "")
        m = re.search(r"boundary=([^\s;]+)", ct)
        if not m:
            return {}, []
        boundary = ("--" + m.group(1).strip('"')).encode()
        length = int(raw_handler.headers.get("Content-Length", 0))
        data = raw_handler.rfile.read(length)
    except Exception:
        return {}, []
    fields, files = {}, []
    for part in data.split(boundary):
        hdr_end = part.find(b"\r\n\r\n")
        if hdr_end == -1:
            continue
        head = part[:hdr_end]
        content = part[hdr_end + 4:]
        if content.endswith(b"\r\n"):
            content = content[:-2]
        fn_m = re.search(rb'filename="([^"]*)"', head)
        if fn_m and fn_m.group(1):
            if content:
                files.append((fn_m.group(1).decode("utf-8", errors="replace"),
                              content, "application/octet-stream"))
            continue
        nm = re.search(rb'name="([^"]*)"', head)
        if nm:
            fields[nm.group(1).decode("utf-8", errors="replace")] = \
                content.decode("utf-8", errors="replace")
    return fields, files


def _read_uploaded_files(raw_handler):
    """Parse a multipart body; return list of (filename, bytes, mime)."""
    return _read_multipart(raw_handler)[1]


def _docx_text(content):
    """Pull ALL current text from a .docx by reading the package XML directly.

    python-docx's paragraph/table model silently drops text in text boxes,
    nested tables and headers/footers — real SOWs keep their fee tables there,
    so it under-reads badly (one 485k-char contract came back as 15k). Reading
    every <w:t> run from document.xml + headers/footers captures all of it, in
    order, with paragraph breaks. <w:delText> (tracked-change deletions) uses a
    different tag, so stale/deleted figures are naturally excluded."""
    import html
    import zipfile
    with zipfile.ZipFile(io.BytesIO(content)) as z:
        parts = [n for n in z.namelist()
                 if re.match(r"word/(document|header\d*|footer\d*)\.xml$", n)]
        parts.sort(key=lambda n: (not n.startswith("word/document"), n))
        out = []
        for n in parts:
            xml = z.read(n).decode("utf-8", "replace")
            for t, pend, tab in re.findall(
                    r"<w:t[^>]*>(.*?)</w:t>|(</w:p>)|(</w:tc>)", xml, re.DOTALL):
                if t:
                    out.append(html.unescape(t))
                elif tab:
                    out.append(" | ")
                elif pend:
                    out.append("\n")
    return "".join(out)


def _html_to_text(html_src):
    """Readable text out of an HTML mail body — Outlook sends HTML far more
    often than plain text, and the fee numbers live in its tables."""
    import html as _html
    s = re.sub(r"(?is)<(script|style)[^>]*>.*?</\1>", " ", html_src or "")
    s = re.sub(r"(?i)<(br|/tr|/p|/div|/h[1-6])[^>]*>", "\n", s)
    s = re.sub(r"(?i)</t[dh]>", " | ", s)
    s = re.sub(r"(?s)<[^>]+>", "", s)
    s = _html.unescape(s).replace(" ", " ")
    s = re.sub(r"[ \t]+", " ", s)
    return re.sub(r"\n\s*\n\s*\n+", "\n\n", s).strip()


def _eml_parts(content):
    """(meta, body) from a saved RFC-822 message (.eml)."""
    import email
    from email import policy
    msg = email.message_from_bytes(content, policy=policy.default)
    meta = {k: str(msg.get(h) or "") for k, h in
            (("from", "From"), ("to", "To"), ("sent", "Date"), ("subject", "Subject"))}
    body = ""
    try:
        part = msg.get_body(preferencelist=("plain", "html"))
        if part is not None:
            body = part.get_content()
            if part.get_content_type() == "text/html":
                body = _html_to_text(body)
    except Exception:
        body = ""
    if not body.strip():
        body = _html_to_text(content.decode("utf-8", errors="replace"))
    return meta, body


def _msg_parts(content):
    """(meta, body) from an Outlook .msg. Needs extract-msg; without it the
    caller still gets a usable message telling the user to paste the body."""
    try:
        import extract_msg
    except ImportError:
        return {}, ("[.msg needs the extract-msg package on this host — "
                    "paste the email body instead]")
    try:
        m = extract_msg.openMsg(io.BytesIO(content))
    except Exception:
        try:
            with tempfile.NamedTemporaryFile(suffix=".msg", delete=False) as tf:
                tf.write(content)
                tmp = tf.name
            m = extract_msg.openMsg(tmp)
        except Exception as e:
            return {}, f"[could not read .msg: {e}]"
    try:
        meta = {"from": str(m.sender or ""), "to": str(m.to or ""),
                "sent": str(m.date or ""), "subject": str(m.subject or "")}
        body = m.body or ""
        if not (body or "").strip():
            body = _html_to_text((m.htmlBody or b"").decode("utf-8", "replace")
                                 if isinstance(m.htmlBody, bytes) else (m.htmlBody or ""))
        return meta, body
    finally:
        try:
            m.close()
        except Exception:
            pass


def _email_parts(content, ext):
    """(meta, body) for any email-ish upload. Non-mail formats (a printed PDF
    or a pasted Word note) carry no headers — the body is all we get."""
    if ext == "eml":
        return _eml_parts(content)
    if ext == "msg":
        return _msg_parts(content)
    return {}, _extract_text(content, ext)


def _sniff_mail_headers(text):
    """From/Sent/Subject out of a PASTED mail — Outlook copies them as plain
    lines, and they are the provenance of the change, so keep them."""
    meta = {}
    for line in (text or "").splitlines()[:20]:
        m = re.match(r"\s*(from|to|sent|date|subject)\s*:\s*(.+?)\s*$", line, re.I)
        if not m:
            continue
        k = m.group(1).lower()
        meta.setdefault("sent" if k == "date" else k, m.group(2))
    return meta


def _email_as_text(meta, body):
    """Header block + body — what both the AI extractor and the on-screen
    preview read, so what the user reviews is what the model saw."""
    head = "\n".join(f"{k.capitalize()}: {v}" for k, v in
                     (("from", meta.get("from")), ("to", meta.get("to")),
                      ("sent", meta.get("sent")), ("subject", meta.get("subject")))
                     if v)
    return (head + "\n\n" + (body or "")).strip()


def _extract_text(content, ext):
    """Best-effort plain text from an uploaded contract (docx/pdf/txt/eml/msg)."""
    if ext in ("eml", "msg"):
        return _email_as_text(*_email_parts(content, ext))
    try:
        if ext == "docx":
            text = _docx_text(content)
            if len(text.strip()) < 40:  # empty/odd package → fall back
                from docx import Document
                doc = Document(io.BytesIO(content))
                text = "\n".join(p.text for p in doc.paragraphs)
            return text
        if ext == "pdf":
            import fitz
            with fitz.open(stream=content, filetype="pdf") as doc:
                return "\n".join(page.get_text() for page in doc)
        return content.decode("utf-8", errors="replace")
    except Exception as e:
        return f"[could not read {ext}: {e}]"


def _extract_fields(text):
    """Heuristic pull of parties / amount / period / project from contract text.
    Deliberately forgiving — the user confirms & corrects on the popup."""
    low = text.lower()
    flat = re.sub(r"\s+", " ", text)

    # parties — try "by and between X and Y", else entity-like names
    client = agency = vendor = ""
    m = re.search(r"by and between\s+(.+?)\s+(?:\(|,)", flat, re.IGNORECASE)
    n = re.search(r"\band\s+([A-Z][\w&.,'\- ]+?(?:Inc|LLC|L\.L\.C|Corp|Corporation|Ltd|Company)\.?)",
                  flat)
    p1 = (m.group(1).strip() if m else "")
    p2 = (n.group(1).strip() if n else "")
    has_samsung = "samsung" in low
    has_cheil = "cheil" in low
    if has_samsung:
        client = SAMSUNG_ENTITY
        agency = CHEIL_ENTITY
    elif has_cheil:
        agency = CHEIL_ENTITY
    # vendor = whichever named party is neither Cheil nor Samsung
    for cand in (p1, p2):
        cl = cand.lower()
        if cand and "cheil" not in cl and "samsung" not in cl and len(cand) < 90:
            vendor = cand
            break

    # amount — take the largest $ figure (usually the contract total)
    amounts = []
    for a in _AMOUNT_RE.findall(text):
        try:
            amounts.append(float(a.replace(",", "")))
        except ValueError:
            pass
    amount = _money(max(amounts)) if amounts else ""

    # dates — effective/start first, an end/expiry date second
    dates = _DATE_RE.findall(text)
    period_start = dates[0] if dates else ""
    period_end = ""
    em = re.search(r"(?:end date|expir\w*|through|terminat\w*)[^\n]{0,40}?(" +
                   _DATE_RE.pattern + ")", text, re.IGNORECASE)
    if em:
        period_end = em.group(1)
    elif len(dates) > 1:
        period_end = dates[-1]

    # project name — a "Project Name:" label, else the doc title-ish first line
    project = ""
    pm = re.search(r"project\s*name\s*[:\-]\s*(.+)", text, re.IGNORECASE)
    if pm:
        project = pm.group(1).strip().split("\n")[0][:120]
    else:
        pm = re.search(r"project\s*[:\-]\s*(.+)", text, re.IGNORECASE)
        if pm:
            project = pm.group(1).strip().split("\n")[0][:120]

    side = "sea" if has_samsung else ("vendor" if vendor else "sea")
    return {"client": client, "agency": agency, "vendor": vendor,
            "amount": amount, "period_start": period_start,
            "period_end": period_end, "project_name": project, "side": side,
            "people": []}


_EXTRACT_MODEL = os.environ.get("SOW_EXTRACT_MODEL", "claude-sonnet-5")
_EXTRACT_SYSTEM = (
    "You extract structured fields from a contract or Statement of Work. "
    "Return ONLY a JSON object (no prose, no code fences) with exactly these keys:\n"
    '{"side":"sea"|"vendor","client":str,"agency":str,"vendor":str,'
    '"amount":str,"period_start":str,"period_end":str,"project_name":str,'
    '"people":[{"name":str,"role":str,"location":str,"rate":str,'
    '"rate_basis":"hour"|"month"|""}]}\n'
    "Definitions:\n"
    "- side='sea' when the contract is between Samsung (Samsung Electronics "
    "America / SEA) and Cheil, where Samsung pays Cheil (Cheil is the agency).\n"
    "- side='vendor' when it is between Cheil (as the paying client) and a "
    "downstream vendor/contractor that Cheil pays.\n"
    "- client = full legal name of the paying party. agency = 'Cheil USA, Inc.' "
    "when Cheil is a party. vendor = the contractor entity that is neither Cheil "
    "nor Samsung, else ''.\n"
    "- amount = the total contract value/fee including its currency symbol, e.g. "
    "'$240,000'. Prefer the grand total.\n"
    "- period_start / period_end = the service period start and end dates, copied "
    "verbatim as written in the document (e.g. 'June 1, 2024').\n"
    "- project_name = the project or SOW name/title.\n"
    "- people = individuals named as project resources/personnel (resource "
    "tables, staffing sections, key personnel lists). Read the WHOLE resource "
    "row, not just the name:\n"
    "    role      = role / title / profile / function as written.\n"
    "    location  = region / country / office for that person, '' if absent.\n"
    "    rate      = that person's rate with its currency symbol, e.g. '$25' "
    "or '$12,000'. Take the per-person rate, never the contract total.\n"
    "    rate_basis= 'hour' when the rate is hourly, 'month' when it is a "
    "monthly rate/retainer, '' when the document does not say.\n"
    "  EXCLUDE signatories, witnesses and legal representatives who only sign "
    "the document. Empty list if none.\n"
    "Use an empty string '' for anything not present. Do not invent values.")


_EMAIL_SYSTEM = (
    "You read an email in which a client or vendor communicates a change to an "
    "existing contract or Statement of Work — a fee revision, an extension, a "
    "scope cut, an early termination. Contract changes are often agreed over "
    "email before (or instead of) a signed amendment.\n"
    "Return ONLY a JSON object (no prose, no code fences) with exactly these keys:\n"
    '{"is_change":true|false,"amount":str,"period_start":str,"period_end":str,'
    '"project_name":str,"change_summary":str}\n'
    "Definitions:\n"
    "- is_change = false when the email agrees no contractual change (a status "
    "update, a meeting request, an invoice reminder). Everything else stays ''.\n"
    "- amount = the value this change is worth, with its currency symbol, e.g. "
    "'$120,000'. If the email states a NEW TOTAL for the contract, return that "
    "total. '' when the money does not change.\n"
    "- period_start = the date the change takes effect (the effective date). If "
    "the email states no date, use the date the email was sent.\n"
    "- period_end = the end of the service period AFTER this change, '' when "
    "the end date does not move.\n"
    "- project_name = the project/SOW the email is about, '' if not stated.\n"
    "- change_summary = one or two sentences, in English, saying exactly what "
    "changed and who agreed to it. Quote the figures.\n"
    "Copy dates verbatim as written in the email. Use '' for anything not "
    "present. Do not invent values.")


def _claude_json(system, text, max_tokens=1200):
    """POST text to Claude with a JSON-only system prompt; return the parsed
    object, or None when the key is missing / the call or the JSON fails."""
    key = None
    for p in (os.path.join(os.path.dirname(__file__), os.pardir, ".env"),
              os.path.expanduser("~/.claude/env")):
        try:
            with open(p) as f:
                m = re.search(r'ANTHROPIC_API_KEY=([^\s"\']+)', f.read())
                if m:
                    key = m.group(1)
                    break
        except OSError:
            pass
    key = key or os.environ.get("ANTHROPIC_API_KEY")
    if not key or not (text or "").strip():
        return None
    req = urllib.request.Request(
        "https://api.anthropic.com/v1/messages", method="POST",
        headers={"x-api-key": key, "anthropic-version": "2023-06-01",
                 "content-type": "application/json"},
        data=json.dumps({
            "model": _EXTRACT_MODEL, "max_tokens": max_tokens,
            "system": system,
            "messages": [{"role": "user", "content": (text or "")[:24000]}],
        }).encode())
    try:
        with urllib.request.urlopen(req, timeout=60) as r:
            resp = json.loads(r.read())
    except Exception:
        return None
    out = "".join(b.get("text", "") for b in resp.get("content", []))
    m = re.search(r"\{.*\}", out, re.DOTALL)
    if not m:
        return None
    try:
        return json.loads(m.group(0))
    except Exception:
        return None


def _extract_fields_llm(text):
    """Structured extraction via Claude. Returns the same dict shape as
    _extract_fields, or None if unavailable/failed (caller falls back to regex)."""
    d = _claude_json(_EXTRACT_SYSTEM, text)
    if not isinstance(d, dict):
        return None
    people = []
    for p in (d.get("people") or [])[:40]:
        if isinstance(p, dict) and str(p.get("name") or "").strip():
            basis = str(p.get("rate_basis") or "").lower()
            people.append({"name": str(p.get("name"))[:80].strip(),
                           "role": str(p.get("role") or "")[:80].strip(),
                           "location": str(p.get("location") or "")[:60].strip(),
                           "rate": str(p.get("rate") or "")[:40].strip(),
                           "rate_basis": basis if basis in ("hour", "month") else ""})
    return {
        "side": "vendor" if d.get("side") == "vendor" else "sea",
        "client": str(d.get("client") or ""),
        "agency": str(d.get("agency") or ""),
        "vendor": str(d.get("vendor") or ""),
        "amount": str(d.get("amount") or ""),
        "period_start": str(d.get("period_start") or ""),
        "period_end": str(d.get("period_end") or ""),
        "project_name": str(d.get("project_name") or ""),
        "people": people,
    }


def _extract_email_change(text):
    """What an email changes about a contract. Empty dict when the AI is
    unavailable — the user then types the figures in the same form."""
    d = _claude_json(_EMAIL_SYSTEM, text)
    if not isinstance(d, dict):
        return {}
    return {
        "is_change": bool(d.get("is_change", True)),
        "amount": str(d.get("amount") or ""),
        "period_start": str(d.get("period_start") or ""),
        "period_end": str(d.get("period_end") or ""),
        "project_name": str(d.get("project_name") or ""),
        "change_summary": str(d.get("change_summary") or "")[:600],
    }


def _extract_fields_best(text):
    """Prefer LLM extraction; fall back to regex, and backfill LLM blanks
    from the regex pass so we never lose a field the heuristics did catch."""
    llm = _extract_fields_llm(text)
    if not llm:
        return _extract_fields(text)
    rx = _extract_fields(text)
    for k, v in llm.items():
        if not v and rx.get(k):
            llm[k] = rx[k]
    return llm


def _norm_tokens(s):
    return set(re.findall(r"[a-z0-9]+", (s or "").lower())) - {
        "the", "a", "an", "of", "for", "and", "sow", "project", "cheil", "samsung"}


def _suggest_parents(data, vendor):
    """SEA contracts ranked by project-name overlap — candidate parents for a
    vendor contract (one SEA↔Cheil deal can hold many Cheil↔Vendor deals)."""
    mine = _norm_tokens(vendor.get("project_name"))
    out = []
    for c in data.get("contracts", []):
        if c.get("side") != "sea":
            continue
        theirs = _norm_tokens(c.get("project_name"))
        score = len(mine & theirs) / max(1, len(mine | theirs)) if (mine or theirs) else 0
        out.append((score, c))
    out.sort(key=lambda x: x[0], reverse=True)
    return out


def _contract_by_id(data, cid):
    return next((c for c in data.get("contracts", []) if c["id"] == cid), None)


def _lifecycle(data, c):
    """Contract lifecycle: cancelled (manual) or active. Amendments OVERRIDE by
    effective date (강프로 2026-07-25): every document in a chain stays on file
    and bills at its own rate, but only until the next document takes effect —
    executed amendments revise the fee section going forward rather than adding
    a second parallel schedule. Full removal = Cancel the old contract."""
    return "cancelled" if c.get("cancelled") else "active"


def _base_contract(data, c):
    """Follow the amends chain down to the base (original) document."""
    by_id = {x["id"]: x for x in data.get("contracts", [])}
    seen = set()
    while c.get("amends_id") in by_id and c["id"] not in seen:
        seen.add(c["id"])
        c = by_id[c["amends_id"]]
    return c


def _amendments_of(data, c):
    """Live amendment documents pointing (directly or via chain) at base c."""
    return [o for o in data.get("contracts", [])
            if o.get("amends_id") and o["id"] != c["id"]
            and _base_contract(data, o)["id"] == c["id"]
            and not o.get("cancelled")]


def _chain_docs(data, c):
    """The whole live amend chain c belongs to, ordered by effective date
    (period_start, falling back to upload order). Undated docs sort last so a
    half-extracted upload never silently truncates a dated one."""
    base = _base_contract(data, c)
    docs = [base] + _amendments_of(data, base)
    docs = [d for d in docs if not d.get("cancelled")]
    return sorted(docs, key=lambda d: (_parse_any_date(d.get("period_start")) or date.max,
                                       d.get("uploaded") or ""))


def _effective_end(data, c):
    """Where c actually stops billing: its own period_end, or the day before
    the next document in its chain takes effect — whichever comes first
    (강프로 2026-07-25, later document overrides earlier). None when c has no
    parseable end date."""
    own_end = _parse_any_date(c.get("period_end"), end=True)
    if own_end is None:
        return None
    nxt = None
    for d in _chain_docs(data, c):
        if d["id"] == c["id"]:
            continue
        s = _parse_any_date(d.get("period_start"))
        if s is None or s <= (_parse_any_date(c.get("period_start")) or date.min):
            continue
        if nxt is None or s < nxt:
            nxt = s
    if nxt is None:
        return own_end
    return min(own_end, nxt - timedelta(days=1))


def _display_state(data, c):
    """What to emphasise on screen (강프로 2026-07-27): the document that
    governs TODAY reads normally, everything behind it tones down.
      current    — in effect right now
      superseded — a later document in the chain has already taken over
      ended      — its own term ran out
      upcoming   — signed, but its effective date has not arrived
      cancelled  — struck off entirely
    """
    if c.get("cancelled"):
        return "cancelled"
    today = date.today()
    start = _parse_any_date(c.get("period_start"))
    own_end = _parse_any_date(c.get("period_end"), end=True)
    eff_end = _effective_end(data, c)
    if eff_end is not None and eff_end < today:
        # cut short by the next document, or simply over
        cut = own_end is not None and eff_end < own_end
        return "superseded" if cut else "ended"
    if start is not None and start > today:
        return "upcoming"
    return "current"


_STATE_CHIP = {
    "superseded": ("↺ Superseded", "A later document took over from this date — "
                                   "kept on file, no longer billing"),
    "ended": ("⏳ Ended", "Its term is over — kept on file for the record"),
    "upcoming": ("🕓 Not started", "Signed, but its effective date has not arrived yet"),
}


def _state_chip(state, eff_end=None, start=None):
    if state not in _STATE_CHIP:
        return ""
    label, tip = _STATE_CHIP[state]
    when = eff_end if state in ("superseded", "ended") else start
    tail = f" {when.isoformat()}" if when else ""
    return (f'<span class="dir-chip ctr-state" title="{tip}">{label}{tail}</span>')


def _auto_register_vendor(data, name):
    """Register the extracted vendor entity right away (강프로 2026-07-24)
    unless a vendor already matches by normalized name."""
    name = (name or "").strip()
    if not name:
        return
    norm = _norm_tokens(name)
    for v in data.get("vendors", []):
        vn = (v.get("name") or "").strip()
        if vn.lower() == name.lower() or (norm and _norm_tokens(vn) == norm):
            return
    data.setdefault("vendors", []).append(
        {"id": uuid.uuid4().hex[:8], "name": name, "entity_line": "", "msa_date": ""})


def _find_person(data, name):
    n = (name or "").strip().lower()
    return next((p for p in data.get("people", [])
                 if (p.get("name") or "").strip().lower() == n), None)


def _contract_groups(data):
    """1:many grouping. Returns (groups, orphans):
      groups  = [(sea_contract, [vendor children…]) …] for every SEA contract
      orphans = vendor contracts with no valid SEA parent
    A vendor's `linked_id` points to its SEA parent."""
    cs = data.get("contracts", [])
    by_id = {c["id"]: c for c in cs}
    # amendments attach to their base's group instead of forming their own
    seas = [c for c in cs if c.get("side") == "sea"
            and not (c.get("amends_id") in by_id)]
    sea_ids = {s["id"] for s in seas}
    children = {s["id"]: [] for s in seas}
    orphans = []
    for c in cs:
        if c.get("side") == "sea":
            continue
        pid = c.get("linked_id")
        if not pid and c.get("amends_id"):
            pid = _base_contract(data, c).get("linked_id")
        if pid in by_id and by_id[pid].get("side") == "sea" and pid not in sea_ids:
            pid = _base_contract(data, by_id[pid])["id"]  # parent is a SEA amendment
        if pid in sea_ids:
            children[pid].append(c)
        else:
            orphans.append(c)
    groups = [(s, children[s["id"]]) for s in seas]
    return groups, orphans


_SIDE_META = {
    "sea": ("SEA ↔ Cheil", "dir-samsung", "var(--accent)", "🔵"),
    "vendor": ("Cheil ↔ Vendor", "dir-agency", "var(--group-4)", "🟠"),
}


_SOURCE_META = {
    "email": ("✉️ Email change",
              "Change agreed by email — counts from its effective date, "
              "no signed amendment on file"),
    "schedule": ("📅 Monthly update",
                 "Revised monthly billing figures — no document name, the sheet "
                 "itself is the schedule from its first month"),
    "": ("↺ Amendment", "Overrides the earlier document from its effective date"),
}


def _source_chip(c):
    label, tip = _SOURCE_META.get(c.get("source") or "", _SOURCE_META[""])
    return ('<span class="dir-chip ctr-src" '
            f'title="{tip}">{label}</span>')


def _contract_card(c, draggable=False, show_title=True, status="active", data=None):
    """Contract box: counterparty, period, total amount as labeled rows
    (강프로 2026-07-24) — the group header above carries the contract name.
    Anything that no longer governs (superseded, ended, cancelled) is toned
    down so the live document is the one that reads (강프로 2026-07-27)."""
    label, chip, color, icon = _SIDE_META.get(c.get("side"), _SIDE_META["sea"])
    state = _display_state(data, c) if data is not None else (
        "cancelled" if status == "cancelled" else "current")
    state_chip = _state_chip(state, _effective_end(data, c) if data is not None else None,
                             _parse_any_date(c.get("period_start")))
    dim = " is-dim" if state in ("superseded", "ended", "cancelled") else ""
    party = c.get("vendor") or c.get("client") or "—"
    period = "—"
    if c.get("period_start") or c.get("period_end"):
        period = f'{_esc(c.get("period_start") or "…")} ~ {_esc(c.get("period_end") or "…")}'
    drag = ' draggable="true"' if draggable else ""
    title = ""
    if show_title:
        title = (f'<div class="ctr-title">'
                 f'{_esc(c.get("project_name") or c.get("filename") or "(untitled contract)")}</div>')
    return (
        f'<div class="ctr-card{dim}" data-cid="{c["id"]}"{drag} '
        f'onclick="openContract(\'{c["id"]}\')" style="cursor:pointer">'
        + ('<span class="ctr-grip">⠿</span>' if draggable else '')
        + f'<div class="ctr-top"><span class="dir-chip {chip}">{icon} {label}</span>'
        + ('<span class="dir-chip" style="color:var(--danger);background:rgba(248,113,113,.12)">❌ Cancelled</span>'
           if status == "cancelled" else state_chip)
        + (_source_chip(c) if c.get("amends_id") else '')
        + ('' if c.get("confirmed") or status != "active" else
           '<span title="Needs review & confirmation" style="font-size:.8rem">⚠️</span>')
        + '</div>'
        f'{title}'
        f'<div class="ctr-rows">'
        f'<span class="ctr-lb">Party</span><span>{_esc(party)}</span>'
        f'<span class="ctr-lb">Period</span><span>{period}</span>'
        f'<span class="ctr-lb">Amount</span><span class="ctr-amt">{_esc(c.get("amount") or "—")}</span>'
        f'</div></div>')


def _parse_any_date(s, end=False):
    """Contract dates are stored verbatim ('November 15, 2023',
    '5-September-2023', '2023-11-15') — try the shapes we actually see.
    Month-only terms ('June 2025', seen in executed amendments) resolve to the
    first of the month, or the last day of it when `end` is set."""
    s = (s or "").strip()
    if not s:
        return None
    d = _parse_date(s)
    if d:
        return d
    s2 = re.sub(r"[,\.]", " ", s).replace("-", " ").replace("/", " ")
    s2 = re.sub(r"\s+", " ", s2).strip()
    for fmt in ("%B %d %Y", "%d %B %Y", "%b %d %Y", "%d %b %Y",
                "%m %d %Y", "%Y %m %d"):
        try:
            return datetime.strptime(s2, fmt).date()
        except ValueError:
            pass
    for fmt in ("%B %Y", "%b %Y", "%Y %m"):
        try:
            d = datetime.strptime(s2, fmt).date()
        except ValueError:
            continue
        return d.replace(day=calendar.monthrange(d.year, d.month)[1]) if end else d
    return None


def _explicit_months(c):
    """{(y,m): amount} typed straight onto the document by a monthly billing
    sheet (강프로 2026-07-27) — no even spread, the sheet IS the schedule.
    None when the document carries no such sheet."""
    raw = c.get("month_amounts")
    if not isinstance(raw, dict) or not raw:
        return None
    out = {}
    for k, v in raw.items():
        ym = _parse_month_label(k)
        n = _num_or_none(v)
        if ym and n is not None:
            out[ym] = n
    return out or None


def _schedule_summary(months):
    """(total, period_start_iso, period_end_iso) for an explicit month map, so
    a sheet-only change still has the amount and dates every other part of the
    app reads off a contract."""
    ks = sorted(months)
    if not ks:
        return 0.0, "", ""
    y0, m0 = ks[0]
    y1, m1 = ks[-1]
    return (sum(months.values()), date(y0, m0, 1).isoformat(),
            date(y1, m1, calendar.monthrange(y1, m1)[1]).isoformat())


def _contract_month_amounts(c, data=None):
    """{(y,m): amount} — the contract total spread across its term with the
    same days/30 partial-month convention the SOW schedules use. None when
    period or amount is missing/unparseable.

    A document carrying its own monthly sheet skips the spread entirely and
    bills exactly what the sheet says (강프로 2026-07-27).

    With `data`, the schedule is cut off where the next document in the amend
    chain takes over (강프로 2026-07-25): the monthly rate still comes from this
    document's own amount over its own full term, but months already governed by
    a later amendment are dropped instead of double-counted."""
    explicit = _explicit_months(c)
    if explicit is not None:
        if data is None:
            return dict(explicit)
        eff_e = _effective_end(data, c)
        if eff_e is None:
            return dict(explicit)
        # the successor takes over from its own effective month onward
        cut = (eff_e.year, eff_e.month)
        return {ym: v for ym, v in explicit.items() if ym <= cut}
    s = _parse_any_date(c.get("period_start"))
    e = _parse_any_date(c.get("period_end"), end=True)
    amt = _num_or_none(c.get("amount"))
    if not s or not e or amt is None:
        return None
    spans = _month_spans(s, e)
    tot = sum(f for _, _, f in spans)
    if not spans or tot <= 0:
        return None
    per = amt / tot
    if data is not None:
        eff_e = _effective_end(data, c)
        if eff_e is not None and eff_e < e:
            spans = _month_spans(s, eff_e)
    return {(y, m): per * f for y, m, f in spans}


def _effective_amount(data, c):
    """What c actually contributes once later amendments override its tail.
    Falls back to the stated amount when the document has no usable period."""
    m = _contract_month_amounts(c, data)
    if m is None:
        return _num_or_none(c.get("amount"))
    return sum(m.values())


def _chain_effective(data, c):
    """(amount, start, end) for a whole amend chain after override — the sum of
    each live document's effective contribution, spanning from the base's start
    to the last document's end."""
    docs = _chain_docs(data, c)
    total, dates = 0.0, []
    got = False
    for d in docs:
        v = _effective_amount(data, d)
        if v is not None:
            total += v
            got = True
        s = _parse_any_date(d.get("period_start"))
        e = _effective_end(data, d)
        dates += [x for x in (s, e) if x]
    return (total if got else None,
            min(dates) if dates else None,
            max(dates) if dates else None)


def _contract_todos(data):
    """Mapping/confirmation gaps that must surface as Home To Dos (강프로
    2026-07-24): every contract needs an explicit user confirmation, every
    vendor contract a SEA parent, and every SEA contract either vendor
    children or a 'Cheil USA delivers itself' mark."""
    todos = []
    groups, orphans = _contract_groups(data)
    live = lambda c: _lifecycle(data, c) == "active"
    for c in data.get("contracts", []):
        if c.get("schedule_preview", {}).get("columns"):
            todos.append((c, "A monthly sheet is waiting — pick its money column "
                             "and apply it, or discard it"))
        if live(c) and not c.get("confirmed"):
            todos.append((c, "Not confirmed yet — review the extracted fields & mapping, then Confirm"))
        elif live(c) and c.get("amends_id") and not (
                c.get("amount") and c.get("period_start")):
            # a change with no figure or no effective date moves no money —
            # say so instead of letting it look reflected (강프로 2026-07-27)
            missing = " and ".join(
                x for x in ("an amount" if not c.get("amount") else "",
                            "an effective date" if not c.get("period_start") else "") if x)
            todos.append((c, f"Change is missing {missing} — it is not in the "
                             f"cashflow until you fill it in"))
    for o in orphans:
        if live(o):
            todos.append((o, "Vendor contract with no SEA deal — drag it under its SEA contract"))
    for sea, kids in groups:
        if live(sea) and not any(live(k) for k in kids) and not sea.get("self_delivered"):
            todos.append((sea, "No vendor mapped — align a vendor contract, or mark it Cheil-USA-self-delivered"))
    return todos


def _group_cashflow_table(sea, kids, data=None):
    """Per-group monthly cashflow: what Cheil bills SEA vs what it pays each
    vendor, month by month over the contract duration (강프로 2026-07-24).
    Only active contract versions are counted."""
    sea_m = _contract_month_amounts(sea, data)
    if data is not None:
        if _lifecycle(data, sea) != "active":
            sea_m = None
        else:
            for a in _amendments_of(data, sea):
                am = _contract_month_amounts(a, data)
                if am:
                    sea_m = sea_m or {}
                    for ym, v in am.items():
                        sea_m[ym] = sea_m.get(ym, 0.0) + v
        kids = [k for k in kids if _lifecycle(data, k) == "active"]
    kid_ms = [(k, _contract_month_amounts(k, data)) for k in kids]
    kid_ms = [(k, m) for k, m in kid_ms if m]
    months = set(sea_m or {})
    for _, m in kid_ms:
        months |= set(m)
    if not months or (not sea_m and not kid_ms):
        return ""
    months = sorted(months)

    def lb(ym):
        return date(ym[0], ym[1], 1).strftime("%b %y")

    def num_cells(mmap, cls=""):
        cells = []
        for ym in months:
            v = (mmap or {}).get(ym)
            cells.append(f'<td class="num {cls}">{_money(v) if v is not None else "–"}</td>')
        total = sum((mmap or {}).values())
        cells.append(f'<td class="num tot {cls}">{_money(total)}</td>')
        return "".join(cells)

    head = ("<tr><th class=\"pin\"></th>"
            + "".join(f"<th>{lb(ym)}</th>" for ym in months)
            + "<th>Total</th></tr>")
    rows = []
    if sea_m:
        rows.append('<tr><td class="pin">🔵 Bill to SEA</td>' + num_cells(sea_m, "bill") + '</tr>')
    for k, m in kid_ms:
        vname = _esc(k.get("vendor") or k.get("project_name") or "vendor")
        rows.append(f'<tr><td class="pin">🟠 {vname}</td>' + num_cells(m, "pay") + '</tr>')
    if sea_m and kid_ms:
        net = {}
        for ym in months:
            net[ym] = (sea_m.get(ym, 0.0)
                       - sum(m.get(ym, 0.0) for _, m in kid_ms))
        net_cells = []
        for ym in months:
            v = net[ym]
            cls = "pos" if v >= 0 else "neg"
            net_cells.append(f'<td class="num {cls}">{_money(v)}</td>')
        tot = sum(net.values())
        net_cells.append(f'<td class="num tot {"pos" if tot >= 0 else "neg"}">{_money(tot)}</td>')
        rows.append('<tr class="net"><td class="pin">Net (margin)</td>' + "".join(net_cells) + '</tr>')
    note = ("Contract totals spread evenly across each term (days/30 partial months) — "
            "except a change that came in as a monthly sheet, which bills its own "
            "figures. An amended document stops at the month its amendment takes "
            "effect. Documents without a parsed period or amount are omitted.")
    return (f'<details class="cf-details" open><summary>📅 Monthly billing &amp; payouts</summary>'
            f'<div class="cf-wrap"><table class="cf-table"><thead>{head}</thead>'
            f'<tbody>{"".join(rows)}</tbody></table></div>'
            f'<div class="sow-meta" style="margin-top:6px">{note}</div></details>')


def _sea_effective_total(sea, data=None):
    """What the SEA deal is worth once its amendments are applied — the number
    that belongs in the group header and drives the margin."""
    if data is None:
        return _num_or_none(sea.get("amount"))
    if _lifecycle(data, sea) != "active":
        return None
    total = _effective_amount(data, sea)
    for a in _amendments_of(data, sea):
        v = _effective_amount(data, a)
        if v is not None:
            total = (total or 0.0) + v
    return total


def _group_rollup(sea, kids, data=None):
    """SEA effective amount (base + live amendments, each cut off where the
    next one takes over) vs the sum of aligned vendor amounts on the same
    basis → margin chips. Cancelled versions are excluded."""
    sea_amt = _sea_effective_total(sea, data)
    if data is not None:
        kids = [k for k in kids if _lifecycle(data, k) == "active"]
    kid_amts = [a for a in ((_effective_amount(data, k) if data is not None
                             else _num_or_none(k.get("amount"))) for k in kids)
                if a is not None]
    parts = []
    if kid_amts:
        vsum = sum(kid_amts)
        parts.append(f'<span class="ctr-chip">Vendors {_money(vsum)}</span>')
        if sea_amt is not None:
            margin = sea_amt - vsum
            pct = f" · {margin / sea_amt * 100:.0f}%" if sea_amt else ""
            cls = "pos" if margin >= 0 else "neg"
            parts.append(f'<span class="ctr-chip {cls}">Margin {_money(margin)}{pct}</span>')
    return "".join(parts)


def _group_is_live(data, sea, kids):
    """A group still going: its SEA chain — the base OR any of its amendments —
    or any vendor card under it is in effect today. Checking the base alone
    would bury a deal the moment an amendment took over from it
    (강프로 2026-07-27)."""
    chain = [sea] + _amendments_of(data, sea) + list(kids)
    return any(_display_state(data, c) in ("current", "upcoming") for c in chain)


def _render_contracts_section(user, data):
    groups, orphans = _contract_groups(data)
    # live groups first, finished ones after — order is stable within each half
    groups = sorted(groups, key=lambda g: not _group_is_live(data, g[0], g[1]))
    gblocks = []
    for sea, kids in groups:
        # no title on grouped vendor cards — the group header already says
        # which deal this is (강프로 2026-07-24); unlinked pool keeps titles
        kid_cards = "".join(
            _contract_card(v, draggable=True, show_title=False,
                           status=_lifecycle(data, v), data=data) for v in kids)
        empty = ('<div class="ctr-drop-hint">Drag vendor contracts here</div>'
                 if not kids else "")
        gname = _esc(sea.get("project_name") or sea.get("filename") or "(untitled contract)")
        live = _group_is_live(data, sea, kids)
        gdim = "" if live else " is-dim"
        # title + value on the summary line so a folded block still says what it
        # is and what it is worth (강프로 2026-07-27); finished deals start folded
        sea_amt = _sea_effective_total(sea, data)
        amt_chip = (f'<span class="ctr-chip ctr-sea-amt">{_money(sea_amt)}</span>'
                    if sea_amt is not None else
                    f'<span class="ctr-chip">{_esc(sea.get("amount") or "no amount")}</span>')
        gblocks.append(
            f'<details class="ctr-group{gdim}" data-gid="{sea["id"]}"{" open" if live else ""}>'
            f'<summary class="ctr-group-hd"><span class="ctr-group-name">{gname}</span>'
            f'{amt_chip}{_group_rollup(sea, kids, data)}</summary>'
            f'<div class="ctr-group-body">'
            f'<div class="ctr-sea-col">{_contract_card(sea, show_title=False, status=_lifecycle(data, sea), data=data)}'
            + "".join(_contract_card(a, show_title=False, status=_lifecycle(data, a), data=data)
                      for a in _amendments_of(data, sea) if a.get("side") == "sea")
            + '</div>'
            f'<div class="ctr-ven-col" data-seadrop data-sea="{sea["id"]}">'
            f'{kid_cards}{empty}</div>'
            f'</div>'
            f'{_change_intake(data, _group_docs(data, sea, kids), sea["id"])}'
            f'{_group_cashflow_table(sea, kids, data)}</details>')
    groups_html = "".join(gblocks) or (
        '<div class="sow-meta" style="padding:22px;text-align:center">'
        'No SEA↔Cheil contracts yet — upload one to start a group.</div>')
    orphan_html = (
        '<div class="ctr-orphans" data-seadrop data-sea="">'
        '<div class="ctr-orphan-hd">Unlinked vendor contracts '
        '<span>· drag onto a group above to align them</span></div>'
        + ("".join(_contract_card(c, draggable=True, status=_lifecycle(data, c), data=data)
                   for c in orphans)
           if orphans else '<div class="ctr-drop-hint">None — drop a vendor card here to unlink it</div>')
        + _change_intake(data, orphans, "orphans")
        + '</div>')
    return f"""
<div class="ctr-dropzone" id="ctrDrop" data-filedrop tabindex="0">
  <input type="file" id="ctrFile" accept=".pdf,.docx,.doc,.txt" hidden>
  <b>⬆ Drop a NEW contract here</b> or click to upload — PDF/Word, parsed
  automatically. Changes to a contract already below go in its own block.
</div>
<div class="ctr-groups">{groups_html}</div>
{orphan_html}"""


def _live_change_targets(data, pool=None):
    """What a change can actually attach to (강프로 2026-07-27): the document
    governing today, or one that has not started yet. A superseded or ended
    document is history — amending it would move no money."""
    pool = data.get("contracts", []) if pool is None else pool
    return [c for c in pool if _display_state(data, c) in ("current", "upcoming")]


def _group_docs(data, sea, kids):
    """Every document that belongs to one deal block — the SEA chain and each
    aligned vendor chain — so the block's own intake targets only its own."""
    return [sea] + _amendments_of(data, sea) + list(kids)


def _doc_label(d):
    """Deal name plus the document's own reference when it adds something —
    a change logged inside a deal inherits the deal name, so the reference is
    the only thing telling two of them apart."""
    name = (d.get("project_name") or "").strip()
    ref = (d.get("filename") or "").strip()
    if name and ref and ref.lower() != name.lower():
        return f"{name} · {ref}"
    return name or ref or "(untitled contract)"


def _target_options(data, docs):
    """<optgroup>s of live documents, SEA side first."""
    out = []
    for side in ("sea", "vendor"):
        label, _chip, _col, icon = _SIDE_META[side]
        rows = []
        for c in sorted((x for x in docs if (x.get("side") or "sea") == side),
                        key=lambda x: (x.get("project_name") or "").lower()):
            party = c.get("vendor") or c.get("client") or ""
            name = c.get("project_name") or c.get("filename") or "(untitled contract)"
            tail = f" · {_esc(party)}" if party else ""
            amd = " ↺ latest amendment" if c.get("amends_id") else ""
            soon = " · not started yet" if _display_state(data, c) == "upcoming" else ""
            rows.append(f'<option value="{c["id"]}">{_esc(name)}{tail}{amd}{soon}</option>')
        if rows:
            out.append(f'<optgroup label="{icon} {label}">{"".join(rows)}</optgroup>')
    return "".join(out)


def _change_intake(data, docs, key):
    """Log a change to an existing contract — rendered INSIDE the deal block it
    belongs to (강프로 2026-07-27), never as one page-wide form, so the document
    being changed is the one you are already looking at and its deal name is
    inherited instead of retyped.

    Two ways in, because both arrive in real life:
      ✉️  the mail that agreed the change
      📅  the revised monthly figures alone — plenty of amendments carry no
          document name at all, only a new billing schedule.
    """
    live = _live_change_targets(data, docs)
    if not live:
        return ""
    opts = _target_options(data, live)
    return f"""
<details class="chg-intake" data-key="{_esc(key)}">
  <summary>✏️ Log a change to this deal <span class="eml-hint">— fee revision,
    extension or scope change on a contract already above</span></summary>
  <div class="chg-tabs">
    <button type="button" class="chg-tab is-on" onclick="chgTab(this,'eml')">✉️ From an email</button>
    <button type="button" class="chg-tab" onclick="chgTab(this,'sch')">📅 Monthly amounts only</button>
  </div>
  <form class="eml-form chg-pane" data-pane="eml"
        onsubmit="return chgSubmit(event,'/sow/contract/email')">
    <label class="eml-field eml-wide"><span>Which contract does this change?</span>
      <select class="slot" name="target" required>{opts}</select></label>
    <label class="eml-field eml-wide"><span>Paste the email — headers and body</span>
      <textarea class="slot" name="text" rows="5" placeholder="From: … / Sent: … / Subject: …

Paste the mail here and the fields below fill themselves."></textarea></label>
    <label class="eml-field eml-wide"><span>…or attach the mail</span>
      <input class="slot" type="file" name="file" accept=".msg,.eml,.pdf,.docx,.doc,.txt"></label>
    <label class="eml-field"><span>New amount</span>
      <input class="slot" name="amount" placeholder="$120,000 — blank = no money change"></label>
    <label class="eml-field"><span>Effective date</span>
      <input class="slot" name="effective" placeholder="August 1, 2026"></label>
    <label class="eml-field"><span>New end date</span>
      <input class="slot" name="end" placeholder="blank = end date unchanged"></label>
    <label class="eml-field eml-wide"><span>Reference (optional)</span>
      <input class="slot" name="name" placeholder="Amendment #2 — blank = the mail subject"></label>
    <label class="eml-field eml-wide"><span>What changed</span>
      <input class="slot" name="note" placeholder="Left blank, the AI writes this from the mail"></label>
    <div class="eml-actions">
      <span class="eml-note">The deal name and both parties are inherited from the
        contract above. Typed values win over the AI read, and you Confirm it on
        the next screen before it counts.</span>
      <button type="submit" class="btn btn-primary btn-sm">✉️ Log change</button>
    </div>
  </form>
  <form class="eml-form chg-pane" data-pane="sch" hidden
        onsubmit="return chgSubmit(event,'/sow/contract/schedule')">
    <label class="eml-field eml-wide"><span>Which contract does this change?</span>
      <select class="slot" name="target" required>{opts}</select></label>
    <label class="eml-field eml-wide"><span>Upload the revised schedule</span>
      <input class="slot" type="file" name="file" accept=".xlsx,.xlsm,.csv,.txt,.tsv"></label>
    <label class="eml-field eml-wide"><span>…or paste the months straight from Excel</span>
      <textarea class="slot" name="text" rows="4" placeholder="Jan-26	120,000
Feb-26	120,000
Mar-26	95,000

Months down a column, or across a header row with the figures underneath."></textarea></label>
    <label class="eml-field eml-wide"><span>Reference (optional)</span>
      <input class="slot" name="name" placeholder="blank = “Monthly update · Jan 2026 – Mar 2026”"></label>
    <label class="eml-field eml-wide"><span>What changed</span>
      <input class="slot" name="note" placeholder="e.g. media retainer reduced from April"></label>
    <div class="eml-actions">
      <span class="eml-note">No document name needed — the sheet is the schedule.
        It bills exactly these figures from its first month, and the contract
        above stops there instead of being billed twice.</span>
      <button type="submit" class="btn btn-primary btn-sm">📅 Apply schedule</button>
    </div>
  </form>
</details>"""


def _month_rows_html(months, cls=""):
    """month → amount list, taking either {(y,m):v} or {'YYYY-MM':v}."""
    items = sorted((_parse_month_label(k) if isinstance(k, str) else k, v)
                   for k, v in months.items())
    return "".join(f'<li><span>{date(y, m, 1).strftime("%b %Y")}</span>'
                   f'<b>{_money(v)}</b></li>' for (y, m), v in items)


def _schedule_preview_box(c, prev):
    """What the uploaded sheet says, BEFORE any of it counts (강프로
    2026-07-28). A real schedule ("15% cut simulation_1.xlsx") carries an
    Original Cost and an Adjusted Cost column and repeats a month; choosing for
    the user there is a wrong answer that looks like a right one, so the read is
    shown, the money column is picked by hand, and only then does it apply."""
    cols = prev.get("columns") or []
    chosen = next((x for x in cols if x["key"] == prev.get("chosen")), cols[0])
    warn = "".join(f'<div class="ppl-sheet-err">⚠ {_esc(w)}</div>'
                   for w in prev.get("warnings") or [])
    picker = ""
    if len(cols) > 1:
        opts = "".join(
            f'<label class="sch-pick{" is-on" if x["key"] == chosen["key"] else ""}">'
            f'<input type="radio" name="schcol_{c["id"]}"'
            f'{" checked" if x["key"] == chosen["key"] else ""} '
            f'onchange="ctrPost(\'/sow/contract/schedule_pick\','
            f'{{id:\'{c["id"]}\',col:\'{x["key"]}\'}})">'
            f'<span>{_esc(x["label"])}</span><b>{_money(x["total"])}</b></label>'
            for x in cols)
        picker = (f'<div class="sow-meta">Which column should this change bill?</div>'
                  f'<div class="sch-picks">{opts}</div>')
    src = prev.get("src")
    return f"""
<div class="ctr-linkbox sch-preview">
  <b>📅 Read from the sheet — not applied yet</b>
  <div class="sow-meta">{_esc(src or "pasted range")} · {_esc(prev.get("note") or "")}</div>
  {warn}
  {picker}
  <div class="sow-meta">Billing <b>{_money(chosen["total"])}</b> across
    {len(chosen["months"])} month(s), exactly as listed — no even spread. Applying
    it stops <b>this</b> contract the month before the first one.</div>
  <ul class="ctr-kidlist sch-list">{_month_rows_html(chosen["months"])}</ul>
  <div style="display:flex;gap:8px;flex-wrap:wrap">
    <button class="btn btn-primary btn-sm" type="button"
      onclick="ctrPost('/sow/contract/schedule_apply',{{id:'{c["id"]}'}})">✅ Apply as a change</button>
    <button class="btn btn-secondary btn-sm" type="button"
      onclick="ctrPost('/sow/contract/schedule_discard',{{id:'{c["id"]}'}})">Discard</button>
  </div>
</div>"""


def _schedule_box(c, months):
    """The monthly figures this document bills, as read off the sheet — the
    document's amount and period are derived from them, so they are shown here
    rather than hidden behind a single total (강프로 2026-07-27)."""
    total, _s, _e = _schedule_summary(months)
    rows = _month_rows_html(months)
    note = c.get("schedule_note")
    src = f'<div class="sow-meta">Read from: <b>{_esc(note)}</b></div>' if note else ""
    return f"""
<div class="ctr-linkbox">
  <b>📅 Monthly billing schedule ({len(months)} month(s))</b>
  {src}
  <div class="sow-meta">These figures bill as they stand — no even spread — and the
    document they change stops the month before the first one. Totals
    <b>{_money(total)}</b>.</div>
  <ul class="ctr-kidlist sch-list">{rows}</ul>
  <div>
    <button class="btn btn-secondary btn-sm" type="button"
      onclick="if(confirm('Drop the monthly schedule? The total stays as the contract amount and goes back to being spread evenly across the period.'))ctrPost('/sow/contract/schedule_clear',{{id:'{c["id"]}'}})">
      🗑 Clear schedule</button>
  </div>
</div>"""


def _people_sheet_box(c):
    """One sheet intake per contract, for BOTH sheets that turn up against a
    contract (강프로 2026-07-28): the team roster and the revised monthly
    billing schedule. It used to take only the roster and reject a schedule
    with "No name column found", which is a dead end when the schedule is what
    you have in your hand — so what the sheet contains decides where it goes."""
    axis = "Contract rate (what Cheil pays)" if c.get("side") == "vendor" else \
           "Selling rate (what Cheil bills)"
    err = c.get("people_sheet_error")
    note = c.get("people_sheet_note")
    status = ""
    if err:
        status = f'<div class="ppl-sheet-err">⚠ {_esc(err)}</div>'
    elif note:
        status = f'<div class="sow-meta">Last sheet read: <b>{_esc(note)}</b></div>'
    return f"""
<details class="ctr-linkbox ppl-sheet">
  <summary><b>📥 Upload a sheet for this contract</b>
    <span class="sow-meta">— the team, or a revised monthly schedule</span></summary>
  {status}
  <div class="sow-meta" style="margin:6px 0 8px">The sheet decides what it is:
    <br>👥 <b>Team</b> — a <b>Resource/Name</b> column, plus any of Role/Function,
    Rate, Location, Email. Rates land as the <b>{axis}</b>.
    <br>📅 <b>Monthly schedule</b> — month labels (Jan-26, 2026-01) with the
    amount against each. It comes back for review as a change to this contract.
    <br>Nothing is saved either way until you review it.</div>
  <form class="ppl-sheet-form" onsubmit="return pplSheet(event,'{c["id"]}')">
    <input class="slot" type="file" name="file" accept=".xlsx,.xlsm,.csv,.txt,.tsv">
    <textarea class="slot" name="text" rows="3"
      placeholder="…or paste the rows straight from Excel (header row included)"></textarea>
    <button class="btn btn-primary btn-sm" type="submit">📥 Read sheet</button>
  </form>
</details>"""


def _render_contract_frag(user, data, cid):
    """Popup body for one contract: extracted fields (editable) + text preview
    + original download + link controls."""
    c = _contract_by_id(data, cid)
    if not c:
        return '<div class="cmodal-body"><p>Contract not found.</p></div>'
    label, chip, color, icon = _SIDE_META.get(c.get("side"), _SIDE_META["sea"])

    # link controls — 1:many (a vendor belongs under one SEA; a SEA holds many)
    if c.get("side") == "vendor":
        parent = _contract_by_id(data, c.get("linked_id")) if c.get("linked_id") else None
        if parent and parent.get("side") == "sea":
            link_html = (
                f'<div class="ctr-linked">🔗 Under <b>{_esc(parent.get("project_name") or parent.get("filename"))}</b> '
                '(SEA ↔ Cheil)'
                f'<button class="btn btn-danger btn-sm" onclick="ctrPost(\'/sow/contract/unlink\',{{id:\'{c["id"]}\'}})">Unlink</button></div>')
        else:
            opts = []
            for score, cand in _suggest_parents(data, c):
                tag = " ★ suggested" if score > 0 else ""
                opts.append(f'<option value="{cand["id"]}">{_esc(cand.get("project_name") or cand.get("filename"))}{tag}</option>')
            if opts:
                link_html = (
                    '<div class="ctr-linkbox"><b>Align under a SEA ↔ Cheil contract:</b>'
                    f'<select id="lnk_{c["id"]}" class="slot">{"".join(opts)}</select>'
                    f'<button class="btn btn-primary btn-sm" onclick="ctrAssign(\'{c["id"]}\')">🔗 Confirm</button>'
                    '<div class="sow-meta">★ = project-name match. You can also drag the card onto a SEA contract.</div></div>')
            else:
                link_html = '<div class="sow-meta">No SEA ↔ Cheil contract to align under yet.</div>'
    else:
        _, kids = next(((s, k) for s, k in _contract_groups(data)[0] if s["id"] == c["id"]),
                       (c, []))
        if kids:
            items = "".join(
                f'<li><span>{_esc(k.get("project_name") or k.get("filename"))}'
                f' <span class="sow-meta">{_esc(k.get("vendor") or "")}</span></span>'
                f'<b>{_esc(k.get("amount") or "—")}</b></li>' for k in kids)
            rollup = _group_rollup(c, kids, data)
            link_html = (f'<div class="ctr-linkbox"><b>Aligned vendor contracts ({len(kids)})'
                         f'{("&nbsp;" + rollup) if rollup else ""}</b>'
                         f'<ul class="ctr-kidlist">{items}</ul>'
                         '<div class="sow-meta">Drag a vendor card onto this group to add more.</div></div>')
        else:
            link_html = ('<div class="sow-meta" style="margin-top:16px;padding-top:14px;'
                         'border-top:1px solid var(--border)">No vendor contracts aligned yet — '
                         'drag a vendor card onto this group on the main screen.</div>')

    def fld(lbl, key, val, ro=False, tip=""):
        extra = ' readonly' + (f' title="{tip}"' if tip else "") if ro else ""
        return (f'<label class="ctr-fld"><span>{lbl}</span>'
                f'<input class="slot" name="{key}" value="{_esc(val)}"{extra}></label>')

    # side-aware fields: the Cheil entity is a constant on both sides, and a
    # vendor deal's payer is always Cheil — neither is worth an editable box.
    if c.get("side") == "vendor":
        party_fld = fld("Vendor / Contractor", "vendor", c.get("vendor"))
        parties_note = f"{CHEIL_ENTITY} (payer) ↔ {c.get('vendor') or 'vendor'}"
    else:
        party_fld = fld("Client (payer)", "client", c.get("client") or SAMSUNG_ENTITY)
        parties_note = f"{c.get('client') or SAMSUNG_ENTITY} ↔ {CHEIL_ENTITY} (agency)"
    # a document billing off its own monthly sheet has no single figure to type —
    # the amount, the start and the end are all read off the schedule below
    months = _explicit_months(c)
    sch_tip = ("Derived from the monthly schedule below — clear the schedule to "
               "type these in by hand")
    if months:
        amount_fld = fld("Contract amount", "amount", c.get("amount"), ro=True, tip=sch_tip)
        schedule_html = _schedule_box(c, months)
    else:
        amount_fld = fld("Contract amount", "amount", c.get("amount"))
        schedule_html = ""
    if c.get("schedule_preview", {}).get("columns"):
        schedule_html = _schedule_preview_box(c, c["schedule_preview"]) + schedule_html
    if c.get("schedule_error"):
        schedule_html = (f'<div class="ctr-linkbox"><div class="ppl-sheet-err">⚠ '
                         f'{_esc(c["schedule_error"])}</div></div>') + schedule_html
    # people found in the contract — 1차 정리 후 리뷰 요청 후 저장 (강프로 2026-07-24):
    # nothing lands in the roster until the user confirms here.
    people_html = ""
    pending = c.get("people_pending") or []
    if pending:
        aff = c.get("vendor") if c.get("side") == "vendor" else "Cheil"
        rows = []
        rate_axis = "cost" if c.get("side") == "vendor" else "selling"
        for i, p in enumerate(pending):
            known = _find_person(data, p.get("name"))
            note = (' <span class="sow-meta">already in roster — fills blanks & links this contract</span>'
                    if known else "")
            # what the extractor read beyond the name — shown so the user can
            # spot a bad rate before it lands in the roster (강프로 2026-07-27)
            bits = []
            if p.get("rate"):
                per = {"hour": "/hr", "month": "/mo"}.get(p.get("rate_basis"), "")
                bits.append(f'{_esc(p["rate"])}{per} {rate_axis}')
            if p.get("location"):
                bits.append(_esc(p["location"]))
            extra = (f' <span class="ppl-extra">{" · ".join(bits)}</span>' if bits else "")
            rows.append(
                f'<label class="ppl-row"><input type="checkbox" value="{i}" checked>'
                f'<span><b>{_esc(p.get("name"))}</b>'
                + (f' · {_esc(p.get("role"))}' if p.get("role") else "")
                + f'{extra} <span class="sow-meta">({_esc(aff)})</span>{note}</span></label>')
        people_html = (
            f'<div class="ctr-linkbox" id="pplBox"><b>👥 People found in this contract ({len(pending)})</b>'
            f'<div class="sow-meta">Name, role, rate and location as read from the resource table — '
            f'rates land as the <b>{rate_axis} rate</b> because this is a '
            f'{"Cheil ↔ Vendor" if c.get("side") == "vendor" else "SEA ↔ Cheil"} contract. '
            f'Review, uncheck any noise, then save to the People roster.</div>'
            + "".join(rows) +
            '<div style="display:flex;gap:8px;margin-top:4px">'
            f'<button class="btn btn-primary btn-sm" type="button" onclick="ctrPeopleSave(\'{c["id"]}\')">💾 Save selected to People</button>'
            f'<button class="btn btn-secondary btn-sm" type="button" onclick="ctrPost(\'/sow/contract/people_dismiss\',{{id:\'{c["id"]}\'}})">Dismiss</button>'
            '</div></div>')
    people_html += _people_sheet_box(c)
    if c.get("confirmed"):
        confirm_btn = (f'<button class="btn btn-secondary btn-sm" type="button" '
                       f'title="Confirmed {_esc((c.get("confirmed_at") or "")[:10])} — click to un-confirm" '
                       f'onclick="ctrPost(\'/sow/contract/confirm\',{{id:\'{c["id"]}\'}})">✔ Confirmed</button>')
    else:
        confirm_btn = (f'<button class="btn btn-primary btn-sm" type="button" '
                       f'title="Mark the fields and SEA↔vendor mapping as reviewed" '
                       f'onclick="ctrPost(\'/sow/contract/confirm\',{{id:\'{c["id"]}\'}})">✅ Confirm contract</button>')
    # lifecycle: a later document overrides the earlier one from its effective
    # date; Cancel drops a doc from every total (강프로 2026-07-25)
    lc = _lifecycle(data, c)
    amds = _amendments_of(data, c) if not c.get("amends_id") else []
    prev = _contract_by_id(data, c.get("amends_id")) if c.get("amends_id") else None
    amend_cands = [o for o in data.get("contracts", [])
                   if o["id"] != c["id"] and o.get("side") == c.get("side")
                   and o.get("amends_id") != c["id"]]
    amend_opts = '<option value="">— none (standalone / original contract) —</option>' + "".join(
        f'<option value="{o["id"]}"{" selected" if c.get("amends_id") == o["id"] else ""}>'
        f'{_esc(o.get("project_name") or o.get("filename") or o["id"])}</option>'
        for o in amend_cands)
    amend_lb, amend_tip = _SOURCE_META.get(c.get("source") or "", _SOURCE_META[""])
    lc_chip = ('<span class="ctr-chip neg">❌ Cancelled — excluded from totals</span>'
               if lc == "cancelled" else
               (f'<span class="ctr-chip ctr-src-chip" title="{amend_tip}">{amend_lb}'
                f' — overrides the earlier document</span>'
                if c.get("amends_id") else '<span class="ctr-chip pos">● Active</span>'))
    eff_html = ""
    if amds:
        eff_amt, cs, ce = _chain_effective(data, c)
        span = (f'{cs.isoformat()} ~ {ce.isoformat()}' if cs and ce else "")
        amd_list = ""
        for d in _chain_docs(data, c):
            own_e = _parse_any_date(d.get("period_end"), end=True)
            de = _effective_end(data, d)
            cut = de is not None and own_e is not None and de < own_e
            note = (f' <span style="color:var(--muted)">— superseded from '
                    f'{(de + timedelta(days=1)).isoformat()}, counts '
                    f'{_money(_effective_amount(data, d) or 0)}</span>' if cut else "")
            link = ("" if d["id"] == c["id"] else
                    f' <a href="#" onclick="openContract(\'{d["id"]}\');return false"'
                    f' style="color:var(--accent)">open →</a>')
            icon_d = ("📅" if d.get("source") == "schedule" else
                      "✉️" if d.get("source") == "email" else
                      "↺" if d.get("amends_id") else "📄")
            amd_list += (
                f'<li><span>{icon_d} '
                f'{_esc(_doc_label(d))}{link}{note}</span>'
                f'<b>{_esc(d.get("amount") or "—")}</b></li>')
        eff_html = (f'<div class="sow-meta"><b>Effective after {len(amds)} amendment(s):</b> '
                    f'{_money(eff_amt or 0)}{(" · " + span) if span else ""}'
                    f' <span style="color:var(--muted)">— each document bills until the '
                    f'next one takes effect</span></div>'
                    f'<ul class="ctr-kidlist">{amd_list}</ul>')
    prev_html = (f'<div class="sow-meta">↺ This document amends <b>{_esc(prev.get("project_name") or prev.get("filename"))}</b> '
                 f'<a href="#" onclick="openContract(\'{prev["id"]}\');return false" style="color:var(--accent)">open →</a>'
                 f' — that document bills up to this one&rsquo;s start date, then this one governs.</div>'
                 if prev else "")
    cancel_btn = (f'<button class="btn btn-secondary btn-sm" type="button" '
                  f'onclick="ctrPost(\'/sow/contract/cancel\',{{id:\'{c["id"]}\'}})">↩ Reactivate</button>'
                  if c.get("cancelled") else
                  f'<button class="btn btn-danger btn-sm" type="button" '
                  f'onclick="if(confirm(\'Cancel this contract? It stays on file but drops out of all totals. For a re-contract, upload the new document (as a fresh contract, or as an Amendment if it only adds to a live one).\'))ctrPost(\'/sow/contract/cancel\',{{id:\'{c["id"]}\'}})">❌ Cancel contract</button>')
    lifecycle_html = (
        f'<div class="ctr-linkbox"><b>Lifecycle</b>'
        f'<div style="display:flex;gap:8px;align-items:center;flex-wrap:wrap">{lc_chip}{cancel_btn}</div>'
        f'{eff_html}{prev_html}'
        f'<label style="display:flex;gap:8px;align-items:center;font-size:.8rem;flex-wrap:wrap">Amendment of (base contract):'
        f'<select class="slot" style="flex:1;min-width:200px" '
        f'onchange="ctrPost(\'/sow/contract/amends\',{{id:\'{c["id"]}\',prev:this.value}})">{amend_opts}</select></label>'
        f'<div class="sow-meta">Amendment flow: upload the amendment document as a new contract, open it, and set "Amendment of" to the base — documents are then ordered by effective date and each one bills at its own rate only until the next takes over, so a revised fee schedule replaces the old one instead of stacking on top of it. Dropping a document from every total entirely = Cancel it instead.</div>'
        f'</div>')

    if c.get("side") != "vendor":
        link_html += (
            f'<label style="display:flex;gap:8px;align-items:center;font-size:.8rem;'
            f'margin-top:10px;cursor:pointer;color:var(--text)">'
            f'<input type="checkbox"{" checked" if c.get("self_delivered") else ""} '
            f'onchange="ctrPost(\'/sow/contract/selfdeliver\',{{id:\'{c["id"]}\'}})">'
            f'🏠 Cheil USA delivers this itself — no vendor contract expected</label>')

    # actual document viewer (강프로 2026-07-24): PDFs render inline (lazy —
    # the iframe only loads when opened); Word files can't render in-browser,
    # so they get a clear pointer to the original download.
    ext = (c.get("ext") or "").lower()
    if ext == "pdf":
        viewer_html = (
            '<details class="ctr-prev" '
            'ontoggle="if(this.open){var f=this.querySelector(\'iframe\');'
            'if(f&&!f.src)f.src=f.dataset.src;}">'
            '<summary>📄 View contract (original PDF)</summary>'
            f'<iframe data-src="/sow/contract/file?id={c["id"]}" '
            'style="width:100%;height:72vh;border:1px solid var(--border);'
            'border-radius:8px;margin-top:8px;background:#fff"></iframe></details>')
    elif ext in ("docx", "doc"):
        viewer_html = (
            '<div class="sow-meta" style="margin-top:14px">📄 Word document — '
            'browsers can\'t render it inline; use '
            f'<a href="/sow/contract/file?id={c["id"]}" target="_blank" '
            'style="color:var(--accent)">⬇ Original</a> to open it in Word.</div>')
    else:
        viewer_html = ""
    uploaded = (c.get("uploaded") or "")[:10]
    preview = _esc(_contract_text(user, c)[:6000])
    is_email = c.get("source") == "email"
    is_sched = c.get("source") == "schedule"
    em = c.get("email_meta") or {}
    email_html = ""
    if is_sched and c.get("change_note"):
        email_html = (f'<div class="eml-note-box"><b>📅 What changed</b><br>'
                      f'{_esc(c.get("change_note"))}</div>')
    if is_email:
        rows = "".join(
            f'<span class="lb">{k.capitalize()}</span><span>{_esc(v)}</span>'
            for k, v in (("from", em.get("from")), ("sent", em.get("sent")),
                         ("subject", em.get("subject"))) if v)
        email_html = (f'<div class="eml-meta">{rows}</div>' if rows else "")
        if c.get("change_note"):
            email_html += (f'<div class="eml-note-box"><b>✉️ What changed</b><br>'
                           f'{_esc(c.get("change_note"))}</div>')
    prev_label = ("✉️ Email text" if is_email else
                  "📅 Schedule as read" if is_sched else "🔤 Extracted text")
    # only offer the download when a file was actually kept — a logged change
    # often has none, and a dead ⬇ Original reads as a broken contract
    orig_btn = (f'<a class="btn btn-secondary btn-sm" href="/sow/contract/file?id={c["id"]}"'
                f' target="_blank">⬇ Original</a>'
                if not c.get("source") or c.get("has_file") else "")
    return f"""
<div class="cmodal-head">
  <span class="dir-chip {chip}">{icon} {label}</span>
  <span class="cmodal-file">{_esc(c.get("filename") or "")}{(' · ' + uploaded) if uploaded else ''}</span>
  <button class="cmodal-x" onclick="closeContract()">✕</button>
</div>
<div class="cmodal-body">
  <div class="ctr-parties">{_esc(parties_note)}</div>
  {email_html}
  <form onsubmit="ctrSave('{c['id']}');return false" class="ctr-form">
    <div class="ctr-grid">
      {fld("Project name", "project_name", c.get("project_name"))}
      {party_fld}
      {amount_fld}
      <label class="ctr-fld"><span>Side</span>
        <select class="slot" name="side">
          <option value="sea"{' selected' if c.get('side')=='sea' else ''}>SEA ↔ Cheil</option>
          <option value="vendor"{' selected' if c.get('side')=='vendor' else ''}>Cheil ↔ Vendor</option>
        </select></label>
      {fld("Effective date" if is_email else "Period start", "period_start",
           c.get("period_start"), ro=bool(months), tip=sch_tip)}
      {fld("Period end", "period_end", c.get("period_end"), ro=bool(months), tip=sch_tip)}
    </div>
    <div class="ctr-actions">
      <button class="btn btn-primary btn-sm" type="submit">💾 Save fields</button>
      {confirm_btn}
      <button class="btn btn-secondary btn-sm" type="button" onclick="ctrReparse('{c['id']}',this)">🪄 Re-read with AI</button>
      {orig_btn}
      <button class="btn btn-danger btn-sm" type="button" onclick="ctrPost('/sow/contract/delete',{{id:'{c['id']}'}})">🗑 Delete</button>
    </div>
  </form>
  {link_html}
  {lifecycle_html}
  {schedule_html}
  {people_html}
  {viewer_html}
  <details class="ctr-prev"><summary>{prev_label}</summary><pre>{preview}</pre></details>
</div>"""


_CTR_CSS = """
.ctr-dropzone{display:block;width:100%;text-align:center;background:var(--surface-2,var(--surface));border:2px dashed var(--border-bright);border-radius:var(--radius-lg);padding:16px 18px;margin-bottom:16px;cursor:pointer;color:var(--text-muted);font-size:.82rem;transition:.15s}
.ctr-dropzone:hover{border-color:var(--accent);color:var(--text)}
.ctr-dropzone b{color:var(--text)}
.ctr-dropzone.drag-over{border-color:var(--accent);background:rgba(56,189,248,.10);color:var(--text)}
/* ── tone-down: superseded / ended / cancelled documents step back so the
   document governing today is the one that reads (강프로 2026-07-27) ── */
.ctr-card.is-dim{opacity:.5;filter:saturate(.55)}
.ctr-card.is-dim:hover,.ctr-card.is-dim:focus-within{opacity:1;filter:none}
.ctr-card.is-dim .ctr-amt{color:var(--text-muted);font-weight:600}
.ctr-group.is-dim>.ctr-group-hd,.ctr-group.is-dim>.ctr-group-body,
.ctr-group.is-dim .cf-wrap,.ctr-group.is-dim .cf-details{opacity:.62}
.ctr-group.is-dim:hover>.ctr-group-hd,.ctr-group.is-dim:hover>.ctr-group-body,
.ctr-group.is-dim:hover .cf-wrap,.ctr-group.is-dim:hover .cf-details{opacity:1}
.ctr-group.is-dim .ctr-card.is-dim{opacity:.72}
.dir-chip.ctr-state{color:var(--text-muted);background:var(--surface-3);
  border:1px solid var(--border);font-weight:600}
/* ── per-deal change intake: lives inside the block it changes ── */
.dir-chip.ctr-src{color:var(--info);background:var(--surface-3);border:1px solid var(--border)}
.ctr-chip.ctr-src-chip{color:var(--info);border-color:var(--info)}
.chg-intake{margin-top:12px;background:var(--surface);border:1px solid var(--border);border-radius:var(--radius-lg);padding:10px 14px}
.chg-intake>summary{cursor:pointer;list-style:none;font-size:.8rem;font-weight:700;color:var(--text)}
.chg-intake>summary::-webkit-details-marker{display:none}
.chg-intake>summary::before{content:"▸ ";color:var(--text-muted)}
.chg-intake[open]>summary::before{content:"▾ "}
.chg-tabs{display:flex;gap:6px;margin-top:10px;flex-wrap:wrap}
.chg-tab{background:var(--surface-2,var(--surface));border:1px solid var(--border);color:var(--text-muted);border-radius:var(--radius-full);padding:5px 14px;font-size:.74rem;font-weight:700;cursor:pointer;transition:.15s}
.chg-tab:hover{color:var(--text);border-color:var(--border-bright)}
.chg-tab.is-on{background:var(--accent-glow);color:var(--accent);border-color:var(--accent)}
.chg-pane[hidden]{display:none}
.sch-list li span{color:var(--text-muted);font-size:.78rem}
.sch-preview{border:1px solid var(--accent);border-radius:var(--radius-lg);padding:12px 14px;background:var(--accent-glow)}
.sch-picks{display:flex;gap:8px;flex-wrap:wrap}
.sch-pick{display:flex;align-items:center;gap:8px;padding:8px 12px;min-height:44px;border:1px solid var(--border);border-radius:var(--radius-md);background:var(--surface);cursor:pointer;font-size:.8rem;flex:1;min-width:180px}
.sch-pick.is-on{border-color:var(--accent);color:var(--text)}
.sch-pick b{margin-left:auto;font-variant-numeric:tabular-nums;color:var(--success)}
.eml-hint{font-weight:500;color:var(--text-muted);font-size:.78rem}
.eml-form{display:grid;grid-template-columns:repeat(3,minmax(0,1fr));gap:10px;margin-top:12px}
.eml-field{display:flex;flex-direction:column;gap:4px;min-width:0}
.eml-field>span{font-size:.64rem;text-transform:uppercase;letter-spacing:.05em;color:var(--text-muted)}
.eml-field.eml-wide{grid-column:1/-1}
.eml-form textarea.slot{width:100%;font-size:.82rem}
.eml-actions{grid-column:1/-1;display:flex;align-items:center;gap:12px;flex-wrap:wrap}
.eml-note{flex:1;min-width:200px;font-size:.7rem;color:var(--text-muted)}
.eml-meta{display:grid;grid-template-columns:auto minmax(0,1fr);gap:3px 10px;background:var(--surface-2);border:1px solid var(--border);border-radius:var(--radius-md);padding:10px 12px;margin-bottom:12px;font-size:.76rem}
.eml-meta .lb{color:var(--text-muted);text-transform:uppercase;font-size:.64rem;letter-spacing:.05em}
.eml-note-box{background:rgba(129,140,248,.10);border:1px solid rgba(129,140,248,.35);border-radius:var(--radius-md);padding:10px 12px;margin-bottom:12px;font-size:.8rem;line-height:1.5}
.ctr-groups{display:flex;flex-direction:column;gap:16px}
.ctr-group{background:var(--surface-2,var(--surface));border:1px solid var(--border);border-radius:var(--radius-xl);padding:14px 16px}
details.ctr-group>summary{list-style:none;cursor:pointer}
details.ctr-group>summary::-webkit-details-marker{display:none}
details.ctr-group>summary::before{content:"▸";color:var(--text-muted);font-size:.8rem;flex-shrink:0}
details.ctr-group[open]>summary::before{content:"▾"}
details.ctr-group>summary:hover .ctr-group-name{color:var(--accent)}
.ctr-group-hd{display:flex;align-items:center;gap:10px;flex-wrap:wrap;margin-bottom:0}
details.ctr-group[open]>.ctr-group-hd{margin-bottom:12px}
.ctr-chip.ctr-sea-amt{font-weight:800;color:var(--success);border-color:rgba(52,211,153,.35)}
.ctr-group-name{font-size:1.02rem;font-weight:800;letter-spacing:-.01em;color:var(--text)}
.ctr-chip{font-size:.68rem;font-weight:700;padding:3px 10px;border-radius:10px;background:var(--surface);border:1px solid var(--border);color:var(--text-muted);font-variant-numeric:tabular-nums}
.ctr-chip.pos{color:var(--success);border-color:rgba(52,211,153,.35)}
.ctr-chip.neg{color:var(--danger);border-color:rgba(248,113,113,.4)}
.ctr-group-body{display:grid;grid-template-columns:minmax(0,5fr) minmax(0,7fr);gap:14px;align-items:start}
.ctr-sea-col .ctr-card{border-left:3px solid var(--accent)}
.ctr-ven-col{display:flex;flex-direction:column;gap:8px;min-height:40px;padding:6px;border:1px dashed transparent;border-radius:var(--radius-md);transition:.15s}
.ctr-ven-col.drag-over,.ctr-orphans.drag-over{background:rgba(56,189,248,.10);border-color:var(--accent)}
.ctr-ven-col .ctr-card{border-left:3px solid var(--group-4)}
.ctr-drop-hint{font-size:.72rem;color:var(--text-muted);font-style:italic;padding:6px 4px}
/* on a wide screen the vendor cards sit side by side instead of stretching to
   850px for three short rows (강프로 2026-07-28) */
@media(min-width:1400px){
  .ctr-ven-col{display:grid;grid-template-columns:repeat(auto-fill,minmax(330px,1fr));align-content:start}
  .ctr-ven-col .ctr-drop-hint{grid-column:1/-1}
}
.ctr-orphans{margin-top:16px;padding:12px;border:1px dashed var(--border-bright);border-radius:var(--radius-lg);display:flex;flex-direction:column;gap:8px;transition:.15s}
.ctr-orphan-hd{font-size:.74rem;font-weight:700;color:var(--text)}
.ctr-orphan-hd span{font-weight:400;color:var(--text-muted)}
.ctr-card{position:relative;background:var(--surface);border:1px solid var(--border);border-radius:var(--radius-lg);padding:14px 16px;transition:.15s}
.ctr-card:hover{border-color:var(--accent);box-shadow:var(--shadow-md)}
.ctr-card[draggable="true"]{padding-left:30px}
.ctr-card.dragging{opacity:.45}
.ctr-grip{position:absolute;left:9px;top:50%;transform:translateY(-50%);color:var(--text-muted);cursor:grab;font-size:1rem;letter-spacing:-2px}
.ctr-top{display:flex;align-items:center;justify-content:space-between;gap:8px;margin-bottom:6px;flex-wrap:wrap;min-width:0}
.ctr-amt{font-weight:800;color:var(--success);font-variant-numeric:tabular-nums;font-size:.86rem}
.ctr-title{font-weight:700;font-size:.9rem;color:var(--text);margin-bottom:6px}
.ctr-rows{display:grid;grid-template-columns:52px 1fr;gap:3px 10px;font-size:.78rem;color:var(--text);align-items:baseline}
.ctr-lb{font-size:.6rem;text-transform:uppercase;letter-spacing:.05em;color:var(--text-muted);font-weight:700}
.ctr-meta{font-size:.74rem;color:var(--text-muted)}
/* ── monthly cashflow table per group ── */
.cf-details{margin-top:12px}
.cf-details summary{cursor:pointer;font-size:.78rem;font-weight:700;color:var(--text-muted);user-select:none}
.cf-details summary:hover{color:var(--text)}
.cf-wrap{overflow-x:auto;margin-top:10px;border:1px solid var(--border);border-radius:var(--radius-md)}
.cf-table{border-collapse:collapse;font-size:.72rem;min-width:100%}
.cf-table th{border-bottom:1px solid var(--border);padding:6px 9px;font-size:.6rem;text-transform:uppercase;letter-spacing:.04em;color:var(--text-muted);background:var(--surface-2);text-align:right;white-space:nowrap}
.cf-table td{border-bottom:1px solid var(--border);padding:5px 9px;white-space:nowrap}
.cf-table tr:last-child td{border-bottom:none}
.cf-table .num{text-align:right;font-variant-numeric:tabular-nums;color:var(--text)}
.cf-table .pin{position:sticky;left:0;background:var(--surface-2);font-weight:700;font-size:.72rem;color:var(--text);z-index:1;max-width:190px;overflow:hidden;text-overflow:ellipsis}
.cf-table th.pin{z-index:2}
.cf-table .bill{color:var(--success)}
.cf-table .pay{color:var(--group-4)}
.cf-table .tot{font-weight:800;border-left:1px solid var(--border-bright)}
.cf-table .pos{color:var(--success)}
.cf-table .neg{color:var(--danger)}
.cf-table tr.net td{background:rgba(56,189,248,.05);font-weight:700}
.ctr-kidlist{margin:8px 0 0;padding:0;list-style:none;font-size:.82rem}
.ctr-kidlist li{display:flex;justify-content:space-between;gap:10px;padding:5px 0;border-bottom:1px dashed var(--border)}
.ctr-kidlist li b{font-variant-numeric:tabular-nums;color:var(--success);white-space:nowrap}
.ctr-parties{font-size:.78rem;color:var(--text-muted);margin-bottom:12px}
.cmodal-ov{position:fixed;inset:0;background:rgba(0,0,0,.55);display:none;align-items:flex-start;justify-content:center;z-index:200;padding:40px 16px;overflow-y:auto}
.cmodal-ov.show{display:flex}
.cmodal{background:var(--surface);border:1px solid var(--border-bright);border-radius:var(--radius-xl);max-width:720px;width:100%;box-shadow:var(--shadow-lg)}
.cmodal-head{display:flex;align-items:center;gap:10px;padding:14px 18px;border-bottom:1px solid var(--border)}
.cmodal-file{font-size:.76rem;color:var(--text-muted);margin-left:auto;overflow:hidden;text-overflow:ellipsis;white-space:nowrap;max-width:50%}
.cmodal-x{background:none;border:none;color:var(--text-muted);font-size:1.1rem;cursor:pointer;padding:2px 6px}
.cmodal-x:hover{color:var(--danger)}
.cmodal-body{padding:18px}
.ctr-grid{display:grid;grid-template-columns:1fr 1fr;gap:10px 14px}
.ctr-fld{display:flex;flex-direction:column;gap:3px}
.ctr-fld>span{font-size:.68rem;text-transform:uppercase;letter-spacing:.04em;color:var(--text-muted)}
.ctr-fld .slot{width:100%}
/* derived from the monthly sheet — say so instead of inviting a doomed edit */
.ctr-fld .slot[readonly]{background:var(--surface-2);color:var(--text-muted);cursor:not-allowed;border-style:dashed}
.ctr-actions{display:flex;gap:8px;margin-top:14px;flex-wrap:wrap}
.ctr-linkbox,.ctr-linked{margin-top:16px;padding-top:14px;border-top:1px solid var(--border);display:flex;flex-direction:column;gap:8px;font-size:.82rem}
.ppl-row{display:flex;gap:8px;align-items:flex-start;font-size:.82rem;padding:2px 0;cursor:pointer}
.ppl-row input{margin-top:3px}
.ppl-sheet>summary{cursor:pointer;list-style:none;font-size:.82rem}
.ppl-sheet>summary::-webkit-details-marker{display:none}
.ppl-sheet>summary::before{content:"▸ ";color:var(--text-muted)}
.ppl-sheet[open]>summary::before{content:"▾ "}
.ppl-sheet-form{display:flex;flex-direction:column;gap:8px}
.ppl-sheet-form textarea.slot{width:100%;font-size:.8rem}
.ppl-sheet-form .btn{align-self:flex-start}
.ppl-sheet-err{font-size:.78rem;color:var(--danger);background:rgba(248,113,113,.10);
  border:1px solid rgba(248,113,113,.35);border-radius:var(--radius-md);padding:8px 10px;margin:6px 0}
.ppl-extra{font-size:.72rem;color:var(--accent);background:var(--accent-glow);border-radius:var(--radius-full);padding:1px 8px;white-space:nowrap}
.ctr-linked{flex-direction:row;align-items:center;gap:10px}
.ctr-prev{margin-top:14px}
.ctr-prev summary{cursor:pointer;font-size:.8rem;color:var(--text-muted)}
.ctr-prev pre{white-space:pre-wrap;font-size:.72rem;max-height:300px;overflow:auto;background:var(--surface-2,var(--surface));border:1px solid var(--border);border-radius:8px;padding:10px;margin-top:8px}
@media(max-width:768px){
  .ctr-group-body{grid-template-columns:1fr}.ctr-grid{grid-template-columns:1fr}
  /* grid items default to min-width:auto — long parties and the state
     chips would otherwise widen the column past the viewport */
  .ctr-sea-col,.ctr-ven-col,.ctr-card{min-width:0}
  .ctr-rows>span{overflow-wrap:anywhere}
  .eml-form{grid-template-columns:1fr}
  .eml-actions .btn{width:100%;min-height:44px}
  .chg-tab{flex:1;min-height:44px}
  .sch-preview .btn{width:100%;min-height:44px}
  .chg-intake>summary{min-height:44px;display:flex;align-items:center;flex-wrap:wrap}
}
"""

_CTR_JS = """<script>
function closeContract(){document.getElementById('cmodalOv').classList.remove('show');}
function openContract(id){
  fetch('/sow/contract?frag=1&id='+encodeURIComponent(id))
    .then(function(r){return r.text();})
    .then(function(h){document.getElementById('cmodal').innerHTML=h;
      document.getElementById('cmodalOv').classList.add('show');
      // a sheet waiting for review is why the popup opened — don't make the
      // user hunt for it down a long contract form
      var s=document.querySelector('#cmodal .sch-preview');
      if(s)s.scrollIntoView({block:'start'});});
}
function ctrPost(url,obj){
  var b=Object.keys(obj).map(function(k){return k+'='+encodeURIComponent(obj[k]);}).join('&');
  fetch(url,{method:'POST',headers:{'Content-Type':'application/x-www-form-urlencoded'},body:b})
    // follow where the route sent us — picking a column has to come back to
    // the same contract, not drop the popup on the floor
    .then(function(r){closeContract();location.href=r.url||'/sow/contracts';});
}
function ctrSave(id){
  var f=document.querySelector('#cmodal form.ctr-form');
  var b='id='+encodeURIComponent(id);
  f.querySelectorAll('input,select').forEach(function(el){b+='&'+el.name+'='+encodeURIComponent(el.value);});
  fetch('/sow/contract/save',{method:'POST',headers:{'Content-Type':'application/x-www-form-urlencoded'},body:b})
    .then(function(){closeContract();location.href='/sow/contracts';});
}
// Fold state is remembered per block. Persist on the CLICK, not on toggle:
// a <details open> fires a toggle while it is being parsed, which would
// overwrite what the user chose last time.
document.addEventListener('click', function(e){
  var sm = e.target.closest('details.ctr-group > summary'); if(!sm) return;
  var g = sm.parentElement;
  setTimeout(function(){
    try{ localStorage.setItem('ddGroup:' + g.dataset.gid, g.open ? '1' : '0'); }catch(err){}
  }, 0);
});
document.addEventListener('DOMContentLoaded', function(){
  document.querySelectorAll('details.ctr-group').forEach(function(g){
    try{
      var v = localStorage.getItem('ddGroup:' + g.dataset.gid);
      if(v === '1') g.open = true; else if(v === '0') g.open = false;
    }catch(err){}
  });
});
// a card dragged onto a folded block opens it instead of falling through
document.addEventListener('dragover', function(e){
  var g = e.target.closest('details.ctr-group');
  if(g && !g.open) g.open = true;
});
function pplSheet(e,id){
  e.preventDefault();
  var f=e.target, fd=new FormData(f);
  if(!(fd.get('text')||'').trim() && !(fd.get('file')&&fd.get('file').name)){
    alert('Attach a sheet or paste the rows.'); return false;
  }
  fd.append('id', id);
  var b=f.querySelector('button[type=submit]');
  if(b){b.disabled=true;b.textContent='📥 Reading…';}
  fetch('/sow/contract/people_upload',{method:'POST',body:fd})
    .then(function(r){ location.href = r.url || '/sow/contracts?newc='+id; })
    .catch(function(){ if(b){b.disabled=false;b.textContent='📥 Read sheet';} });
  return false;
}
function ctrPeopleSave(id){
  var box=document.getElementById('pplBox');if(!box)return;
  var sel=[];
  box.querySelectorAll('input[type=checkbox]:checked').forEach(function(cb){sel.push(cb.value);});
  var b='id='+encodeURIComponent(id)+'&sel='+encodeURIComponent(sel.join('\\n'));
  fetch('/sow/contract/people_save',{method:'POST',headers:{'Content-Type':'application/x-www-form-urlencoded'},body:b})
    .then(function(){location.href='/sow/contracts?newc='+encodeURIComponent(id);});
}
function ctrAssign(id){
  var sel=document.getElementById('lnk_'+id);
  if(!sel||!sel.value)return;
  ctrPost('/sow/contract/assign',{vendor:id,sea:sel.value});
}
function ctrReparse(id,btn){
  if(btn){btn.disabled=true;btn.textContent='🪄 Reading…';}
  fetch('/sow/contract/reparse',{method:'POST',headers:{'Content-Type':'application/x-www-form-urlencoded'},body:'id='+encodeURIComponent(id)})
    .then(function(){location.href='/sow/contracts?newc='+encodeURIComponent(id);});
}
function chgTab(btn,pane){
  var box=btn.closest('.chg-intake'); if(!box)return;
  box.querySelectorAll('.chg-tab').forEach(function(b){b.classList.toggle('is-on',b===btn);});
  box.querySelectorAll('.chg-pane').forEach(function(f){f.hidden=(f.dataset.pane!==pane);});
}
function chgSubmit(e,url){
  e.preventDefault();
  var f=e.target, fd=new FormData(f), b=f.querySelector('button[type=submit]');
  var hasText=(fd.get('text')||'').trim(), hasFile=(fd.get('file')&&fd.get('file').name);
  if(f.dataset.pane==='sch'){
    if(!hasText && !hasFile){ alert('Attach the schedule or paste the months.'); return false; }
  }else if(!hasText && !hasFile && !(fd.get('amount')||'').trim() && !(fd.get('note')||'').trim()){
    alert('Paste the email, attach it, or type what changed.'); return false;
  }
  var was=b?b.textContent:'';
  if(b){b.disabled=true;b.textContent='⏳ Reading…';}
  fetch(url,{method:'POST',body:fd})
    .then(function(r){ location.href = r.url || '/sow/contracts'; })
    .catch(function(){ if(b){b.disabled=false;b.textContent=was;} });
  return false;
}
function ctrUpload(files){
  if(!files||!files.length)return;
  var fd=new FormData(); fd.append('file', files[0]);
  fetch('/sow/contract/upload',{method:'POST',body:fd})
    .then(function(r){ location.href = r.url || '/sow'; });
}
var ctrDragId=null;
document.addEventListener('dragstart',function(e){
  var card=e.target.closest('.ctr-card[draggable="true"]');
  if(!card){return;}
  ctrDragId=card.getAttribute('data-cid');
  e.dataTransfer.effectAllowed='move';
  try{e.dataTransfer.setData('text/plain',ctrDragId);}catch(err){}
  card.classList.add('dragging');
});
document.addEventListener('dragend',function(e){
  var card=e.target.closest('.ctr-card'); if(card)card.classList.remove('dragging');
  ctrDragId=null;
});
document.addEventListener('dragover',function(e){
  var t=e.target.closest('[data-seadrop],[data-filedrop]');
  if(t){e.preventDefault(); t.classList.add('drag-over');}
});
document.addEventListener('dragleave',function(e){
  var t=e.target.closest('[data-seadrop],[data-filedrop]');
  if(t && !t.contains(e.relatedTarget))t.classList.remove('drag-over');
});
document.addEventListener('drop',function(e){
  var fzone=e.target.closest('[data-filedrop]');
  if(fzone){e.preventDefault(); fzone.classList.remove('drag-over');
    if(e.dataTransfer.files&&e.dataTransfer.files.length)ctrUpload(e.dataTransfer.files);
    return;}
  var t=e.target.closest('[data-seadrop]');
  if(!t)return;
  e.preventDefault(); t.classList.remove('drag-over');
  if(e.dataTransfer.files&&e.dataTransfer.files.length){ctrUpload(e.dataTransfer.files); return;}
  var vid=ctrDragId||(e.dataTransfer&&e.dataTransfer.getData('text/plain'));
  if(!vid)return;
  ctrPost('/sow/contract/assign',{vendor:vid,sea:t.getAttribute('data-sea')||''});
});
document.addEventListener('DOMContentLoaded',function(){
  var ov=document.getElementById('cmodalOv');
  if(ov)ov.addEventListener('click',function(e){if(e.target===ov)closeContract();});
  var dz=document.getElementById('ctrDrop'), fi=document.getElementById('ctrFile');
  if(dz&&fi){
    dz.addEventListener('click',function(){fi.click();});
    fi.addEventListener('change',function(){ctrUpload(fi.files);});
  }
  var m=location.search.match(/[?&]newc=([^&]+)/);
  if(m)openContract(decodeURIComponent(m[1]));
});
</script>"""


def handle(method, path, body, ctx):
    user = ctx.get("user", "guest")

    if method == "GET" and path == "/sow":
        return ("html", _render_landing(user))

    if method == "GET" and path == "/sow/contracts":
        return ("html", _render_contracts_page(user))

    if method == "GET" and path == "/sow/docs":
        return ("html", _render_docs_page(user))

    # ── uploaded contracts ────────────────────────────────────────────────
    if method == "GET" and path == "/sow/contract":
        data = _load(user)
        return ("html", _render_contract_frag(user, data, _f(body, "id")))

    if method == "POST" and path == "/sow/contract/upload":
        raw = body.get("__raw__") or body.get("__raw_handler__")
        files = _read_uploaded_files(raw) if raw else []
        if not files:
            return ("redirect", "/sow")
        fn, content, _mime = files[0]
        fn = _safe_filename(fn)
        ext = _safe_ext(fn)
        text = _extract_text(content, ext)
        fields = _extract_fields_best(text)
        cid = uuid.uuid4().hex[:8]
        data = _load(user)
        rec = {"id": cid, "filename": fn, "ext": ext,
               "uploaded": datetime.now().isoformat(timespec="seconds"),
               "linked_id": None}
        rec.update({k: fields.get(k, "") for k in
                    ("client", "agency", "vendor", "amount",
                     "period_start", "period_end", "project_name", "side")})
        if fields.get("people"):
            rec["people_pending"] = fields["people"]
        try:
            os.makedirs(_contracts_dir(user), exist_ok=True)
            with open(_contract_file_path(user, rec), "wb") as fp:
                fp.write(content)
            _store_contract_text(user, rec, text)
        except OSError:
            return ("redirect", "/sow/contracts")
        _auto_register_vendor(data, rec.get("vendor"))
        data.setdefault("contracts", []).append(rec)
        _save(user, data)
        return ("redirect", f"/sow/contracts?newc={cid}")

    if method == "POST" and path == "/sow/contract/email":
        # A change agreed over email, not in a signed amendment (강프로
        # 2026-07-27). Stored as a normal amendment record so the effective-date
        # override and the cashflow treat it exactly like a papered one — the
        # ✉️ source stays visible on the card.
        raw = body.get("__raw__") or body.get("__raw_handler__")
        fields, files = _read_multipart(raw) if raw else ({}, [])
        if not fields and not files:
            fields = {k: _f(body, k) for k in
                      ("target", "text", "amount", "effective", "end", "name", "note")}
        data = _load(user)
        target = _contract_by_id(data, (fields.get("target") or "").strip())
        if not target:
            return ("redirect", "/sow/contracts")
        meta, fn, ext, content = {}, "", "", None
        body_text = (fields.get("text") or "").strip()
        if files:
            fn, content, _mime = files[0]
            fn = _safe_filename(fn)
            ext = _safe_ext(fn)
            meta, file_body = _email_parts(content, ext)
            body_text = "\n\n".join(x for x in (body_text, file_body) if x.strip())
        if not meta:
            meta = _sniff_mail_headers(body_text)
        text = _email_as_text(meta, body_text)
        ai = _extract_email_change(text) if text.strip() else {}
        cid = uuid.uuid4().hex[:8]
        rec = {"id": cid, "source": "email", "amends_id": target["id"],
               "side": target.get("side") or "sea",
               "client": target.get("client") or "",
               "agency": target.get("agency") or "",
               "vendor": target.get("vendor") or "",
               "filename": ((fields.get("name") or "").strip() or fn
                            or meta.get("subject") or "Email change"),
               "ext": ext or "txt", "linked_id": None,
               "has_file": bool(content),
               "uploaded": datetime.now().isoformat(timespec="seconds"),
               "confirmed": False,
               "email_meta": {k: v for k, v in meta.items() if v},
               # the deal name is inherited from the contract this change sits
               # under — the intake lives inside its block, so it is never
               # retyped (강프로 2026-07-27); the reference above names the change
               "project_name": (fields.get("project") or ai.get("project_name")
                                or target.get("project_name") or ""),
               # typed values win over the AI read — the user is looking at the
               # mail, the model is looking at a forward chain
               "amount": (fields.get("amount") or ai.get("amount") or "").strip(),
               "period_start": (fields.get("effective") or ai.get("period_start") or "").strip(),
               "period_end": (fields.get("end") or ai.get("period_end") or "").strip(),
               "change_note": (fields.get("note") or ai.get("change_summary") or "").strip()}
        if ai and ai.get("is_change") is False and not fields.get("amount"):
            rec["change_note"] = (rec["change_note"] or
                                  "No contractual change found in this email — "
                                  "check before confirming.")
        try:
            os.makedirs(_contracts_dir(user), exist_ok=True)
            if content:
                with open(_contract_file_path(user, rec), "wb") as fp:
                    fp.write(content)
            _store_contract_text(user, rec, text)
        except OSError:
            pass
        data.setdefault("contracts", []).append(rec)
        _save(user, data)
        return ("redirect", f"/sow/contracts?newc={cid}")

    if method == "POST" and path == "/sow/contract/schedule":
        # A change that arrives as monthly figures and nothing else (강프로
        # 2026-07-27) — no document, often no name. Stored as an amendment so
        # the effective-date override treats it like any other, except that it
        # bills the sheet's own months instead of an even spread.
        raw = body.get("__raw__") or body.get("__raw_handler__")
        fields, files = _read_multipart(raw) if raw else ({}, [])
        if not fields and not files:
            fields = {k: _f(body, k) for k in ("target", "text", "name", "note")}
        data = _load(user)
        target = _contract_by_id(data, (fields.get("target") or "").strip())
        if not target:
            return ("redirect", "/sow/contracts")
        rows, src = [], ""
        if files:
            fn, content, _mime = files[0]
            try:
                rows = _sheet_rows(content, _sheet_ext(fn))
                src = _safe_filename(fn)
            except Exception:
                rows = []
        elif (fields.get("text") or "").strip():
            rows = _text_rows(fields["text"])
            src = "pasted range"
        read = _month_table_read(rows) if rows else None
        if not read or not read["columns"]:
            people, _n = _people_from_table(rows) if rows else ([], "")
            target["schedule_error"] = (
                f"That looks like a team sheet ({len(people)} people), not a "
                f"billing schedule — upload it from the contract itself, under "
                f"“Upload a sheet for this contract”." if people else
                "No monthly figures found — the sheet needs month labels "
                "(Jan-26, January 2026, 2026-01) with an amount against each.")
            _save(user, data)
            return ("redirect", f"/sow/contracts?newc={target['id']}")
        target.pop("schedule_error", None)
        # Nothing is written into the cashflow yet: a real sheet often has more
        # than one money column (Original vs Adjusted) and repeated months, so
        # what was read is staged for review first (강프로 2026-07-28).
        target["schedule_preview"] = dict(
            read, src=src, name=(fields.get("name") or "").strip(),
            note_text=(fields.get("note") or "").strip())
        _save(user, data)
        return ("redirect", f"/sow/contracts?newc={target['id']}")

    if method == "POST" and path == "/sow/contract/schedule_pick":
        data = _load(user)
        c = _contract_by_id(data, _f(body, "id"))
        prev = (c or {}).get("schedule_preview")
        if prev and any(col["key"] == _f(body, "col") for col in prev["columns"]):
            prev["chosen"] = _f(body, "col")
            _save(user, data)
        return ("redirect", f"/sow/contracts?newc={_f(body, 'id')}")

    if method == "POST" and path == "/sow/contract/schedule_discard":
        data = _load(user)
        c = _contract_by_id(data, _f(body, "id"))
        if c:
            c.pop("schedule_preview", None)
            _save(user, data)
        return ("redirect", f"/sow/contracts?newc={_f(body, 'id')}")

    if method == "POST" and path == "/sow/contract/schedule_apply":
        data = _load(user)
        target = _contract_by_id(data, _f(body, "id"))
        prev = (target or {}).get("schedule_preview")
        if not prev or not prev.get("columns"):
            return ("redirect", "/sow/contracts")
        col = next((x for x in prev["columns"] if x["key"] == prev.get("chosen")),
                   prev["columns"][0])
        months = {_parse_month_label(k): v for k, v in col["months"].items()}
        months = {k: v for k, v in months.items() if k}
        if not months:
            return ("redirect", "/sow/contracts")
        total, p_start, p_end = _schedule_summary(months)
        span = (f'{date(*sorted(months)[0], 1).strftime("%b %Y")} – '
                f'{date(*sorted(months)[-1], 1).strftime("%b %Y")}')
        src_bits = [x for x in (prev.get("src"), col["label"] if len(prev["columns"]) > 1
                                else "", prev.get("note")) if x]
        cid = uuid.uuid4().hex[:8]
        rec = {"id": cid, "source": "schedule", "amends_id": target["id"],
               "side": target.get("side") or "sea",
               "client": target.get("client") or "",
               "agency": target.get("agency") or "",
               "vendor": target.get("vendor") or "",
               "project_name": target.get("project_name") or "",
               "filename": (prev.get("name") or f"Monthly update · {span}"),
               "ext": "txt", "linked_id": None, "has_file": False,
               "uploaded": datetime.now().isoformat(timespec="seconds"),
               "confirmed": False,
               "month_amounts": {f"{y:04d}-{m:02d}": round(v, 2)
                                 for (y, m), v in months.items()},
               "schedule_note": " · ".join(src_bits),
               "amount": _money(total),
               "period_start": p_start, "period_end": p_end,
               "change_note": prev.get("note_text") or ""}
        try:
            _store_contract_text(user, rec, "\n".join(
                f'{date(y, m, 1).strftime("%b %Y")}\t{v:,.2f}'
                for (y, m), v in sorted(months.items())))
        except OSError:
            pass
        target.pop("schedule_preview", None)
        data.setdefault("contracts", []).append(rec)
        _save(user, data)
        return ("redirect", f"/sow/contracts?newc={cid}")

    if method == "POST" and path == "/sow/contract/schedule_clear":
        data = _load(user)
        c = _contract_by_id(data, _f(body, "id"))
        if c:
            c.pop("month_amounts", None)
            c.pop("schedule_note", None)
            _save(user, data)
        return ("redirect", "/sow/contracts")

    if method == "POST" and path == "/sow/contract/reparse":
        data = _load(user)
        c = _contract_by_id(data, _f(body, "id"))
        if c:
            # re-read the ORIGINAL file so extractor improvements reach already
            # uploaded contracts (the stored text may be an old under-read)
            text = _contract_text(user, c)
            try:
                with open(_contract_file_path(user, c), "rb") as fp:
                    text = _extract_text(fp.read(), c.get("ext", "bin"))
                _store_contract_text(user, c, text)
            except OSError:
                pass
            if c.get("source") == "email":
                # an email is read by the change extractor, not the contract one
                ai = _extract_email_change(text)
                for k in ("amount", "period_start", "period_end", "project_name"):
                    if ai.get(k):
                        c[k] = ai[k]
                if ai.get("change_summary"):
                    c["change_note"] = ai["change_summary"]
                c["confirmed"] = False
                _save(user, data)
                return ("redirect", f"/sow/contracts?newc={_f(body, 'id')}")
            fields = _extract_fields_best(text)
            for k in ("client", "agency", "vendor", "amount",
                      "period_start", "period_end", "project_name", "side"):
                c[k] = fields.get(k, c.get(k, ""))
            if fields.get("people"):
                c["people_pending"] = fields["people"]
            _auto_register_vendor(data, c.get("vendor"))
            c["confirmed"] = False  # re-read fields need a fresh confirmation
            _save(user, data)
        return ("redirect", f"/sow/contracts?newc={_f(body, 'id')}")

    if method == "POST" and path == "/sow/contract/save":
        data = _load(user)
        c = _contract_by_id(data, _f(body, "id"))
        if c:
            # only fields the (side-aware) form actually submitted — the other
            # side's party box isn't rendered and must not be wiped
            for k in ("client", "vendor", "amount",
                      "period_start", "period_end", "project_name"):
                if k in body:
                    c[k] = _f(body, k)
            c["side"] = "vendor" if _f(body, "side") == "vendor" else "sea"
            c["agency"] = CHEIL_ENTITY
            if c["side"] == "vendor":
                c["client"] = CHEIL_ENTITY
            _save(user, data)
        return ("redirect", "/sow/contracts")

    if method == "POST" and path == "/sow/contract/assign":
        # Align a vendor contract under a SEA↔Cheil parent (1:many). An empty
        # or invalid `sea` unlinks. `vendor` may currently be tagged either
        # side; assigning it under a SEA also forces its side to vendor.
        data = _load(user)
        v = _contract_by_id(data, _f(body, "vendor"))
        sea = _contract_by_id(data, _f(body, "sea"))
        if v:
            if sea and sea.get("side") == "sea" and sea["id"] != v["id"]:
                v["side"] = "vendor"
                v["linked_id"] = sea["id"]
            else:
                v["linked_id"] = None
            _save(user, data)
        return ("redirect", "/sow/contracts")

    if method == "POST" and path == "/sow/contract/confirm":
        data = _load(user)
        c = _contract_by_id(data, _f(body, "id"))
        if c:
            c["confirmed"] = not c.get("confirmed")
            c["confirmed_at"] = (datetime.now().isoformat(timespec="seconds")
                                 if c["confirmed"] else None)
            _save(user, data)
        return ("redirect", "/sow/contracts")

    if method == "POST" and path == "/sow/contract/selfdeliver":
        data = _load(user)
        c = _contract_by_id(data, _f(body, "id"))
        if c:
            c["self_delivered"] = not c.get("self_delivered")
            _save(user, data)
        return ("redirect", "/sow/contracts")

    if method == "POST" and path == "/sow/contract/cancel":
        data = _load(user)
        c = _contract_by_id(data, _f(body, "id"))
        if c:
            c["cancelled"] = not c.get("cancelled")
            _save(user, data)
        return ("redirect", "/sow/contracts")

    if method == "POST" and path == "/sow/contract/amends":
        data = _load(user)
        c = _contract_by_id(data, _f(body, "id"))
        if c:
            pid = _f(body, "prev")
            prev = _contract_by_id(data, pid) if pid else None
            # guard: valid target, not itself, no 2-cycle
            if prev and prev["id"] != c["id"] and prev.get("amends_id") != c["id"]:
                c["amends_id"] = prev["id"]
            else:
                c["amends_id"] = None
            _save(user, data)
        return ("redirect", "/sow/contracts")

    if method == "POST" and path == "/sow/contract/people_upload":
        # One sheet intake against a contract (강프로 2026-07-28). The team for a
        # deal usually arrives as a sheet rather than inside the contract text,
        # and so does a revised monthly billing schedule — the sheet's own
        # contents say which it is, instead of one being rejected out of hand.
        raw = body.get("__raw__") or body.get("__raw_handler__")
        fields, files = _read_multipart(raw) if raw else ({}, [])
        if not fields and not files:
            fields = {k: _f(body, k) for k in ("id", "text")}
        data = _load(user)
        c = _contract_by_id(data, (fields.get("id") or "").strip())
        if not c:
            return ("redirect", "/sow/contracts")
        rows, src = [], ""
        if files:
            fn, content, _mime = files[0]
            try:
                rows = _sheet_rows(content, _sheet_ext(fn))
                src = _safe_filename(fn)
            except Exception:
                rows = []
        elif (fields.get("text") or "").strip():
            rows = _text_rows(fields["text"])
            src = "pasted range"
        found, note = _people_from_table(rows) if rows else ([], "")
        read = _month_table_read(rows) if rows else None
        if found:
            # a roster names people; that wins whenever a name column is there
            c["people_pending"] = _merge_pending(c.get("people_pending"), found)
            c["people_sheet_note"] = note
            c.pop("people_sheet_error", None)
        elif read and read["columns"]:
            c["schedule_preview"] = dict(read, src=src, name="", note_text="")
            c["people_sheet_note"] = (
                f"read as a monthly billing schedule — {read['note']}")
            c.pop("people_sheet_error", None)
        else:
            c["people_sheet_error"] = (
                "Neither a team nor a schedule could be read — a team sheet "
                "needs a Resource/Name column, a monthly schedule needs month "
                "labels (Jan-26, 2026-01) with an amount against each.")
        _save(user, data)
        return ("redirect", f"/sow/contracts?newc={c['id']}")

    if method == "POST" and path == "/sow/contract/people_save":
        data = _load(user)
        c = _contract_by_id(data, _f(body, "id"))
        if c:
            pending = c.get("people_pending") or []
            aff = (c.get("vendor") if c.get("side") == "vendor" else "Cheil") or "Cheil"
            for tok in _f(body, "sel").split("\n"):
                tok = tok.strip()
                if not tok.isdigit() or int(tok) >= len(pending):
                    continue
                ext = pending[int(tok)]
                name = (ext.get("name") or "").strip()
                if not name:
                    continue
                p = _find_person(data, name)
                if not p:
                    p = _migrate_person({"id": uuid.uuid4().hex[:8], "name": name,
                                         "sell_hr": ""})
                    p["affiliation"] = aff
                    data.setdefault("people", []).append(p)
                _apply_extracted_person(p, ext, c)
                lc = p.setdefault("linked_contracts", [])
                if c["id"] not in lc:
                    lc.append(c["id"])
            c.pop("people_pending", None)
            _save(user, data)
        return ("redirect", "/sow/contracts")

    if method == "POST" and path == "/sow/contract/people_dismiss":
        data = _load(user)
        c = _contract_by_id(data, _f(body, "id"))
        if c:
            c.pop("people_pending", None)
            _save(user, data)
        return ("redirect", "/sow/contracts")

    if method == "POST" and path == "/sow/contract/unlink":
        data = _load(user)
        a = _contract_by_id(data, _f(body, "id"))
        if a:
            a["linked_id"] = None
            _save(user, data)
        return ("redirect", "/sow/contracts")

    if method == "POST" and path == "/sow/contract/delete":
        data = _load(user)
        c = _contract_by_id(data, _f(body, "id"))
        if c:
            # deleting a SEA parent orphans its vendor children
            for other in data["contracts"]:
                if other.get("linked_id") == c["id"]:
                    other["linked_id"] = None
            for p in (_contract_file_path(user, c), _contract_text_path(user, c)):
                try:
                    os.remove(p)
                except OSError:
                    pass
            data["contracts"] = [x for x in data["contracts"] if x["id"] != c["id"]]
            _save(user, data)
        return ("redirect", "/sow/contracts")

    if method == "GET" and path == "/sow/contract/file":
        data = _load(user)
        c = _contract_by_id(data, _f(body, "id"))
        if not c:
            return ("redirect", "/sow")
        try:
            with open(_contract_file_path(user, c), "rb") as fp:
                blob = fp.read()
        except OSError:
            return ("redirect", "/sow")
        mime = _CONTRACT_MIME.get(c.get("ext"), "application/octet-stream")
        return ("file_inline", blob, mime, c.get("filename") or f"{c['id']}.{c.get('ext')}")

    if method == "GET" and path.startswith("/sow/types"):
        d = _f(body, "dir")
        return ("html", _render_types(user, "agy" if d == "agy" else "sea"))

    if method == "GET" and path == "/sow/vendors":
        return ("html", _render_vendors(user, saved=_f(body, "saved") == "1"))

    if method == "GET" and path == "/sow/people":
        return ("html", _render_people(user, saved=_f(body, "saved") == "1",
                                       view=_f(body, "view") or "roster"))

    if method == "GET" and path == "/sow/person":
        pid = _f(body, "id")
        data = _load(user)
        person = next((p for p in data["people"] if p["id"] == pid), None)
        return ("html", _render_person_detail(user, person or {},
                                              saved=_f(body, "saved") == "1"))

    if method == "POST" and path == "/sow/person/save":
        data = _load(user)
        pid = _f(body, "id")
        name = _f(body, "name")
        if not name:
            return ("redirect", "/sow/people")
        cur = next((p for p in data["people"] if p["id"] == pid), None)
        linked = body.get("linked_sows") or []
        if not isinstance(linked, list):
            linked = [linked]
        rec = _migrate_person({"id": pid or uuid.uuid4().hex[:8], "name": name,
                               "sell_hr": ""})
        for k in ("affiliation", "role_title", "project", "location",
                  "sell_hr", "sell_mo", "client_duration", "client_budget", "client_po",
                  "cost_hr", "cost_mo", "partner_duration", "partner_cost", "partner_po",
                  "salary_mo", "cheil_since", "salary_oh", "ebita",
                  "email_cheil", "email_samsung", "pc", "svpn"):
            rec[k] = _f(body, k)
        rec["affiliation"] = rec["affiliation"] or "Cheil"
        rec["linked_sows"] = [str(x) for x in linked if x]
        linked_c = body.get("linked_contracts") or []
        if not isinstance(linked_c, list):
            linked_c = [linked_c]
        # a vendor contract always drags its SEA parent along (강프로 2026-07-24)
        expanded = []
        for x in linked_c:
            x = str(x)
            if not x:
                continue
            if x not in expanded:
                expanded.append(x)
            c0 = _contract_by_id(data, x)
            pid0 = (c0 or {}).get("linked_id")
            if pid0 and pid0 not in expanded:
                expanded.append(pid0)
        pick = _f(body, "project_pick")
        if pick.startswith("ctr:"):
            c0 = _contract_by_id(data, pick[4:])
            if c0:
                rec["project"] = (c0.get("project_name") or "").strip()
                for cid2 in (c0["id"], c0.get("linked_id")):
                    if cid2 and cid2 not in expanded:
                        expanded.append(cid2)
        elif pick == "__keep__":
            rec["project"] = (cur or {}).get("project") or ""
        else:
            rec["project"] = ""
        rec["linked_contracts"] = expanded
        if not rec.get("project") and expanded:
            first = _contract_by_id(data, expanded[0])
            if first and first.get("project_name"):
                rec["project"] = first["project_name"]
        if cur:
            data["people"][data["people"].index(cur)] = rec
        else:
            data["people"].append(rec)
        _save(user, data)
        return ("redirect", f"/sow/person?id={rec['id']}&saved=1")

    if method == "POST" and path == "/sow/person/duplicate":
        # a second developer on the same deal differs by name and login, not by
        # rate or project — clone the deal side, blank the personal side
        data = _load(user)
        src = next((p for p in data["people"] if p["id"] == _f(body, "id")), None)
        if not src:
            return ("redirect", "/sow/people")
        new = dict(src)
        new["id"] = uuid.uuid4().hex[:8]
        new["name"] = f'{src.get("name") or "Person"} (copy)'
        for k in ("email_cheil", "email_samsung", "pc", "svpn", "cheil_since",
                  "salary_mo", "salary_oh", "ebita"):
            new[k] = ""
        new["linked_contracts"] = list(src.get("linked_contracts") or [])
        new["linked_sows"] = []
        data["people"].append(_migrate_person(new))
        _save(user, data)
        return ("redirect", f"/sow/person?id={new['id']}")

    if method == "POST" and path == "/sow/person/cell":
        # inline roster edit — one cell, no page reload (강프로 2026-07-27)
        data = _load(user)
        person = next((p for p in data["people"] if p["id"] == _f(body, "id")), None)
        field = _f(body, "field")
        if not person or field not in _PP_EDITABLE:
            return ("json", {"ok": False})
        person[field] = _f(body, "value").strip()
        _save(user, data)
        # echo back the canonical rendering so the cell shows "$6,720", not
        # whatever shape it was typed in
        disp = (_money_input(person[field]) if field in _PP_MONEY else person[field])
        return ("json", {"ok": True, "value": person[field], "display": disp})

    if method == "POST" and path == "/sow/person/delete":
        data = _load(user)
        pid = _f(body, "id")
        person = next((p for p in data["people"] if p["id"] == pid), None)
        if person and _person_doc_count(data, person) == 0:
            data["people"] = [p for p in data["people"] if p["id"] != pid]
            _save(user, data)
        return ("redirect", "/sow/people")

    if method == "GET" and path.startswith("/sow/xlsx"):
        sid = _f(body, "id")
        data = _load(user)
        sow = next((s for s in data["sows"] if s["id"] == sid), None)
        if not sow or TYPES.get(_sow_type(sow), {}).get("kind") != "est":
            return ("redirect", "/sow")
        blob = _build_est_xlsx(sow)
        safe = "".join(c if c.isalnum() or c in " ._-" else "_" for c in (sow.get("title") or "Estimation"))[:60].strip()
        return ("file_inline", blob,
                "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                f"Cheil_Estimation_{safe}.xlsx")

    if method == "POST" and path == "/sow/vendor/save":
        data = _load(user)
        vid = _f(body, "id")
        name = _f(body, "name")
        if not name:
            return ("redirect", "/sow/vendors")
        rec = {"id": vid or uuid.uuid4().hex[:8], "name": name,
               "entity_line": _f(body, "entity_line"), "msa_date": _f(body, "msa_date")}
        cur = next((v for v in data["vendors"] if v["id"] == vid), None)
        if cur:
            data["vendors"][data["vendors"].index(cur)] = rec
        else:
            data["vendors"].append(rec)
        _save(user, data)
        return ("redirect", "/sow/vendors?saved=1")

    if method == "POST" and path == "/sow/vendor/delete":
        data = _load(user)
        vid = _f(body, "id")
        if any(s.get("vendor_id") == vid for s in data["sows"]):
            return ("redirect", "/sow/vendors")
        data["vendors"] = [v for v in data["vendors"] if v["id"] != vid]
        _save(user, data)
        return ("redirect", "/sow/vendors")

    if method == "GET" and path == "/sow/asset/logo":
        logo = os.path.join(_ASSETS, "cheil_logo.png")
        if os.path.exists(logo):
            return ("binary", open(logo, "rb").read(), "image/png")
        return ("html", "", )

    def _editor_for(kind):
        return {"sow": _render_doc_editor, "est": _render_est_editor}.get(kind, _render_agreement_editor)

    if method == "GET" and path.startswith("/sow/new"):
        tkey = _f(body, "type")
        if tkey not in TYPES:
            return ("redirect", "/sow")
        return ("html", _editor_for(TYPES[tkey]["kind"])(user, {}, tkey))

    if method == "GET" and path.startswith("/sow/edit"):
        # GET dispatch strips the query string; params arrive as the body dict.
        sid = _f(body, "id")
        data = _load(user)
        sow = next((s for s in data["sows"] if s["id"] == sid), None)
        if not sow:
            return ("redirect", "/sow")
        tkey = _sow_type(sow)
        return ("html", _editor_for(TYPES[tkey]["kind"])(user, sow, tkey, saved=_f(body, "saved") == "1"))

    if method == "POST" and path == "/sow/save":
        data = _load(user)
        sid = _f(body, "id") or uuid.uuid4().hex[:10]
        tkey = _f(body, "type")
        t = TYPES.get(tkey) or TYPES["sea_sow"]
        try:
            resources = json.loads(_f(body, "resources_json") or "[]")
            assert isinstance(resources, list)
        except Exception:
            resources = []
        try:
            overrides = json.loads(_f(body, "schedule_overrides") or "{}")
            assert isinstance(overrides, dict)
            overrides = {str(k): float(v) for k, v in overrides.items()}
        except Exception:
            overrides = {}
        # Inline vendor registration rides along with the save.
        vendor_id = _f(body, "vendor_id")
        v_name = _f(body, "v_name")
        if v_name:
            vendor_id = uuid.uuid4().hex[:8]
            data["vendors"].append({
                "id": vendor_id, "name": v_name,
                "entity_line": _f(body, "v_entity"),
                "msa_date": _f(body, "v_msa"),
            })
        sow = next((s for s in data["sows"] if s["id"] == sid), None)
        if t["kind"] == "est":
            try:
                est_rows = json.loads(_f(body, "rows_json") or "[]")
                assert isinstance(est_rows, list)
            except Exception:
                est_rows = []
            rec = {
                "id": sid, "type": tkey, "direction": t["dir"], "kind": "est",
                "title": _f(body, "title"), "project_name": _f(body, "project_name"),
                "period_label": _f(body, "period_label"),
                "months": _f(body, "months") or "0",
                "rows": est_rows,
                "created": (sow or {}).get("created") or datetime.now().isoformat(),
                "updated": datetime.now().isoformat(),
            }
            if sow:
                data["sows"][data["sows"].index(sow)] = rec
            else:
                data["sows"].append(rec)
            _save(user, data)
            return ("redirect", f"/sow/edit?id={sid}&saved=1")
        if t["kind"] != "sow":
            vend = next((v for v in data["vendors"] if v["id"] == vendor_id), None)
            label = "MSA" if t["kind"] == "msa" else "NDA"
            rec = {
                "id": sid, "type": tkey, "direction": t["dir"], "kind": t["kind"],
                "date": _f(body, "date"), "vendor_id": vendor_id,
                "title": f"{label} — {(vend or {}).get('name') or 'vendor TBD'}",
                "created": (sow or {}).get("created") or datetime.now().isoformat(),
                "updated": datetime.now().isoformat(),
            }
            if sow:
                data["sows"][data["sows"].index(sow)] = rec
            else:
                data["sows"].append(rec)
            _save(user, data)
            return ("redirect", f"/sow/edit?id={sid}&saved=1")
        rec = {
            "id": sid,
            "type": tkey if tkey in TYPES else "sea_sow",
            "direction": t["dir"],
            "res_mode": _f(body, "res_mode") if _f(body, "res_mode") in ("monthly", "hourly") else t["mode"],
            "date": _f(body, "date"),
            "title": _f(body, "title"),
            "project_name": _f(body, "project_name"),
            "prepared_by": _f(body, "prepared_by"),
            "prepared_for": _f(body, "prepared_for"),
            "vendor_id": vendor_id,
            "exec_summary": _f(body, "exec_summary"),
            "deliverables": _f(body, "deliverables"),
            "stk_c_name": _f(body, "stk_c_name"), "stk_c_email": _f(body, "stk_c_email"),
            "stk_c_loc": _f(body, "stk_c_loc"),
            "stk_a_name": _f(body, "stk_a_name"), "stk_a_email": _f(body, "stk_a_email"),
            "stk_a_loc": _f(body, "stk_a_loc"),
            "start": _f(body, "start"), "end": _f(body, "end"),
            "invoice_rule": _f(body, "invoice_rule") or "next_first",
            "resources": resources,
            "schedule_overrides": overrides,
            "created": (sow or {}).get("created") or datetime.now().isoformat(),
            "updated": datetime.now().isoformat(),
        }
        if sow:
            data["sows"][data["sows"].index(sow)] = rec
        else:
            data["sows"].append(rec)
        _save(user, data)
        return ("redirect", f"/sow/edit?id={sid}&saved=1")

    if method == "POST" and path == "/sow/delete":
        data = _load(user)
        sid = _f(body, "id")
        data["sows"] = [s for s in data["sows"] if s["id"] != sid]
        _save(user, data)
        return ("redirect", "/sow")

    if method == "GET" and path.startswith("/sow/docx"):
        sid = _f(body, "id")
        data = _load(user)
        sow = next((s for s in data["sows"] if s["id"] == sid), None)
        if not sow:
            return ("redirect", "/sow")
        vendor = next((v for v in data["vendors"] if v["id"] == sow.get("vendor_id")), None)
        kind = TYPES[_sow_type(sow)]["kind"]
        builder = {"sow": _build_docx, "msa": _build_msa_docx, "nda": _build_nda_docx}[kind]
        blob = builder(sow, vendor)
        safe = "".join(c if c.isalnum() or c in " ._-" else "_" for c in (sow.get("title") or "SOW"))[:60].strip()
        prefix = {"sow": "Cheil_SOW", "msa": "Cheil_MSA", "nda": "Cheil_NDA"}[kind]
        fname = f"{prefix}_{safe}_{sow.get('date') or date.today().isoformat()}.docx"
        return ("file_inline", blob,
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document", fname)

    return ("html", "<h2>404 Not Found</h2>")
